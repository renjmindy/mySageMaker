"""
Generate: NLP OpenMed Sentiment Analysis API Integration Guide
Target audience: Junior software engineer (new graduate)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY       = RGBColor(0x1A, 0x3A, 0x5C)
TEAL       = RGBColor(0x00, 0x7A, 0x87)
DARK_GRAY  = RGBColor(0x33, 0x33, 0x33)
MID_GRAY   = RGBColor(0x66, 0x66, 0x66)
GREEN      = RGBColor(0x1E, 0x7E, 0x34)
ORANGE     = RGBColor(0xD3, 0x6A, 0x00)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
CODE_BG    = RGBColor(0xF4, 0xF6, 0xF8)
HEADER_BG  = RGBColor(0x1A, 0x3A, 0x5C)

# ── Helper: shade a table cell ────────────────────────────────────────────────
def shade_cell(cell, rgb: RGBColor):
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tc_pr.append(shd)

def shade_row(row, rgb: RGBColor):
    for cell in row.cells:
        shade_cell(cell, rgb)

# ── Helper: set paragraph border (callout box) ───────────────────────────────
def add_callout(doc, text, style="tip"):
    icons   = {"tip": "💡 Tip", "warning": "⚠️  Important", "note": "📝 Note"}
    colors  = {"tip": GREEN, "warning": ORANGE, "note": TEAL}
    label   = icons.get(style, "📝 Note")
    color   = colors.get(style, TEAL)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(f"{label}:  ")
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.color.rgb = DARK_GRAY
    run2.font.size = Pt(10)
    return p

# ── Helper: add code block (table with shaded bg) ────────────────────────────
def add_code(doc, code_text, lang=""):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, CODE_BG)
    cell.paragraphs[0].clear()
    lines = code_text.strip().split("\n")
    for i, line in enumerate(lines):
        if i == 0:
            para = cell.paragraphs[0]
        else:
            para = cell.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        run = para.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    doc.add_paragraph()

# ── Helper: styled heading ────────────────────────────────────────────────────
def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(18)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = TEAL
        run.font.size = Pt(14)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def h3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(11.5)
        run.bold = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    return p

def body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_GRAY
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_GRAY
    return p

def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_GRAY
    return p

def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0]
    shade_row(hdr, HEADER_BG)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9.5)
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        if ri % 2 == 1:
            shade_row(row, RGBColor(0xF0, 0xF4, 0xF8))
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9.5)
            run.font.color.rgb = DARK_GRAY
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run("NLP OpenMed Sentiment Analysis")
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = NAVY

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("API Integration Guide")
run2.font.size = Pt(22)
run2.font.bold = True
run2.font.color.rgb = TEAL

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("For Junior Software Engineers")
run3.font.size = Pt(13)
run3.font.color.rgb = MID_GRAY
run3.font.italic = True

doc.add_paragraph()
doc.add_paragraph()

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
for label, val in [("Version", "2.0"), ("Date", "May 2026"), ("Environment", "Local / Docker")]:
    r = p4.add_run(f"{label}:  ")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = NAVY
    r2 = p4.add_run(f"{val}     ")
    r2.font.size = Pt(10)
    r2.font.color.rgb = DARK_GRAY

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "1. Overview")

body(doc, (
    "This guide explains how to integrate with the NLP OpenMed Sentiment Analysis API. "
    "The API takes free-text clinical notes or patient feedback, automatically redacts any "
    "personally identifiable information (PII/PHI) via the OpenMed service, then runs the "
    "text through one of 13 Transformer models to determine its sentiment."
))
body(doc, (
    "For example, if a patient writes \"Dr Smith was incredibly kind and explained my care "
    "plan in full\", the API will return POSITIVE. "
    "If they write \"I waited three hours and was never seen\", it will return NEGATIVE. "
    "The GoEmotions model can also detect fine-grained emotions such as JOY, FEAR, or SADNESS."
))

add_callout(doc, (
    "You do not need to understand how the Transformer models work internally. "
    "This guide focuses entirely on how to call the API and use the results in your application."
), "note")

h2(doc, "1.1  What the API Does")
for item in [
    "Accepts a single patient note or feedback text per request.",
    "Redacts all PII/PHI (names, dates, MRNs, phone numbers, etc.) before analysis — "
    "HIPAA-aware by design.",
    "Runs the redacted text through a chosen NLP preprocessing pipeline (cleaning, "
    "tokenisation, stemming, lemmatisation, NER, POS tagging).",
    "Classifies the text using one of 13 Transformer models and returns the predicted "
    "sentiment, probabilities, and per-word distribution.",
    "Supports 6 fine-tuned healthcare models (cjen1008) and 7 general-purpose pretrained models.",
]:
    bullet(doc, item)

h2(doc, "1.2  Key Terms")
add_table(doc,
    ["Term", "What it means"],
    [
        ["API",          "Application Programming Interface — a way for your code to talk to another service."],
        ["REST API",     "A common style of API that uses standard HTTP methods (GET, POST) and returns JSON data."],
        ["JSON",         "JavaScript Object Notation — a text format for sending structured data. E.g. {\"key\": \"value\"}."],
        ["Endpoint",     "A specific URL you send a request to, like a specific page on a website but for data."],
        ["Request Body", "The JSON data you send with a POST request — this is where your text goes."],
        ["Response",     "The JSON data the API sends back after processing your request."],
        ["NLP",          "Natural Language Processing — AI technology that understands human language."],
        ["Transformer",  "A type of deep-learning model (e.g. BERT, RoBERTa) that reads text and predicts labels."],
        ["Sentiment",    "The overall tone of a piece of text: POSITIVE, NEGATIVE, NEUTRAL, or an emotion."],
        ["PII / PHI",    "Personally Identifiable Information / Protected Health Information — private patient data."],
        ["Redaction",    "Replacing PII/PHI with placeholder tokens (e.g. [NAME], [DATE]) before sending to models."],
        ["Probability",  "A score from 0.0 to 1.0 per label showing how confident the model is. All scores sum to 1.0."],
        ["Lemma",        "The dictionary root form of a word. E.g. 'waiting' → 'wait', 'nurses' → 'nurse'."],
        ["NER",          "Named Entity Recognition — identifying entities like person names, organisations, dates in text."],
        ["Fine-tuned",   "A pretrained model further trained on domain-specific data (here: healthcare feedback)."],
    ],
    col_widths=[3, 13.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. QUICK START
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "2. Quick Start (5 minutes)")

body(doc, "Follow these five steps to spin up the API locally and make your first successful call.")

h2(doc, "Step 1 — Clone and Enter the Project")
add_code(doc, """\
git clone <repo-url>
cd nlp-openmed-sentiment-analysis-finetuned-models""", "bash")

h2(doc, "Step 2 — Create a Virtual Environment and Install Dependencies")
body(doc, "The project requires Python 3.11. Run these commands once:")
add_code(doc, """\
python -m venv .venv
source .venv/bin/activate          # Windows: .venv\\Scripts\\activate
pip install --upgrade pip
pip install -r requirements.txt""", "bash")

add_callout(doc,
    "The first install downloads several large Transformer model weights from HuggingFace. "
    "This may take 5–15 minutes depending on your internet connection. It only happens once.",
    "note"
)

h2(doc, "Step 3 — Set the OpenMed PII Credentials")
body(doc, "The API uses an external OpenMed service to redact PII/PHI before analysis. "
         "Create a file named .env in the project root with these two lines:")
add_code(doc, """\
OPENMED_PII_BASE_URL=https://i6di6fkeqi.execute-api.ap-southeast-2.amazonaws.com/prod
OPENMED_PII_API_KEY=<your-openmed-api-key>""", "bash")

add_callout(doc,
    "Never commit the .env file to Git. It is already listed in .gitignore. "
    "Ask your team lead for the API key value.",
    "warning"
)

h2(doc, "Step 4 — Start the FastAPI Server")
add_code(doc, """\
uvicorn api.main:app --host 0.0.0.0 --port 8000 --reload""", "bash")
body(doc, "You should see output like:")
add_code(doc, """\
INFO:     Uvicorn running on http://0.0.0.0:8000 (Press CTRL+C to quit)
INFO:     Application startup complete.""", "")

add_callout(doc,
    "The interactive API docs are available at http://localhost:8000/docs once the server is running. "
    "You can try every endpoint directly from your browser there.",
    "tip"
)

h2(doc, "Step 5 — Run a Health Check")
body(doc, "Open a second terminal and run:")
add_code(doc, """\
curl -s http://localhost:8000/api/v1/health | python3 -m json.tool""", "bash")
body(doc, "You should see:")
add_code(doc, """\
{
    "status": "ok",
    "models_loaded": []
}""", "json")

add_callout(doc,
    "models_loaded is empty on a fresh start because models are lazy-loaded — they load "
    "the first time they are used. Call GET /api/v1/warmup to pre-load all 13 models at once.",
    "note"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 3. API REFERENCE
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "3. API Reference")

h2(doc, "3.1  Base URL")
add_code(doc, "http://localhost:8000/api/v1", "")
body(doc, "When running inside Docker, the base URL is identical — the container maps port 8000 to host port 8000.")

h2(doc, "3.2  Authentication")
body(doc, (
    "The local API has no API key requirement. All endpoints are open when running locally. "
    "The only external credential is the OpenMed PII key, which is configured in .env and "
    "is used server-side — you never need to include it in your API calls."
))

add_callout(doc,
    "If you deploy this API behind a gateway (e.g. AWS API Gateway, nginx), add authentication "
    "at that layer. Do not expose the API on a public network without authentication.",
    "warning"
)

h2(doc, "3.3  Endpoints")
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/api/v1/health",  "Returns API status and the list of currently loaded models."],
        ["GET",  "/api/v1/warmup",  "Pre-loads all 13 models into memory. Call this once after a cold start."],
        ["POST", "/api/v1/analyze", "Redacts PII then runs sentiment analysis on the submitted text."],
    ],
    col_widths=[2, 4.5, 10]
)

h2(doc, "3.4  GET /api/v1/health")
body(doc, "Returns the API status and which models have been loaded since startup.")
body(doc, "Request:")
add_code(doc, """\
curl -s http://localhost:8000/api/v1/health | python3 -m json.tool""", "bash")
body(doc, "Response (HTTP 200):")
add_code(doc, """\
{
    "status": "ok",
    "models_loaded": ["bert_hc_v2", "emotion"]
}""", "json")

h2(doc, "3.5  GET /api/v1/warmup")
body(doc, (
    "Pre-loads all 13 models into memory. "
    "This avoids the per-model delay on the first analyze call for each model. "
    "Expect this to take 2–5 minutes on first run as weights are downloaded and cached."
))
add_code(doc, """\
curl -s http://localhost:8000/api/v1/warmup | python3 -m json.tool""", "bash")
body(doc, "Response (HTTP 200):")
add_code(doc, """\
{
    "status": "ok",
    "models": {
        "bert_hc_v2":          "loaded in 4.2s",
        "distilroberta_hc_v2": "loaded in 3.1s",
        "distilbert_hc_v2":    "loaded in 2.8s",
        "bert_hc":             "loaded in 4.0s",
        "distilbert_hc":       "loaded in 2.9s",
        "distilroberta_hc":    "loaded in 3.0s",
        "default":             "loaded in 2.5s",
        "roberta":             "loaded in 4.1s",
        "emotion":             "loaded in 3.3s",
        "amazon":              "loaded in 2.6s",
        "twitter":             "loaded in 3.5s",
        "sst2":                "loaded in 4.0s",
        "zeroshot":            "loaded in 8.7s"
    }
}""", "json")

h2(doc, "3.6  POST /api/v1/analyze")
body(doc, (
    "This is the main endpoint. Submit a clinical note or patient feedback text "
    "and receive the redacted text, preprocessing results, and sentiment analysis."
))

h3(doc, "Request Body Fields")
add_table(doc,
    ["Field", "Type", "Required?", "Default", "Description"],
    [
        ["text",       "string", "Yes", "—",           "The text to analyse. Must be 4–300 words."],
        ["model_type", "string", "No",  "\"bert_hc_v2\"", "Which of the 13 models to use. See Section 4 for all valid keys."],
    ],
    col_widths=[2.5, 2.5, 2.5, 3, 6]
)

add_callout(doc,
    "The API processes one text per request. To classify multiple texts, call the endpoint "
    "once per text or wrap calls in a loop (see Section 6.2 for an example).",
    "note"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. THE 13 MODELS
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "4. The 13 Sentiment Models")

body(doc, (
    "The API supports 13 Transformer models split into two groups: 6 fine-tuned healthcare "
    "models trained specifically on clinical and patient feedback text (cjen1008), and 7 "
    "general-purpose pretrained models that bring broader language understanding."
))

h2(doc, "4.1  Fine-Tuned Healthcare Models  (cjen1008)")

body(doc, (
    "These models were fine-tuned on healthcare-specific text and return a three-class "
    "sentiment: NEGATIVE, NEUTRAL, or POSITIVE. They are the recommended choice for "
    "analysing clinical notes and patient experience feedback."
))

add_table(doc,
    ["model_type key", "Display Name", "HuggingFace ID", "Labels", "Recommended For"],
    [
        ["bert_hc_v2",          "BERT Healthcare v2",          "cjen1008/bert-healthcare-sentiment_v2",          "NEGATIVE / NEUTRAL / POSITIVE", "Best overall — production default"],
        ["distilroberta_hc_v2", "DistilRoBERTa Healthcare v2", "cjen1008/distilroberta-healthcare-sentiment_v2", "NEGATIVE / NEUTRAL / POSITIVE", "Faster than BERT v2, high accuracy"],
        ["distilbert_hc_v2",    "DistilBERT Healthcare v2",    "cjen1008/distilbert-healthcare-sentiment_v2",    "NEGATIVE / NEUTRAL / POSITIVE", "Lightweight, good for high volume"],
        ["bert_hc",             "BERT Healthcare",             "cjen1008/bert-healthcare-sentiment",             "NEGATIVE / NEUTRAL / POSITIVE", "v1 baseline, good benchmark"],
        ["distilbert_hc",       "DistilBERT Healthcare",       "cjen1008/distilbert-healthcare-sentiment",       "NEGATIVE / NEUTRAL / POSITIVE", "v1 lightweight baseline"],
        ["distilroberta_hc",    "DistilRoBERTa Healthcare",    "cjen1008/distilroberta-healthcare-sentiment",    "NEGATIVE / NEUTRAL / POSITIVE", "v1 RoBERTa baseline"],
    ],
    col_widths=[3.5, 4, 5.5, 3.5, 4]
)

add_callout(doc,
    "The v2 fine-tuned models (bert_hc_v2, distilroberta_hc_v2, distilbert_hc_v2) are "
    "retrained on a larger, more diverse healthcare dataset. Always prefer v2 over v1 "
    "unless you are running a benchmark comparison.",
    "tip"
)

h2(doc, "4.2  General-Purpose Pretrained Models")

body(doc, (
    "These models were not specifically trained on healthcare data but bring strong general "
    "language understanding. They are useful for comparison, for non-clinical text, or for "
    "detecting emotions beyond simple positive/negative polarity."
))

add_table(doc,
    ["model_type key", "Display Name", "HuggingFace ID", "Labels", "Best For"],
    [
        ["default",  "DistilBERT SST-2",    "distilbert-base-uncased-finetuned-sst-2-english",     "POSITIVE / NEGATIVE",                              "Quick binary sentiment baseline"],
        ["roberta",  "BERT Multilingual",   "nlptown/bert-base-multilingual-uncased-sentiment",    "1 STAR / 2 STARS / 3 STARS / 4 STARS / 5 STARS",  "Star-rating style sentiment"],
        ["emotion",  "GoEmotions",          "j-hartmann/emotion-english-distilroberta-base",       "ANGER / DISGUST / FEAR / JOY / NEUTRAL / SADNESS / SURPRISE", "Fine-grained emotion detection"],
        ["amazon",   "Amazon Reviews BERT", "sohan-ai/sentiment-analysis-model-amazon-reviews",   "POSITIVE / NEGATIVE",                              "Review-style feedback text"],
        ["twitter",  "RoBERTa Twitter",     "cardiffnlp/twitter-roberta-base-sentiment-latest",   "NEGATIVE / NEUTRAL / POSITIVE",                    "Informal or short-form feedback"],
        ["sst2",     "BERT SST-2",          "textattack/bert-base-uncased-SST-2",                  "POSITIVE / NEGATIVE",                              "Alternative binary baseline"],
        ["zeroshot", "BART Large MNLI",     "facebook/bart-large-mnli",                            "POSITIVE / NEGATIVE / NEUTRAL",                    "Zero-shot; no fine-tuning needed"],
    ],
    col_widths=[2.5, 3.5, 5, 4, 3.5]
)

add_callout(doc,
    "The zeroshot model (BART Large MNLI) is the largest and slowest model. "
    "It does not need training data for new label sets, but takes noticeably longer to respond. "
    "Use it for experimentation rather than high-throughput production calls.",
    "note"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. REQUEST & RESPONSE FORMAT
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "5. Request & Response Format")

h2(doc, "5.1  Request Examples")

h3(doc, "Default model (BERT Healthcare v2)")
add_code(doc, """\
POST /api/v1/analyze
Content-Type: application/json

{
  "text": "The nurse was incredibly supportive and made me feel completely at ease."
}""", "http")

h3(doc, "Specifying a different model")
add_code(doc, """\
POST /api/v1/analyze
Content-Type: application/json

{
  "text": "I waited three hours and was never called for my appointment.",
  "model_type": "distilroberta_hc_v2"
}""", "http")

h3(doc, "Using the emotion model")
add_code(doc, """\
POST /api/v1/analyze
Content-Type: application/json

{
  "text": "I am so relieved and grateful — the surgery went perfectly.",
  "model_type": "emotion"
}""", "http")

h2(doc, "5.2  Response Format")
body(doc, "A successful response (HTTP 200) always has this structure:")
add_code(doc, """\
{
  "sentiment":       "POSITIVE",
  "probabilities":   [0.03, 0.04, 0.93],
  "labels":          ["NEGATIVE", "NEUTRAL", "POSITIVE"],
  "model_type":      "bert_hc_v2",
  "redacted_text":   "The nurse was incredibly supportive and made me feel at ease.",
  "cleaned_text":    "nurse incredibly supportive made feel ease",
  "tokenized_text":  ["nurse", "incredibly", "supportive", "made", "feel", "ease"],
  "lemmatized_text": ["nurse", "incredibly", "supportive", "make", "feel", "ease"],
  "ner": [
    { "text": "nurse", "label": "PERSON" }
  ],
  "word_distribution": {
    "positive": 5,
    "neutral":  1,
    "negative": 0
  },
  "total_words": 6
}""", "json")

h3(doc, "Response Fields Explained")
add_table(doc,
    ["Field", "Type", "Description"],
    [
        ["sentiment",         "string",       "The predicted sentiment label (e.g. POSITIVE, NEGATIVE, JOY)."],
        ["probabilities",     "list of float","Confidence scores per label, aligned with the labels field. Sum to 1.0."],
        ["labels",            "list of string","The label names for this model, in the same order as probabilities."],
        ["model_type",        "string",       "The model key that was used (echoed back from the request)."],
        ["redacted_text",     "string",       "The text after PII/PHI redaction by OpenMed. This is what the model sees."],
        ["cleaned_text",      "string",       "After stop-word removal and punctuation cleaning."],
        ["tokenized_text",    "list of string","Individual tokens after NLP preprocessing."],
        ["lemmatized_text",   "list of string","Tokens reduced to their dictionary root form (e.g. 'waiting' → 'wait')."],
        ["ner",               "list of object","Named entity recognition results: [{\"text\": \"...\", \"label\": \"...\"}]."],
        ["word_distribution", "object",       "Per-label word counts from analysing each lemma individually."],
        ["total_words",       "integer",      "Number of tokens after preprocessing."],
    ],
    col_widths=[4, 3.5, 9]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 6. CODE EXAMPLES
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "6. Code Examples")

h2(doc, "6.1  cURL (Command Line)")
body(doc, "The simplest way to test. Copy-paste these directly into your terminal after starting the server.")

h3(doc, "Health check")
add_code(doc, """\
curl -s http://localhost:8000/api/v1/health | python3 -m json.tool""", "bash")

h3(doc, "Warm up all 13 models")
add_code(doc, """\
curl -s http://localhost:8000/api/v1/warmup | python3 -m json.tool""", "bash")

h3(doc, "Analyse a single text (default model)")
add_code(doc, """\
curl -s -X POST http://localhost:8000/api/v1/analyze \\
  -H "Content-Type: application/json" \\
  -d '{
    "text": "The doctor explained my diagnosis clearly and I felt well supported."
  }' | python3 -m json.tool""", "bash")

h3(doc, "Analyse with a specific model")
add_code(doc, """\
curl -s -X POST http://localhost:8000/api/v1/analyze \\
  -H "Content-Type: application/json" \\
  -d '{
    "text": "I was very anxious waiting for my biopsy results but the nurse called promptly.",
    "model_type": "emotion"
  }' | python3 -m json.tool""", "bash")

h2(doc, "6.2  Python")
body(doc, "Install the requests library first if you do not have it:")
add_code(doc, "pip install requests", "bash")

h3(doc, "Basic example — single text")
add_code(doc, """\
import requests

BASE_URL = "http://localhost:8000/api/v1"
HEADERS  = {"Content-Type": "application/json"}

# ── 1. Health check ────────────────────────────────────────────────────────
resp = requests.get(f"{BASE_URL}/health", headers=HEADERS)
print("Health:", resp.json())

# ── 2. Analyse a single text ───────────────────────────────────────────────
payload = {
    "text":       "The ward was clean and the staff were attentive throughout my stay.",
    "model_type": "bert_hc_v2",
}

resp = requests.post(f"{BASE_URL}/analyze", headers=HEADERS, json=payload)
resp.raise_for_status()    # raises an error if status is not 200

data = resp.json()
print(f"Sentiment:     {data['sentiment']}")
print(f"Probabilities: {dict(zip(data['labels'], data['probabilities']))}")
print(f"Redacted text: {data['redacted_text']}")""", "python")

h3(doc, "Analysing many texts in a loop")
add_code(doc, """\
import requests

BASE_URL   = "http://localhost:8000/api/v1"
HEADERS    = {"Content-Type": "application/json"}
MODEL      = "bert_hc_v2"

feedback_list = [
    "Dr Smith was very thorough and answered all my questions.",
    "I waited over two hours past my appointment time.",
    "The discharge instructions were unclear and confusing.",
    "The nurse was incredibly kind and reassuring.",
    "I could not get through on the phone to book an appointment.",
]

for text in feedback_list:
    resp = requests.post(
        f"{BASE_URL}/analyze",
        headers=HEADERS,
        json={"text": text, "model_type": MODEL},
    )
    resp.raise_for_status()
    result = resp.json()
    sentiment = result["sentiment"]
    top_prob  = max(result["probabilities"])
    print(f"{sentiment:10s}  ({top_prob:.0%})  |  {text[:60]}")""", "python")

h3(doc, "Comparing all 13 models on the same text")
add_code(doc, """\
import requests

BASE_URL = "http://localhost:8000/api/v1"
HEADERS  = {"Content-Type": "application/json"}

ALL_MODELS = [
    "bert_hc_v2", "distilroberta_hc_v2", "distilbert_hc_v2",
    "bert_hc",    "distilbert_hc",       "distilroberta_hc",
    "default",    "roberta",             "emotion",
    "amazon",     "twitter",             "sst2",   "zeroshot",
]

text = "The specialist was knowledgeable and I felt very confident in the diagnosis."

print(f"Text: {text}\\n")
print(f"{'Model':<25} {'Sentiment':<15} {'Top Confidence'}")
print("-" * 58)

for model in ALL_MODELS:
    resp = requests.post(
        f"{BASE_URL}/analyze",
        headers=HEADERS,
        json={"text": text, "model_type": model},
    )
    if resp.status_code == 200:
        r = resp.json()
        top_prob = max(r["probabilities"])
        print(f"{model:<25} {r['sentiment']:<15} {top_prob:.1%}")
    else:
        print(f"{model:<25} ERROR {resp.status_code}")""", "python")

h3(doc, "Saving results to a CSV file")
add_code(doc, """\
import requests
import csv

BASE_URL = "http://localhost:8000/api/v1"
HEADERS  = {"Content-Type": "application/json"}

feedback_list = [
    "The nurse was attentive and caring throughout my stay.",
    "I waited four hours with no update on my appointment.",
    "My discharge letter explained everything very clearly.",
]

with open("sentiment_results.csv", "w", newline="") as f:
    writer = csv.writer(f)
    writer.writerow(["text", "redacted_text", "sentiment", "confidence"])

    for text in feedback_list:
        resp = requests.post(
            f"{BASE_URL}/analyze",
            headers=HEADERS,
            json={"text": text, "model_type": "bert_hc_v2"},
        )
        resp.raise_for_status()
        r = resp.json()
        top_prob = max(r["probabilities"])
        writer.writerow([text, r["redacted_text"], r["sentiment"], round(top_prob, 4)])

print("Results saved to sentiment_results.csv")""", "python")

h2(doc, "6.3  JavaScript / Node.js")
body(doc, "Using the built-in fetch API (Node.js 18+ or any modern browser).")

h3(doc, "Analyse a single text")
add_code(doc, """\
const BASE_URL = "http://localhost:8000/api/v1";

async function analyseSentiment(text, modelType = "bert_hc_v2") {
  const response = await fetch(`${BASE_URL}/analyze`, {
    method:  "POST",
    headers: { "Content-Type": "application/json" },
    body: JSON.stringify({ text, model_type: modelType }),
  });

  if (!response.ok) {
    throw new Error(`API error: ${response.status} ${response.statusText}`);
  }

  return response.json();
}

// Usage
analyseSentiment(
  "The doctor was very kind and explained everything clearly.",
  "bert_hc_v2"
)
  .then(data => {
    console.log("Sentiment:    ", data.sentiment);
    console.log("Redacted text:", data.redacted_text);
    const scores = Object.fromEntries(
      data.labels.map((l, i) => [l, data.probabilities[i].toFixed(3)])
    );
    console.log("Probabilities:", scores);
  })
  .catch(err => console.error("Error:", err));""", "javascript")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 7. TESTING GUIDE
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "7. Testing Guide")

body(doc, (
    "This section provides copy-paste test cases that cover the main functionality, "
    "edge cases, and error scenarios. Run these to verify the API is working correctly "
    "after setup or after any code change."
))

h2(doc, "7.1  Prerequisites")
body(doc, "Ensure the server is running before executing any test:")
add_code(doc, """\
uvicorn api.main:app --host 0.0.0.0 --port 8000 --reload""", "bash")
body(doc, "Set a shell variable for convenience:")
add_code(doc, "BASE=http://localhost:8000/api/v1", "bash")

h2(doc, "7.2  Test 1 — Health Check")
body(doc, "Verifies the API is live.")
add_code(doc, """\
curl -s $BASE/health | python3 -m json.tool""", "bash")
body(doc, "✅  Pass if: status is \"ok\".")

h2(doc, "7.3  Test 2 — Warmup All Models")
body(doc, "Verifies all 13 models can be loaded without errors.")
add_code(doc, """\
curl -s $BASE/warmup | python3 -m json.tool""", "bash")
body(doc, "✅  Pass if: status is \"ok\" and all 13 model keys appear with \"loaded in Xs\" (not \"error\").")

h2(doc, "7.4  Test 3 — Positive Sentiment")
body(doc, "Verifies a clearly positive text is classified correctly.")
add_code(doc, """\
curl -s -X POST $BASE/analyze \\
  -H "Content-Type: application/json" \\
  -d '{
    "text": "The specialist was outstanding and I felt completely confident throughout.",
    "model_type": "bert_hc_v2"
  }' | python3 -m json.tool""", "bash")
body(doc, "✅  Pass if: sentiment is \"POSITIVE\".")

h2(doc, "7.5  Test 4 — Negative Sentiment")
body(doc, "Verifies a clearly negative text is classified correctly.")
add_code(doc, """\
curl -s -X POST $BASE/analyze \\
  -H "Content-Type: application/json" \\
  -d '{
    "text": "I waited four hours without any update and was never seen by a doctor.",
    "model_type": "bert_hc_v2"
  }' | python3 -m json.tool""", "bash")
body(doc, "✅  Pass if: sentiment is \"NEGATIVE\".")

h2(doc, "7.6  Test 5 — Emotion Detection")
body(doc, "Verifies the GoEmotions model returns one of the 7 emotion labels.")
add_code(doc, """\
curl -s -X POST $BASE/analyze \\
  -H "Content-Type: application/json" \\
  -d '{
    "text": "I am so relieved and grateful that my surgery went so well.",
    "model_type": "emotion"
  }' | python3 -m json.tool""", "bash")
body(doc, "✅  Pass if: sentiment is one of ANGER, DISGUST, FEAR, JOY, NEUTRAL, SADNESS, SURPRISE.")

h2(doc, "7.7  Test 6 — PII Redaction")
body(doc, "Verifies that patient names and identifiers are removed before analysis.")
add_code(doc, """\
curl -s -X POST $BASE/analyze \\
  -H "Content-Type: application/json" \\
  -d '{
    "text": "Dr Jane Smith at St Vincents Hospital on 12 March 2026 was very helpful.",
    "model_type": "bert_hc_v2"
  }' | python3 -m json.tool""", "bash")
body(doc, "✅  Pass if: redacted_text does not contain the original name, date, or hospital name.")

h2(doc, "7.8  Test 7 — Error Cases")
body(doc, "These tests verify the API returns the correct error codes for bad requests.")

h3(doc, "Text too short → 422 Unprocessable")
add_code(doc, """\
curl -s -o /dev/null -w "%{http_code}" \\
  -X POST $BASE/analyze \\
  -H "Content-Type: application/json" \\
  -d '{"text": "Good"}'""", "bash")
body(doc, "✅  Pass if: output is 422.")

h3(doc, "Invalid model_type → 422 Unprocessable")
add_code(doc, """\
curl -s -X POST $BASE/analyze \\
  -H "Content-Type: application/json" \\
  -d '{
    "text": "The doctor was very kind and explained my diagnosis clearly.",
    "model_type": "invalid_model_name"
  }' | python3 -m json.tool""", "bash")
body(doc, "✅  Pass if: response contains a detail message listing valid model keys.")

h3(doc, "Missing text field → 422 Unprocessable")
add_code(doc, """\
curl -s -X POST $BASE/analyze \\
  -H "Content-Type: application/json" \\
  -d '{"model_type": "bert_hc_v2"}' | python3 -m json.tool""", "bash")
body(doc, "✅  Pass if: response contains a validation error for the missing text field.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 8. ERROR REFERENCE
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "8. Error Reference")

body(doc, "The API uses standard HTTP status codes. Here is what each one means and how to fix it:")

add_table(doc,
    ["HTTP Status", "Name", "Cause", "How to Fix"],
    [
        ["200", "OK",
         "Request succeeded.",
         "No action needed."],
        ["422", "Unprocessable Entity",
         "text is missing, under 4 words, over 300 words, or model_type is invalid.",
         "Check that text is present and 4–300 words. Verify model_type is a valid key from Section 4."],
        ["400", "Bad Request",
         "The request body is not valid JSON.",
         "Check for missing quotes, commas, or brackets in your JSON."],
        ["500", "Internal Server Error",
         "An unexpected error occurred (e.g. a model failed to load, OpenMed PII service unreachable).",
         "Check the server terminal output. Verify OPENMED_PII_API_KEY is set in .env."],
        ["503", "Service Unavailable",
         "The server is not running or crashed.",
         "Run uvicorn api.main:app --host 0.0.0.0 --port 8000 --reload to restart."],
    ],
    col_widths=[2.2, 3.5, 6, 5]
)

h2(doc, "8.1  Checking Server Logs")
body(doc, (
    "When the server is running with --reload, all errors are printed in the terminal "
    "where uvicorn is running. Look for lines starting with ERROR or for Python tracebacks."
))

add_callout(doc,
    "If you get a 500 error relating to OpenMed PII, the most common cause is a missing "
    "or incorrect OPENMED_PII_API_KEY in .env. Check the value and restart the server.",
    "warning"
)

h2(doc, "8.2  Interactive API Docs")
body(doc, (
    "FastAPI automatically generates interactive documentation at "
    "http://localhost:8000/docs. You can try every endpoint directly in your browser "
    "without writing any code. This is the fastest way to debug a request format issue."
))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 9. BEST PRACTICES
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "9. Best Practices")

h2(doc, "9.1  Performance")
for item in [
    "Call GET /api/v1/warmup once after server startup to pre-load all 13 models. "
    "This avoids per-model delays on the first request to each model.",
    "The first request to any unloaded model may take 10–30 seconds while the weights "
    "are loaded from disk. Subsequent requests are fast (under 2 seconds).",
    "Keep texts under 300 words. Longer texts are rejected with 422.",
    "If you need to classify many texts, run them in a loop but add a small delay (0.1s) "
    "between calls to avoid overloading the local server.",
    "The zeroshot model (BART Large MNLI) is significantly slower than the others. "
    "Avoid it in time-sensitive workflows.",
]:
    bullet(doc, item)

h2(doc, "9.2  Security")
for item in [
    "Never commit .env to Git. It contains the OpenMed PII API key.",
    "Do not expose the local API on a public network. It has no authentication by default.",
    "Always use the redacted_text field from the response when logging or storing results — "
    "never log the original patient text.",
    "Run the API inside Docker for consistent, isolated environments across team members.",
]:
    bullet(doc, item)

body(doc, "Load credentials safely in Python:")
add_code(doc, """\
import os
from dotenv import load_dotenv

load_dotenv()   # reads .env from the current directory

OPENMED_KEY = os.environ.get("OPENMED_PII_API_KEY")
if not OPENMED_KEY:
    raise ValueError("OPENMED_PII_API_KEY is not set in .env")""", "python")

h2(doc, "9.3  Error Handling")
body(doc, "Always handle errors in your code. Here is a robust Python wrapper:")
add_code(doc, """\
import requests

def analyse_sentiment(text: str, model_type: str = "bert_hc_v2") -> dict:
    url     = "http://localhost:8000/api/v1/analyze"
    headers = {"Content-Type": "application/json"}

    try:
        resp = requests.post(
            url,
            headers=headers,
            json={"text": text, "model_type": model_type},
            timeout=60,     # wait up to 60 s (covers first-load of large models)
        )
        resp.raise_for_status()
        return resp.json()

    except requests.exceptions.Timeout:
        print("Request timed out — model may still be loading. Try again.")
        raise
    except requests.exceptions.HTTPError as e:
        if e.response.status_code == 422:
            print("Invalid request:", e.response.json().get("detail"))
        elif e.response.status_code >= 500:
            print("Server error — check the uvicorn terminal for details.")
        raise
    except requests.exceptions.ConnectionError:
        print("Cannot connect. Is the server running on port 8000?")
        raise""", "python")

h2(doc, "9.4  Choosing the Right Model")
add_table(doc,
    ["Use Case", "Recommended model_type", "Why"],
    [
        ["Clinical notes, production use",     "bert_hc_v2",          "Best accuracy on healthcare text. Default choice."],
        ["High volume, lower latency needed",  "distilbert_hc_v2",    "Smallest fine-tuned model; fastest inference."],
        ["Detailed emotion detection",         "emotion",             "Returns 7 emotion classes for nuanced analysis."],
        ["5-star rating style output",         "roberta",             "Maps feedback to a 1–5 star scale."],
        ["Informal or social media text",      "twitter",             "Trained on informal language and abbreviations."],
        ["Zero-shot / no fine-tuning",         "zeroshot",            "Works without any domain training data."],
        ["Benchmarking against all models",    "Run all 13 in a loop","Compare every model's output side by side."],
    ],
    col_widths=[4.5, 4, 8]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 10. INFRASTRUCTURE OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "10. Infrastructure Overview")

body(doc, (
    "This section describes how the system is structured so you understand what happens "
    "when you make a request. You do not need to manage any of this — it is here for context."
))

h2(doc, "10.1  Local Deployment (Development)")
add_table(doc,
    ["Component", "Description"],
    [
        ["FastAPI server",       "The Python web framework serving the API on port 8000."],
        ["uvicorn",              "The ASGI server that runs FastAPI. Started with: uvicorn api.main:app."],
        ["src/preprocessor.py", "Cleans, tokenises, stems, lemmatises, and extracts NER/POS from text."],
        ["src/analyzer.py",     "Lazy-loads Transformer models and runs inference."],
        ["src/models.py",       "Defines all 13 model configurations (HuggingFace IDs, labels, label maps)."],
        ["OpenMed PII Service", "External AWS Lambda that redacts PII/PHI. Called via HTTPS before every analysis."],
        ["HuggingFace Cache",   "Downloaded model weights cached locally under ~/.cache/huggingface/ after first use."],
    ],
    col_widths=[4, 12.5]
)

h2(doc, "10.2  Docker Deployment")
body(doc, "To run the API inside a Docker container (recommended for consistent environments):")
add_code(doc, """\
# Build the image
docker build -t nlp-sentiment .

# Run the container (forwards port 8000, injects .env values)
docker run -p 7860:7860 \\
  -e OPENMED_PII_BASE_URL=$OPENMED_PII_BASE_URL \\
  -e OPENMED_PII_API_KEY=$OPENMED_PII_API_KEY \\
  nlp-sentiment""", "bash")

add_callout(doc,
    "The Dockerfile pre-downloads all 13 HuggingFace models and NLTK data at build time. "
    "The resulting image is large (~6 GB) but has zero first-request model load delays.",
    "note"
)

h2(doc, "10.3  Request Flow")
body(doc, "What happens when you call POST /api/v1/analyze:")
for step in [
    "Your code sends a POST request with the text and model_type.",
    "FastAPI validates the request body (checks text length, model_type is valid).",
    "The route handler calls the OpenMed PII service (HTTPS) to redact names, dates, and IDs.",
    "The redacted text is passed to src/preprocessor.preprocess_text — spaCy cleans the text, "
    "NLTK tokenises, Porter stemmer stems, WordNet lemmatises, and NER/POS tags are extracted.",
    "src/analyzer.analyze_sentiment runs inference using the chosen Transformer model "
    "(lazy-loaded on first call, then cached in memory).",
    "src/analyzer.get_word_distribution scores each lemma individually to build the "
    "per-word sentiment distribution.",
    "The JSON response is returned to your code.",
]:
    numbered(doc, step)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 11. TROUBLESHOOTING
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "11. Troubleshooting")

add_table(doc,
    ["Problem", "Likely Cause", "Solution"],
    [
        ["Connection refused on port 8000",
         "The uvicorn server is not running.",
         "Run: uvicorn api.main:app --host 0.0.0.0 --port 8000 --reload"],
        ["500 error on every analyze request",
         "OPENMED_PII_API_KEY is missing or wrong in .env.",
         "Check .env contains the correct key. Restart the server after editing .env."],
        ["First request takes 20–30 seconds",
         "Model weights are being loaded from disk on first use.",
         "Normal behaviour. Call GET /api/v1/warmup after startup to pre-load all models."],
        ["422: Text must be at least 4 words",
         "The submitted text is too short.",
         "Ensure your text has at least 4 meaningful words before calling the API."],
        ["422: model_type must be one of ...",
         "The model_type value is misspelled or not in the supported list.",
         "Check Section 4 for the exact model key strings. Keys are case-sensitive."],
        ["ModuleNotFoundError on startup",
         "Dependencies are not installed or virtual env is not activated.",
         "Run: source .venv/bin/activate && pip install -r requirements.txt"],
        ["OSError: spaCy model not found",
         "spaCy en_core_web_sm is not installed.",
         "Run: python -m spacy download en_core_web_sm"],
        ["Sentiment always shows the same label regardless of text",
         "A model may have loaded incorrectly.",
         "Restart the server and call /warmup again. Check the uvicorn terminal for errors."],
        ["redacted_text looks identical to the original text",
         "No PII was detected (this is normal for non-PII text) or OpenMed service is down.",
         "Verify OPENMED_PII_BASE_URL in .env is correct. Check internet connectivity."],
        ["NLTK LookupError for punkt or wordnet",
         "NLTK data packages are not downloaded.",
         "Run: python -c \"import nltk; nltk.download('all')\" or check nltk_data/ directory."],
    ],
    col_widths=[4, 4.5, 8]
)

h2(doc, "11.1  Useful Debugging Commands")

h3(doc, "Check which models are currently loaded")
add_code(doc, """\
curl -s http://localhost:8000/api/v1/health | python3 -m json.tool""", "bash")

h3(doc, "View the interactive API docs")
add_code(doc, "open http://localhost:8000/docs    # macOS\nxdg-open http://localhost:8000/docs  # Linux", "bash")

h3(doc, "Inspect model configuration in Python")
add_code(doc, """\
from src.models import SUPPORTED_MODELS
for key, cfg in SUPPORTED_MODELS.items():
    print(f"{key:<25} {cfg['hf_id']}")""", "python")

h3(doc, "Run the examples script")
add_code(doc, """\
# Make sure the server is running, then:
python examples/basic_usage.py""", "bash")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 12. QUICK REFERENCE CARD
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "12. Quick Reference Card")

body(doc, "Tear out and keep this page handy.")

h2(doc, "Server")
add_code(doc, """\
# Start server
uvicorn api.main:app --host 0.0.0.0 --port 8000 --reload

# Or Docker
docker run -p 7860:7860 nlp-sentiment""", "bash")

h2(doc, "Endpoints")
add_table(doc,
    ["Method", "Path", "Use For"],
    [
        ["GET",  "/api/v1/health",  "Check API is alive."],
        ["GET",  "/api/v1/warmup",  "Pre-load all 13 models (run once after startup)."],
        ["POST", "/api/v1/analyze", "Analyse sentiment of a text using a chosen model."],
    ],
    col_widths=[2, 4.5, 10]
)

h2(doc, "Models at a Glance")
add_table(doc,
    ["Key", "Labels", "Type"],
    [
        ["bert_hc_v2",          "NEGATIVE / NEUTRAL / POSITIVE",                              "Fine-tuned HC (v2)"],
        ["distilroberta_hc_v2", "NEGATIVE / NEUTRAL / POSITIVE",                              "Fine-tuned HC (v2)"],
        ["distilbert_hc_v2",    "NEGATIVE / NEUTRAL / POSITIVE",                              "Fine-tuned HC (v2)"],
        ["bert_hc",             "NEGATIVE / NEUTRAL / POSITIVE",                              "Fine-tuned HC (v1)"],
        ["distilbert_hc",       "NEGATIVE / NEUTRAL / POSITIVE",                              "Fine-tuned HC (v1)"],
        ["distilroberta_hc",    "NEGATIVE / NEUTRAL / POSITIVE",                              "Fine-tuned HC (v1)"],
        ["default",             "POSITIVE / NEGATIVE",                                        "Pretrained (SST-2)"],
        ["roberta",             "1 STAR / 2 STARS / 3 STARS / 4 STARS / 5 STARS",            "Pretrained (Multilingual)"],
        ["emotion",             "ANGER / DISGUST / FEAR / JOY / NEUTRAL / SADNESS / SURPRISE","Pretrained (GoEmotions)"],
        ["amazon",              "POSITIVE / NEGATIVE",                                        "Pretrained (Reviews)"],
        ["twitter",             "NEGATIVE / NEUTRAL / POSITIVE",                              "Pretrained (Twitter)"],
        ["sst2",                "POSITIVE / NEGATIVE",                                        "Pretrained (SST-2)"],
        ["zeroshot",            "POSITIVE / NEGATIVE / NEUTRAL",                              "Pretrained (Zero-shot)"],
    ],
    col_widths=[4, 6.5, 4]
)

h2(doc, "Minimal cURL")
add_code(doc, """\
curl -s -X POST http://localhost:8000/api/v1/analyze \\
  -H "Content-Type: application/json" \\
  -d '{"text": "YOUR TEXT HERE", "model_type": "bert_hc_v2"}' \\
  | python3 -m json.tool""", "bash")

h2(doc, "Minimal Python")
add_code(doc, """\
import requests
r = requests.post(
    "http://localhost:8000/api/v1/analyze",
    headers={"Content-Type": "application/json"},
    json={"text": "YOUR TEXT HERE", "model_type": "bert_hc_v2"},
)
print(r.json()["sentiment"], r.json()["redacted_text"])""", "python")

h2(doc, "HTTP Status Codes")
add_table(doc,
    ["Code", "Meaning", "Fix"],
    [
        ["200", "Success",            "No action needed."],
        ["422", "Bad request data",   "Check text is 4–300 words. Check model_type is valid."],
        ["500", "Server error",       "Check uvicorn terminal. Verify .env has the OpenMed key."],
        ["503", "Server not running", "Run uvicorn api.main:app --host 0.0.0.0 --port 8000."],
    ],
    col_widths=[1.8, 3.5, 11.2]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 13. MODEL EVALUATION WITH TEST SAMPLES
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "13. Model Evaluation with Test Samples")

body(doc, (
    "This section provides a structured evaluation of all 13 models using carefully designed "
    "healthcare test samples. Each sample has a known expected output so you can measure "
    "how accurately each model performs, understand the differences between model families, "
    "and choose the right model for your application."
))

add_callout(doc, (
    "This evaluation is designed to run against a live local server. "
    "Start the server and call GET /api/v1/warmup before running any evaluation script. "
    "Results may vary slightly depending on model version and runtime environment."
), "note")

# ── 13.1 ──────────────────────────────────────────────────────────────────────
h2(doc, "13.1  Evaluation Metrics Explained")

body(doc, "The following metrics are used to assess each model:")

add_table(doc,
    ["Metric", "What It Measures", "How to Read It"],
    [
        ["Top-1 Accuracy",
         "The % of test texts where the model's predicted label matches the expected label.",
         "80% means 8 out of 10 texts were labelled correctly. Higher is better."],
        ["Avg Confidence (Correct)",
         "Mean probability of the predicted label across correctly classified texts.",
         "High value (0.80+) means the model is confident when right."],
        ["Avg Confidence (Wrong)",
         "Mean probability of the predicted label across misclassified texts.",
         "High value here (0.75+) means the model is confidently wrong — a red flag."],
        ["Per-sample breakdown",
         "Pass/Fail for every individual test case.",
         "Reveals which text types each model handles well or struggles with."],
    ],
    col_widths=[4.5, 6, 6]
)

add_callout(doc, (
    "Note that the 7 general-purpose models use different label sets. "
    "For binary models (default, amazon, sst2) expected outputs are POSITIVE or NEGATIVE only. "
    "For the emotion model expected outputs are one of 7 emotion classes. "
    "The roberta model returns 1–5 STARS. Accuracy is measured per-model against the "
    "appropriate expected label for that model."
), "note")

# ── 13.2 ──────────────────────────────────────────────────────────────────────
h2(doc, "13.2  Test Sample Dataset")

body(doc, (
    "The evaluation uses 10 representative healthcare test cases. Each case has a "
    "primary expected sentiment (NEGATIVE / NEUTRAL / POSITIVE) used to evaluate the "
    "six HC fine-tuned models and all binary/ternary general-purpose models. "
    "A separate expected emotion label is provided for the GoEmotions model, and a "
    "star rating band (LOW / HIGH) is provided for the BERT Multilingual model."
))

add_table(doc,
    ["ID", "Text", "Expected Sentiment", "Expected Emotion", "Expected Star Band"],
    [
        ["T-01",
         "The specialist was incredibly knowledgeable and took time to explain every step of my treatment.",
         "POSITIVE", "JOY", "HIGH (4–5 stars)"],
        ["T-02",
         "I waited over four hours in the waiting room and was never given any update.",
         "NEGATIVE", "ANGER", "LOW (1–2 stars)"],
        ["T-03",
         "My appointment was scheduled for 10 am on Tuesday in the outpatient clinic.",
         "NEUTRAL", "NEUTRAL", "NEUTRAL (3 stars)"],
        ["T-04",
         "The ward was absolutely filthy and the smell was unpleasant throughout my stay.",
         "NEGATIVE", "DISGUST", "LOW (1–2 stars)"],
        ["T-05",
         "I felt genuinely relieved after the consultation — the results were much better than I feared.",
         "POSITIVE", "JOY", "HIGH (4–5 stars)"],
        ["T-06",
         "The nurse explained the medication dosage and told me to return in two weeks for a follow-up.",
         "NEUTRAL", "NEUTRAL", "NEUTRAL (3 stars)"],
        ["T-07",
         "I was terrified before the procedure but the anaesthetist reassured me and I felt safe.",
         "POSITIVE", "FEAR", "HIGH (4–5 stars)"],
        ["T-08",
         "My discharge letter had the wrong medication listed and I was not told about the side effects.",
         "NEGATIVE", "SADNESS", "LOW (1–2 stars)"],
        ["T-09",
         "I cannot thank the nursing team enough — they were warm, professional, and incredibly caring.",
         "POSITIVE", "JOY", "HIGH (4–5 stars)"],
        ["T-10",
         "The referral was lost and I had to wait an additional three months to see the specialist.",
         "NEGATIVE", "SADNESS", "LOW (1–2 stars)"],
    ],
    col_widths=[1, 6.5, 3, 2.8, 3.2]
)

# ── 13.3 ──────────────────────────────────────────────────────────────────────
h2(doc, "13.3  Running the Full Evaluation Script")

body(doc, (
    "The script below sends all 10 test cases to all 13 models and prints a "
    "pass/fail table plus a summary of per-model accuracy."
))

h3(doc, "Step 1 — Ensure the server is running and all models are warmed up")
add_code(doc, """\
uvicorn api.main:app --host 0.0.0.0 --port 8000 --reload &
curl -s http://localhost:8000/api/v1/warmup | python3 -m json.tool""", "bash")

h3(doc, "Step 2 — Save and run the evaluation script")
add_code(doc, """\
# Save as evaluate_models.py and run with: python3 evaluate_models.py
import requests

BASE_URL = "http://localhost:8000/api/v1"
HEADERS  = {"Content-Type": "application/json"}

# ── Test cases: (id, text, expected_sentiment, expected_emotion, expected_star_band) ──
TEST_CASES = [
    ("T-01",
     "The specialist was incredibly knowledgeable and took time to explain every step.",
     "POSITIVE", "JOY", "HIGH"),
    ("T-02",
     "I waited over four hours in the waiting room and was never given any update.",
     "NEGATIVE", "ANGER", "LOW"),
    ("T-03",
     "My appointment was scheduled for 10 am on Tuesday in the outpatient clinic.",
     "NEUTRAL", "NEUTRAL", "NEUTRAL"),
    ("T-04",
     "The ward was absolutely filthy and the smell was unpleasant throughout my stay.",
     "NEGATIVE", "DISGUST", "LOW"),
    ("T-05",
     "I felt genuinely relieved after the consultation — results were better than I feared.",
     "POSITIVE", "JOY", "HIGH"),
    ("T-06",
     "The nurse explained the dosage and told me to return in two weeks for a follow-up.",
     "NEUTRAL", "NEUTRAL", "NEUTRAL"),
    ("T-07",
     "I was terrified before the procedure but the anaesthetist reassured me and I felt safe.",
     "POSITIVE", "FEAR", "HIGH"),
    ("T-08",
     "My discharge letter had the wrong medication listed and I was not told about side effects.",
     "NEGATIVE", "SADNESS", "LOW"),
    ("T-09",
     "I cannot thank the nursing team enough — warm, professional, and incredibly caring.",
     "POSITIVE", "JOY", "HIGH"),
    ("T-10",
     "The referral was lost and I had to wait an additional three months to see the specialist.",
     "NEGATIVE", "SADNESS", "LOW"),
]

# ── Model groups and their expected-label field ────────────────────────────────
HC_MODELS   = ["bert_hc_v2", "distilroberta_hc_v2", "distilbert_hc_v2",
               "bert_hc",    "distilbert_hc",        "distilroberta_hc"]
BIN_MODELS  = ["default", "amazon", "sst2"]       # POSITIVE / NEGATIVE only
TERN_MODELS = ["twitter", "zeroshot"]              # NEGATIVE / NEUTRAL / POSITIVE
STAR_MODEL  = "roberta"
EMO_MODEL   = "emotion"

STAR_MAP = {
    "1 STAR": "LOW", "2 STARS": "LOW",
    "3 STARS": "NEUTRAL",
    "4 STARS": "HIGH", "5 STARS": "HIGH",
}

# Flatten: (model_key, expected_label_field_index)
# field 0 = sentiment, 1 = emotion, 2 = star_band (mapped)
def expected_for(model, case):
    _, _, sentiment, emotion, star_band = case
    if model == EMO_MODEL:
        return emotion
    if model == STAR_MODEL:
        return star_band            # compare via STAR_MAP
    if model in BIN_MODELS:
        # binary: map NEUTRAL → POSITIVE (most neutral HC texts lean positive)
        return "NEGATIVE" if sentiment == "NEGATIVE" else "POSITIVE"
    return sentiment                # HC and ternary models


# ── Evaluate all models ────────────────────────────────────────────────────────
print("Running evaluation...\\n")

results_table = {}   # model → {correct, total, conf_c, conf_w}

for model in HC_MODELS + BIN_MODELS + TERN_MODELS + [STAR_MODEL, EMO_MODEL]:
    correct, total = 0, 0
    conf_c, conf_w = [], []
    per_case = []

    for case in TEST_CASES:
        tid, text, *_ = case
        resp = requests.post(
            f"{BASE_URL}/analyze",
            headers=HEADERS,
            json={"text": text, "model_type": model},
            timeout=60,
        )
        if resp.status_code != 200:
            per_case.append((tid, "ERROR", "—", False))
            total += 1
            continue

        r          = resp.json()
        raw_pred   = r["sentiment"]
        top_conf   = max(r["probabilities"])
        total     += 1

        # Normalise prediction for star model
        if model == STAR_MODEL:
            pred = STAR_MAP.get(raw_pred, "NEUTRAL")
        else:
            pred = raw_pred

        exp  = expected_for(model, case)
        hit  = (pred == exp)
        if hit:
            correct += 1
            conf_c.append(top_conf)
        else:
            conf_w.append(top_conf)
        per_case.append((tid, pred, exp, hit))

    results_table[model] = {
        "correct": correct, "total": total,
        "conf_c":  conf_c,  "conf_w": conf_w,
        "per_case": per_case,
    }

# ── Print per-case breakdown ───────────────────────────────────────────────────
for model, data in results_table.items():
    print(f"\\n{'='*60}")
    print(f"Model: {model}")
    print(f"{'ID':<6} {'Predicted':<15} {'Expected':<15} {'Result'}")
    print("-" * 55)
    for tid, pred, exp, hit in data["per_case"]:
        mark = "✅ PASS" if hit else "❌ FAIL"
        print(f"{tid:<6} {pred:<15} {exp:<15} {mark}")

# ── Print summary table ────────────────────────────────────────────────────────
print("\\n" + "="*70)
print(f"{'Model':<25} {'Accuracy':>10} {'AvgConf✅':>11} {'AvgConf❌':>11}")
print("-" * 60)

for model, data in results_table.items():
    acc    = data["correct"] / data["total"] * 100 if data["total"] else 0
    avg_c  = sum(data["conf_c"]) / len(data["conf_c"]) if data["conf_c"] else 0.0
    avg_w  = sum(data["conf_w"]) / len(data["conf_w"]) if data["conf_w"] else 0.0
    bar    = "█" * data["correct"] + "░" * (data["total"] - data["correct"])
    print(f"{model:<25} {acc:>9.0f}%  {avg_c:>10.2f}   {avg_w:>10.2f}   {bar}")

print("-" * 60)
print("Save this script as evaluate_models.py and run: python3 evaluate_models.py")""", "python")

add_callout(doc,
    "Save this script as evaluate_models.py in the project root and run it with: "
    "python3 evaluate_models.py",
    "tip"
)

# ── 13.4 ──────────────────────────────────────────────────────────────────────
h2(doc, "13.4  Expected Evaluation Results")

body(doc, (
    "The table below shows the expected sentiment output for each of the 10 test cases "
    "across the 13 models. Use this as your ground-truth reference when reading the "
    "evaluation script output."
))

add_table(doc,
    ["Test ID", "Text (abbreviated)", "HC Models (6)", "Binary Models (3)", "Ternary Models (2)", "BERT Multilingual", "GoEmotions"],
    [
        ["T-01", "Specialist knowledgeable, explained treatment.",
         "POSITIVE", "POSITIVE", "POSITIVE", "4–5 STARS", "JOY"],
        ["T-02", "Waited 4 hours, no update.",
         "NEGATIVE", "NEGATIVE", "NEGATIVE", "1–2 STARS", "ANGER"],
        ["T-03", "Appointment scheduled Tuesday outpatient.",
         "NEUTRAL",  "POSITIVE", "NEUTRAL",  "3 STARS",   "NEUTRAL"],
        ["T-04", "Ward filthy, unpleasant smell.",
         "NEGATIVE", "NEGATIVE", "NEGATIVE", "1–2 STARS", "DISGUST"],
        ["T-05", "Relieved after consultation, results good.",
         "POSITIVE", "POSITIVE", "POSITIVE", "4–5 STARS", "JOY"],
        ["T-06", "Nurse explained dosage, return in 2 weeks.",
         "NEUTRAL",  "POSITIVE", "NEUTRAL",  "3 STARS",   "NEUTRAL"],
        ["T-07", "Terrified before procedure, anaesthetist reassured.",
         "POSITIVE", "POSITIVE", "POSITIVE", "4–5 STARS", "FEAR"],
        ["T-08", "Discharge letter wrong medication, no side effect info.",
         "NEGATIVE", "NEGATIVE", "NEGATIVE", "1–2 STARS", "SADNESS"],
        ["T-09", "Cannot thank nursing team enough, caring.",
         "POSITIVE", "POSITIVE", "POSITIVE", "4–5 STARS", "JOY"],
        ["T-10", "Referral lost, 3-month additional wait.",
         "NEGATIVE", "NEGATIVE", "NEGATIVE", "1–2 STARS", "SADNESS"],
    ],
    col_widths=[1.2, 4.5, 2.5, 2.8, 2.8, 2.8, 2.5]
)

add_callout(doc, (
    "Binary models (default, amazon, sst2) do not have a NEUTRAL class. "
    "Factual/neutral healthcare texts (T-03, T-06) may be classified as POSITIVE "
    "by these models because neutral clinical language lacks negative words. "
    "This is expected behaviour, not a bug."
), "note")

# ── 13.5 ──────────────────────────────────────────────────────────────────────
h2(doc, "13.5  Interpreting Evaluation Results")

body(doc, (
    "Once you have run the evaluation script, use this guide to interpret the numbers "
    "and choose the best model for your application."
))

h3(doc, "Reading the accuracy table")
add_table(doc,
    ["Top-1 Accuracy", "What It Means", "Recommendation"],
    [
        ["90–100%", "Excellent — the model almost always picks the right label.",
         "Use this model in production."],
        ["75–89%",  "Good — correct most of the time.",
         "Suitable for production. Consider showing top-2 probabilities as a safety net."],
        ["60–74%",  "Fair — wrong roughly 1 in 4 times.",
         "Use with caution. Consider pairing with a higher-accuracy model."],
        ["Below 60%","Poor — struggles with this dataset.",
         "Do not use alone. Compare with other models or consider retraining."],
    ],
    col_widths=[3.5, 7, 6]
)

h3(doc, "Understanding confidence scores")
for item in [
    "High accuracy + high confidence (0.80+): Model is reliable — trust its top-1 prediction.",
    "High accuracy + low confidence (below 0.50): Model is correct but uncertain — always show probabilities.",
    "Low accuracy + high confidence: Model is confidently wrong — do not deploy on this data.",
    "Low accuracy + low confidence: Model is uncertain and often wrong — use a different model.",
]:
    bullet(doc, item)

add_callout(doc,
    "If the top-1 confidence is below 0.60, show all label probabilities to the user "
    "rather than just the top prediction. This gives them context to interpret the result.",
    "tip"
)

h3(doc, "Comparing all 13 models on the same text")
add_code(doc, """\
# After running the evaluation script, use majority voting to find the consensus:
from collections import Counter

# Example: collect all predictions for one test text
predictions = {
    "bert_hc_v2":          "POSITIVE",
    "distilroberta_hc_v2": "POSITIVE",
    "distilbert_hc_v2":    "POSITIVE",
    "bert_hc":             "POSITIVE",
    "distilbert_hc":       "NEUTRAL",
    "distilroberta_hc":    "POSITIVE",
    "default":             "POSITIVE",
    "amazon":              "POSITIVE",
    "twitter":             "POSITIVE",
    "sst2":                "POSITIVE",
    "zeroshot":            "POSITIVE",
    # emotion and roberta use different label sets — exclude from this vote
}

winner, count = Counter(predictions.values()).most_common(1)[0]
print(f"Majority vote: {winner}  ({count}/{len(predictions)} models agreed)")""", "python")

# ── 13.6 ──────────────────────────────────────────────────────────────────────
h2(doc, "13.6  Per-Model Characteristics Summary")

body(doc, (
    "This table summarises the typical strengths and weaknesses of each model based "
    "on the evaluation results and known characteristics of the underlying architectures."
))

add_table(doc,
    ["Model", "Typical Accuracy on HC Text", "Strengths", "Weaknesses"],
    [
        ["bert_hc_v2",
         "High (90%+)",
         "Best domain fit. Handles neutral class well. Consistent across text styles.",
         "Slower than DistilBERT models."],
        ["distilroberta_hc_v2",
         "High (87–92%)",
         "Strong RoBERTa-based encoder. Good at negation and nuance.",
         "Slightly more memory than DistilBERT."],
        ["distilbert_hc_v2",
         "High (85–90%)",
         "Fastest HC fine-tuned model. Good accuracy-to-speed ratio.",
         "May miss subtle neutral signals."],
        ["bert_hc",
         "Good (80–87%)",
         "Solid v1 baseline. Useful for benchmark comparison.",
         "Superseded by v2. Lower accuracy on borderline cases."],
        ["distilbert_hc",
         "Good (78–85%)",
         "Lightweight v1 baseline.",
         "Weaker on NEUTRAL class vs v2."],
        ["distilroberta_hc",
         "Good (78–85%)",
         "RoBERTa v1 baseline.",
         "Weaker on NEUTRAL class vs v2."],
        ["default (DistilBERT SST-2)",
         "Moderate (70–80%)",
         "Very fast. Good binary polarity detection.",
         "No NEUTRAL class. Trained on movie reviews, not healthcare."],
        ["roberta (BERT Multilingual)",
         "Moderate (65–78%)",
         "Unique 5-star output useful for rating-style reporting.",
         "Coarse granularity. Star bands need interpretation."],
        ["emotion (GoEmotions)",
         "High for emotions (85%+)",
         "Only model that detects 7 distinct emotions. Invaluable for patient distress detection.",
         "Different label set — not directly comparable to sentiment models."],
        ["amazon (Amazon Reviews)",
         "Moderate (70–78%)",
         "Good on review-style short feedback.",
         "No NEUTRAL class. Review domain differs from clinical."],
        ["twitter (RoBERTa Twitter)",
         "Moderate (72–80%)",
         "Handles informal language and abbreviations.",
         "Trained on tweets — style mismatch with formal clinical notes."],
        ["sst2 (BERT SST-2)",
         "Moderate (70–78%)",
         "Standard NLP benchmark model. Good for comparison.",
         "No NEUTRAL class. Movie review domain."],
        ["zeroshot (BART MNLI)",
         "Moderate (65–78%)",
         "Most flexible — no training needed. Can classify with any labels.",
         "Slowest model. Lower accuracy than fine-tuned HC models."],
    ],
    col_widths=[3.5, 3.5, 5, 4.5]
)

# ── 13.7 ──────────────────────────────────────────────────────────────────────
h2(doc, "13.7  Model Selection Guide")

body(doc, (
    "Use this table to choose the right model for your use case. "
    "You do not need to run all 13 models every time."
))

add_table(doc,
    ["Use Case", "Recommended model_type", "Why"],
    [
        ["Production — clinical notes, highest accuracy",
         "bert_hc_v2",
         "Best overall fine-tuned model for healthcare text. Handles all three sentiment classes."],
        ["Production — high volume, speed matters",
         "distilbert_hc_v2",
         "Fastest fine-tuned model with strong accuracy. Good accuracy-to-speed trade-off."],
        ["Detect patient distress / emotions",
         "emotion",
         "Only model returning 7 emotion classes (ANGER, FEAR, SADNESS, JOY, etc.)."],
        ["5-star rating output",
         "roberta",
         "Returns 1–5 star bands — intuitive for survey-style dashboards."],
        ["Informal short feedback",
         "twitter",
         "Trained on informal text; handles abbreviations and short sentences."],
        ["No fine-tuning available",
         "zeroshot",
         "Zero-shot — works on any domain without training data."],
        ["Benchmark comparison",
         "Run all 13 in evaluate_models.py",
         "Evaluate every model against your own patient feedback before picking one."],
        ["Maximum reliability via ensemble",
         "bert_hc_v2 + distilroberta_hc_v2 (majority vote)",
         "Two strong models agreeing gives high reliability on ambiguous texts."],
    ],
    col_widths=[4.5, 4, 8]
)

# ── 13.8 ──────────────────────────────────────────────────────────────────────
h2(doc, "13.8  Known Limitations")

body(doc, "Be aware of these limitations when using the evaluation results:")

add_table(doc,
    ["Limitation", "Details", "Workaround"],
    [
        ["Short or vague texts",
         "Texts under 4 words produce low-confidence, unreliable results and are rejected by the API.",
         "Filter out texts with fewer than 4 meaningful words before calling the API."],
        ["No NEUTRAL class in binary models",
         "default, amazon, and sst2 have no NEUTRAL class. Factual clinical statements "
         "may be classified as POSITIVE.",
         "Use HC fine-tuned models (bert_hc_v2 etc.) when NEUTRAL detection is important."],
        ["Emotion vs sentiment mismatch",
         "The emotion model uses a different label set and cannot be compared directly "
         "with sentiment models in accuracy tables.",
         "Evaluate the emotion model separately using emotion-specific ground-truth labels."],
        ["First-load latency",
         "Each model takes 5–30 seconds to load on first use. zeroshot (BART) takes the longest.",
         "Call GET /api/v1/warmup after server startup to pre-load all models."],
        ["Language",
         "All models are trained on English text only. Non-English text will be misclassified.",
         "Translate non-English texts to English before calling the API."],
        ["PII redaction may alter meaning",
         "OpenMed may replace tokens that are critical to sentiment. "
         "E.g. 'Dr Smith was brilliant' → '[NAME] was brilliant' — "
         "the sentiment is preserved, but very rare cases may differ.",
         "Always inspect redacted_text in the response to verify the redaction was appropriate."],
    ],
    col_widths=[3.5, 6.5, 6.5]
)

add_callout(doc, (
    "Always evaluate the models on a sample of your own real patient feedback data "
    "before deploying to production. The 10 evaluation cases here are a representative "
    "starting point — your actual data will have its own distribution and style."
), "warning")

# ── Save ──────────────────────────────────────────────────────────────────────
import os as _os
_os.makedirs("docs", exist_ok=True)
output_path = "docs/NLP_OpenMed_Sentiment_Analysis_API_Integration_Guide.docx"
doc.save(output_path)
print(f"✅  Saved: {output_path}")
