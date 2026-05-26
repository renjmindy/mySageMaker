"""
Build NLP_Healthcare_Sentiment_API_Guide.docx
Run: python3 docs/build_guide.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Colours ────────────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1A, 0x37, 0x6C)   # header / title
MID_BLUE   = RGBColor(0x25, 0x63, 0xAE)   # section headings
GREEN      = RGBColor(0x1E, 0x7E, 0x34)   # success / positive
RED        = RGBColor(0xC0, 0x39, 0x2B)   # error / negative
ORANGE     = RGBColor(0xD3, 0x7A, 0x00)   # warning / neutral
GREY_BG    = RGBColor(0xF4, 0xF4, 0xF4)   # code block background
LIGHT_BLUE = RGBColor(0xDB, 0xE9, 0xFF)   # table header background
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    hex_col = str(rgb)
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_col)
    tcPr.append(shd)

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for cell, w in zip(row.cells, widths_cm):
            cell.width = Cm(w)

def hr(doc):
    """Horizontal rule paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2563AE')
    pb.append(bottom)
    pPr.append(pb)
    p.paragraph_format.space_after = Pt(4)
    return p

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 11)
    run.font.color.rgb = DARK_BLUE if level == 1 else MID_BLUE
    return p

def body(doc, text, bold_parts=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if bold_parts:
        remaining = text
        for chunk in bold_parts:
            idx = remaining.find(chunk)
            if idx > 0:
                p.add_run(remaining[:idx])
            r = p.add_run(chunk)
            r.bold = True
            remaining = remaining[idx + len(chunk):]
        if remaining:
            p.add_run(remaining)
    else:
        p.add_run(text)
    return p

def bullet(doc, text, bold_part=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    if bold_part:
        idx = text.find(bold_part)
        if idx >= 0:
            p.add_run(text[:idx])
            p.add_run(bold_part).bold = True
            p.add_run(text[idx + len(bold_part):])
        else:
            p.add_run(text)
    else:
        p.add_run(text)
    return p

def note_box(doc, label, text, colour=ORANGE):
    """Coloured callout box."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, RGBColor(0xFF, 0xF8, 0xE7))
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f'{label}  ')
    r.bold = True
    r.font.color.rgb = colour
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def code_block(doc, lines, caption=None):
    """Monospaced grey code block."""
    if caption:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_after = Pt(2)
        cr = cp.add_run(caption)
        cr.bold = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = MID_BLUE

    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, GREY_BG)

    first = True
    for line in lines:
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.name = 'Courier New'
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(6)

def hdr_row(table, headers, bg=DARK_BLUE, fg=WHITE, widths_cm=None):
    row = table.rows[0]
    for cell, hdr in zip(row.cells, headers):
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(hdr)
        r.bold = True
        r.font.color.rgb = fg
        r.font.size = Pt(9)
    if widths_cm:
        set_col_widths(table, widths_cm)

def data_row(table, values, alt=False):
    row = table.add_row()
    bg = RGBColor(0xF0, 0xF5, 0xFF) if alt else WHITE
    for cell, val in zip(row.cells, values):
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        if isinstance(val, tuple):   # (text, bold, colour)
            text, bold, colour = val
            r = p.add_run(text)
            r.bold = bold
            if colour:
                r.font.color.rgb = colour
        else:
            r = p.add_run(str(val))
        r.font.size = Pt(9)


# ══════════════════════════════════════════════════════════════════════════════
#  BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Cover block ───────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(10)
title_p.paragraph_format.space_after  = Pt(4)
tr = title_p.add_run('NLP Healthcare Sentiment Analysis')
tr.bold = True
tr.font.size = Pt(20)
tr.font.color.rgb = DARK_BLUE

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_after = Pt(2)
sr = sub_p.add_run('API Quick-Start Guide')
sr.font.size = Pt(13)
sr.font.color.rgb = MID_BLUE

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_p.paragraph_format.space_after = Pt(14)
mr = meta_p.add_run('API region: ap-southeast-2  •  6 fine-tuned healthcare models  •  May 2026')
mr.font.size = Pt(9)
mr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
mr.italic = True

hr(doc)

# ── Section 1: What this API does ─────────────────────────────────────────────
heading(doc, '1.  What This API Does')
body(doc,
    'This API analyses a piece of clinical text — such as a patient note, discharge '
    'summary, or progress note — and returns a sentiment label (NEGATIVE, NEUTRAL, or '
    'POSITIVE) together with confidence scores for each label.')
body(doc,
    'Under the hood it runs one of six fine-tuned Transformer models, all trained '
    'specifically on healthcare text by cjen1008 on HuggingFace.',
    space_after=4)
body(doc, 'Labels explained:', bold_parts=['Labels explained:'])
bullet(doc, 'POSITIVE — patient is improving, symptoms resolving, responding well to treatment')
bullet(doc, 'NEUTRAL  — stable condition, no significant change, routine/factual observations')
bullet(doc, 'NEGATIVE — deteriorating, high pain, complications, distress, or poor prognosis')

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── Section 2: Prerequisites ──────────────────────────────────────────────────
heading(doc, '2.  Prerequisites')
body(doc, 'You only need two things before making your first call:')
bullet(doc, 'A terminal  (macOS Terminal, Windows PowerShell, Linux bash, or WSL2)')
bullet(doc, 'curl  installed  —  pre-installed on macOS and most Linux distros', bold_part='curl')

body(doc, 'Check if curl is installed:', bold_parts=['Check if curl is installed:'], space_after=2)
code_block(doc, ['curl --version'])
body(doc, 'You should see something like curl 8.x.x .... If not, install it:',
     bold_parts=['curl 8.x.x'])
bullet(doc, 'macOS:    brew install curl')
bullet(doc, 'Ubuntu:   sudo apt install curl')
bullet(doc, 'Windows:  curl ships with Windows 10/11 — open PowerShell and type curl --version')

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── Section 3: One-Time Setup ─────────────────────────────────────────────────
heading(doc, '3.  One-Time Setup — Save Your Credentials')
body(doc,
    'Copy and paste these two lines into your terminal. You only need to do this once '
    'per terminal session (they reset when you close the window).')

code_block(doc, [
    'API="https://ws1gaxke8i.execute-api.ap-southeast-2.amazonaws.com/prod"',
    'KEY="jRkqkFo7Ap6fbwMCjrgVt5pTk8ICLSXe90Fw8hYB"',
], caption='Paste into your terminal:')

note_box(doc,
    '⚠  Security note:',
    'Never commit the API key to Git or share it in a public channel. '
    'Treat it like a password. Contact Mindy if the key needs to be rotated.')

# ── Section 4: Available Endpoints ────────────────────────────────────────────
heading(doc, '4.  Available Endpoints')

tbl = doc.add_table(rows=1, cols=4)
tbl.style = 'Table Grid'
hdr_row(tbl, ['Method', 'Path', 'Auth required?', 'What it does'],
        widths_cm=[2.2, 5.0, 3.2, 7.0])
data_row(tbl, ['GET',  '/api/v1/health',  'Yes (X-Api-Key)', 'Check API is alive; lists loaded models'], alt=False)
data_row(tbl, ['GET',  '/api/v1/warmup',  'Yes (X-Api-Key)', 'Pre-loads all 6 models into memory (called automatically every 5 min)'], alt=True)
data_row(tbl, ['POST', '/api/v1/analyze', 'Yes (X-Api-Key)', 'Analyse text — returns sentiment label + probabilities'], alt=False)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ── Section 5: Available Models ───────────────────────────────────────────────
heading(doc, '5.  Available Models')
body(doc,
    'Pass one of the following values as model_type in your request body. '
    'If you omit it, bert_hc_v2 is used by default.',
    bold_parts=['model_type', 'bert_hc_v2'])

tbl2 = doc.add_table(rows=1, cols=3)
tbl2.style = 'Table Grid'
hdr_row(tbl2, ['model_type value', 'Display name', 'Notes'], widths_cm=[4.5, 5.5, 7.4])
models = [
    ('bert_hc_v2',          'BERT Healthcare v2',          'Default — best overall accuracy'),
    ('distilroberta_hc_v2', 'DistilRoBERTa Healthcare v2', 'Fast, slightly lower accuracy than BERT v2'),
    ('distilbert_hc_v2',    'DistilBERT Healthcare v2',    'Compact and quick'),
    ('bert_hc',             'BERT Healthcare',             'v1 model — native lowercase labels'),
    ('distilbert_hc',       'DistilBERT Healthcare',       'v1 compact model'),
    ('distilroberta_hc',    'DistilRoBERTa Healthcare',    'v1 RoBERTa-based model'),
]
for i, (key, display, note) in enumerate(models):
    data_row(tbl2, [key, display, note], alt=(i % 2 == 1))

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ── Section 6: Step-by-Step First Call ────────────────────────────────────────
heading(doc, '6.  Your First API Call  (Step by Step)')

body(doc, 'Step 1 — Health check', bold_parts=['Step 1 — Health check'])
body(doc, 'This confirms the API is running and shows which models are already loaded.',
     space_after=2)
code_block(doc, [
    'curl -s "$API/api/v1/health" \\',
    '  -H "X-Api-Key: $KEY" | python3 -m json.tool',
])
body(doc, 'Expected response:', bold_parts=['Expected response:'], space_after=2)
code_block(doc, [
    '{',
    '    "status": "ok",',
    '    "models_loaded": ["bert_hc_v2", "distilroberta_hc_v2", ...]',
    '}',
])

body(doc, 'Step 2 — Analyse a clinical note', bold_parts=['Step 2 — Analyse a clinical note'])
body(doc,
    'Write your JSON request to a file first — this avoids quote and newline '
    'problems in the terminal (explained in Section 8):',
    bold_parts=['Write your JSON request to a file first'])
code_block(doc, [
    "cat > /tmp/req.json << 'EOF'",
    '{"text":"Patient reports significant improvement in pain levels and is responding well to the new medication.","model_type":"bert_hc_v2"}',
    'EOF',
    '',
    'curl -s -X POST "$API/api/v1/analyze" \\',
    '  -H "Content-Type: application/json" \\',
    '  -H "X-Api-Key: $KEY" \\',
    '  --data @/tmp/req.json | python3 -m json.tool',
])

body(doc, 'Example response (trimmed):', bold_parts=['Example response (trimmed):'], space_after=2)
code_block(doc, [
    '{',
    '    "sentiment": "POSITIVE",',
    '    "labels":        ["NEGATIVE", "NEUTRAL", "POSITIVE"],',
    '    "probabilities": [0.049,       0.345,     0.606],',
    '    "model_type":    "bert_hc_v2",',
    '    "cleaned_text":  "patient report significant improvement pain level ...",',
    '    "tokenized_text": ["patient", "reports", "significant", ...],',
    '    "ner": [{"text": "new medication", "label": "TREATMENT"}],',
    '    "word_distribution": {"improvement": 1, "pain": 1, ...},',
    '    "total_words": 17',
    '}',
])

# ── Section 7: More Examples ───────────────────────────────────────────────────
heading(doc, '7.  More Examples')

body(doc, 'Negative clinical note:', bold_parts=['Negative clinical note:'], space_after=2)
code_block(doc, [
    "cat > /tmp/req.json << 'EOF'",
    '{"text":"Patient shows no improvement after two weeks of treatment. Pain has worsened significantly and she is refusing to eat.","model_type":"bert_hc_v2"}',
    'EOF',
    'curl -s -X POST "$API/api/v1/analyze" -H "Content-Type: application/json" -H "X-Api-Key: $KEY" --data @/tmp/req.json | python3 -m json.tool',
])

body(doc, 'Neutral / stable note:', bold_parts=['Neutral / stable note:'], space_after=2)
code_block(doc, [
    "cat > /tmp/req.json << 'EOF'",
    '{"text":"Vital signs remain stable. Blood pressure 128/82, heart rate 76 bpm. No acute distress noted. Plan: continue current medications and follow up in one week.","model_type":"bert_hc_v2"}',
    'EOF',
    'curl -s -X POST "$API/api/v1/analyze" -H "Content-Type: application/json" -H "X-Api-Key: $KEY" --data @/tmp/req.json | python3 -m json.tool',
])

body(doc, 'Test all 6 models in a loop:', bold_parts=['Test all 6 models in a loop:'], space_after=2)
code_block(doc, [
    "cat > /tmp/req.json << 'EOF'",
    '{"text":"Patient is recovering well and discharge is planned for tomorrow.","model_type":"PLACEHOLDER"}',
    'EOF',
    '',
    'for MODEL in bert_hc_v2 distilroberta_hc_v2 distilbert_hc_v2 bert_hc distilbert_hc distilroberta_hc; do',
    '  # update model_type in the file',
    '  sed -i "s/\\"model_type\\":\\"[^\\"]*\\"/\\"model_type\\":\\"$MODEL\\"/" /tmp/req.json',
    '  RESP=$(curl -s -X POST "$API/api/v1/analyze" \\',
    '    -H "Content-Type: application/json" -H "X-Api-Key: $KEY" --data @/tmp/req.json)',
    '  SENTIMENT=$(echo "$RESP" | python3 -c "import sys,json; print(json.load(sys.stdin)[\'sentiment\'])" 2>/dev/null)',
    '  printf "%-30s -> %s\\n" "$MODEL" "${SENTIMENT:-ERROR}"',
    'done',
])

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── Section 8: Common Errors & Fixes ──────────────────────────────────────────
heading(doc, '8.  Common Errors and How to Fix Them')

tbl3 = doc.add_table(rows=1, cols=3)
tbl3.style = 'Table Grid'
hdr_row(tbl3, ['Error / Symptom', 'Cause', 'Fix'], widths_cm=[5.0, 5.5, 6.9])

errors = [
    (
        '{"detail": "JSON decode error"} … "Invalid control character"',
        'Your text value wrapped onto a new line — JSON does not allow raw newlines inside strings.',
        'Keep the text value on ONE continuous line, or use the --data @file.json approach.',
    ),
    (
        'text:: command not found  or  model_type:: command not found',
        'Smart/curly quotes (‘ ’) replaced straight quotes in the curl command when copying from a document or browser.',
        'Re-type the single quotes manually, or use the --data @file.json approach to avoid quoting entirely.',
    ),
    (
        '{"message": "Forbidden"} or 403',
        'Missing or wrong X-Api-Key header.',
        'Make sure you ran the two export lines in Section 3, and that $KEY is set correctly (echo $KEY to check).',
    ),
    (
        '{"message": "Endpoint request timed out"}',
        'The Lambda container was cold (idle for >5 min before the warmup rule fired).',
        'Wait 10 seconds and retry. The EventBridge warmup rule fires every 5 min; once warm, requests take ~1 second.',
    ),
    (
        '"msg": "Text must be at least 4 words"  (422)',
        'The text field has fewer than 4 words.',
        'Use a complete clinical sentence (minimum 4 words).',
    ),
    (
        '"msg": "Text must not exceed 300 words"  (422)',
        'The text field has more than 300 words.',
        'Trim or split the note into shorter passages.',
    ),
    (
        '"msg": "model_type must be one of …"  (422)',
        'You passed a model_type string that is not in the allowed list.',
        'Use exactly one of the six values from Section 5, e.g. bert_hc_v2.',
    ),
]
for i, (err, cause, fix) in enumerate(errors):
    data_row(tbl3, [err, cause, fix], alt=(i % 2 == 1))

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ── Section 9: Understanding the Response ────────────────────────────────────
heading(doc, '9.  Understanding the Response Fields')

tbl4 = doc.add_table(rows=1, cols=3)
tbl4.style = 'Table Grid'
hdr_row(tbl4, ['Field', 'Type', 'Meaning'], widths_cm=[4.2, 2.8, 10.4])

fields = [
    ('sentiment',        'string',       'Top predicted label: NEGATIVE, NEUTRAL, or POSITIVE'),
    ('labels',           'list[string]', 'All three label names in score order'),
    ('probabilities',    'list[float]',  'Confidence scores 0–1 for each label (sums to ~1.0)'),
    ('model_type',       'string',       'Which model processed the request'),
    ('cleaned_text',     'string',       'Text after stop-word removal and normalisation'),
    ('tokenized_text',   'list[string]', 'Individual word tokens'),
    ('lemmatized_text',  'list[string]', 'Tokens reduced to their root form (e.g. "running" → "run")'),
    ('ner',              'list[object]', 'Named entities detected (text + label, e.g. TREATMENT, CONDITION)'),
    ('word_distribution','object',       'Per-word sentiment score counts'),
    ('total_words',      'integer',      'Number of tokens after cleaning'),
]
for i, row_data in enumerate(fields):
    data_row(tbl4, list(row_data), alt=(i % 2 == 1))

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ── Section 10: Quick Reference Card ─────────────────────────────────────────
heading(doc, '10.  Quick Reference Card')

body(doc, 'Setup (once per terminal session):', bold_parts=['Setup (once per terminal session):'], space_after=2)
code_block(doc, [
    'API="https://ws1gaxke8i.execute-api.ap-southeast-2.amazonaws.com/prod"',
    'KEY="jRkqkFo7Ap6fbwMCjrgVt5pTk8ICLSXe90Fw8hYB"',
])

body(doc, 'Health check:', bold_parts=['Health check:'], space_after=2)
code_block(doc, ['curl -s "$API/api/v1/health" -H "X-Api-Key: $KEY" | python3 -m json.tool'])

body(doc, 'Analyse text (recommended pattern):', bold_parts=['Analyse text (recommended pattern):'], space_after=2)
code_block(doc, [
    "cat > /tmp/req.json << 'EOF'",
    '{"text":"YOUR CLINICAL NOTE HERE","model_type":"bert_hc_v2"}',
    'EOF',
    'curl -s -X POST "$API/api/v1/analyze" -H "Content-Type: application/json" -H "X-Api-Key: $KEY" --data @/tmp/req.json | python3 -m json.tool',
])

body(doc, 'Six model_type values you can use:', bold_parts=['Six model_type values you can use:'], space_after=2)
for m in ['bert_hc_v2 (default)', 'distilroberta_hc_v2', 'distilbert_hc_v2',
          'bert_hc', 'distilbert_hc', 'distilroberta_hc']:
    bullet(doc, m)

note_box(doc,
    '💡 Tip:',
    'All six models are kept warm by an automated warmup that runs every 5 minutes. '
    'You should not see cold-start timeouts during normal working hours.',
    colour=MID_BLUE)

# ── Footer note ───────────────────────────────────────────────────────────────
hr(doc)
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_p.paragraph_format.space_before = Pt(4)
fr = footer_p.add_run(
    'Questions? Contact Mindy  •  '
    'Models: cjen1008 on HuggingFace  •  '
    'Deployed: AWS Lambda ap-southeast-2  •  May 2026'
)
fr.font.size = Pt(8)
fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
fr.italic = True

# ── Save ──────────────────────────────────────────────────────────────────────
out = os.path.join(os.path.dirname(__file__), 'NLP_Healthcare_Sentiment_API_Guide.docx')
doc.save(out)
print(f'Saved → {out}')
