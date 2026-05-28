"""
Generate: Healthcare PREM Topic Classifier — Local Docker Deployment Guide
Target audience: Junior software engineer (new graduate, March 2026)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

NAVY      = RGBColor(0x1A, 0x3A, 0x5C)
TEAL      = RGBColor(0x00, 0x7A, 0x87)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MID_GRAY  = RGBColor(0x66, 0x66, 0x66)
GREEN     = RGBColor(0x1E, 0x7E, 0x34)
ORANGE    = RGBColor(0xD3, 0x6A, 0x00)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
CODE_BG   = RGBColor(0xF4, 0xF6, 0xF8)
HEADER_BG = RGBColor(0x1A, 0x3A, 0x5C)

def shade_cell(cell, rgb):
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tc_pr.append(shd)

def shade_row(row, rgb):
    for cell in row.cells:
        shade_cell(cell, rgb)

def add_callout(doc, text, style="tip"):
    icons  = {"tip": "💡 Tip", "warning": "⚠️  Important", "note": "📝 Note"}
    colors = {"tip": GREEN, "warning": ORANGE, "note": TEAL}
    label  = icons.get(style, "📝 Note")
    color  = colors.get(style, TEAL)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(f"{label}:  ")
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.color.rgb = DARK_GRAY
    run2.font.size = Pt(10)
    return p

def add_code(doc, code_text, lang=""):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, CODE_BG)
    cell.paragraphs[0].clear()
    lines = code_text.strip().split("\n")
    for i, line in enumerate(lines):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        run = para.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    doc.add_paragraph()

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(18)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = TEAL
        run.font.size = Pt(14)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def h3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(11.5)
        run.bold = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    return p

def body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_GRAY
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_GRAY
    return p

def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_GRAY
    return p

def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0]
    shade_row(hdr, HEADER_BG)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9.5)
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        if ri % 2 == 1:
            shade_row(row, RGBColor(0xF0, 0xF4, 0xF8))
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9.5)
            run.font.color.rgb = DARK_GRAY
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

# =============================================================================
# TITLE PAGE
# =============================================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run("Healthcare PREM Topic Classifier")
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = NAVY

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Local Docker Deployment Guide")
run2.font.size = Pt(22)
run2.font.bold = True
run2.font.color.rgb = TEAL

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("For Junior Software Engineers")
run3.font.size = Pt(13)
run3.font.color.rgb = MID_GRAY
run3.font.italic = True

doc.add_paragraph()
doc.add_paragraph()

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
for label, val in [("Version", "2.0"), ("Date", "May 2026"), ("Environment", "Local — no cloud account required")]:
    r = p4.add_run(f"{label}:  ")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = NAVY
    r2 = p4.add_run(f"{val}     ")
    r2.font.size = Pt(10)
    r2.font.color.rgb = DARK_GRAY

doc.add_page_break()

# =============================================================================
# 1. OVERVIEW
# =============================================================================
h1(doc, "1. Overview")

body(doc, (
    "This guide explains how to run the Healthcare PREM Topic Classifier on your own computer "
    "using Docker. You will load a pre-built Docker image, start the API, and use it to "
    "automatically classify patient feedback text into one of 15 predefined PREM "
    "(Patient Reported Experience Measure) topic categories."
))
body(doc, (
    "The API works entirely offline — it does not connect to AWS, the internet, or any cloud "
    "service. All six NLP models and their weights are packaged inside the Docker image and "
    "ready to use immediately after loading."
))

add_callout(doc, (
    "You do not need an AWS account, API key, or internet connection after you have loaded "
    "the Docker image. This guide is fully self-contained."
), "note")

h2(doc, "1.1  What the API Does")
for item in [
    "Accepts one or more patient feedback texts in a single request (up to 500).",
    "Classifies each text into one of 15 predefined PREM topic categories using up to six NLP models.",
    "Returns the predicted topic name, a confidence score, and the top N matching topics per model.",
    "Supports BERTopic (transformer-based, zero-shot) and three classical models pre-trained on "
    "a 225-sentence PREM seed corpus: LDA, LSI, and NMF.",
]:
    bullet(doc, item)

h2(doc, "1.2  The 15 PREM Topic Categories")
body(doc, (
    "Every piece of patient feedback is classified into one of these 15 predefined categories:"
))
add_table(doc,
    ["ID", "Topic Name", "Example Words / Phrases"],
    [
        ["0",  "Clinical Staff",                          "doctor, nurse, GP, specialist, physiotherapist, psychologist, midwife"],
        ["1",  "Administrative Staff",                    "reception, booking, billing, front desk"],
        ["2",  "Efficiency and Flow",                     "waiting, waited, queue, delay, on time, late"],
        ["3",  "Booking & Access",                        "book, appointment, portal, online, Zedoc, couldn't access"],
        ["4",  "Information & Shared Decision-Making",    "explained, told, informed, listened, dismissed, didn't listen"],
        ["5",  "Dignity & Respect",                       "respectful, rude, kind, ignored, valued, humiliated, courteous"],
        ["6",  "Care Quality",                            "thorough, careful, rushed, excellent, poor, skilled"],
        ["7",  "Care Coordination",                       "handover, referral, transferred, GP, coordinated"],
        ["8",  "Results & Information",                   "results, report, pathology, test, didn't understand"],
        ["9",  "Facilities & Environment & Hygiene",      "parking, car park, clean, comfortable, building"],
        ["10", "Telehealth",                              "telehealth, video, online consult, remote, phone"],
        ["11", "Emotional Support",                       "anxious, reassured, felt safe, valued, dismissed"],
        ["12", "Cultural & Accessibility Needs",          "cultural, language, disability, female practitioner"],
        ["13", "Discharge & Follow-up",                   "discharge, follow-up, home, after care, instructions"],
        ["14", "Feeling in Good Hands",                   "confident, trust, safe, overall positive, highly recommend"],
    ],
    col_widths=[0.8, 5, 10.7]
)

h2(doc, "1.3  Key Terms")
add_table(doc,
    ["Term", "What it means"],
    [
        ["Docker",       "A tool that packages an application and everything it needs into a portable 'container'."],
        ["Docker Image", "The packaged snapshot — like a zip file. Load it once and run it many times."],
        ["Container",    "A running instance of a Docker image."],
        ["API",          "A way for your code to send data to the application and receive results back over HTTP."],
        ["REST API",     "A common style of API using HTTP methods (GET, POST) that returns JSON data."],
        ["JSON",         "JavaScript Object Notation — a text format for structured data: {\"key\": \"value\"}."],
        ["Endpoint",     "A specific URL path you send a request to. E.g. /classify is the main endpoint."],
        ["Port",         "A number identifying a channel on your computer. This API runs on port 8001."],
        ["PREM",         "Patient Reported Experience Measure — a framework of 15 healthcare topic categories."],
        ["Confidence",   "A score from 0.0 to 1.0 showing how sure the model is. 0.9 = very confident."],
    ],
    col_widths=[3.5, 13]
)

doc.add_page_break()

# =============================================================================
# 2. PREREQUISITES
# =============================================================================
h1(doc, "2. Prerequisites")

body(doc, "You only need Docker installed. No Python, no libraries, no internet connection after setup.")

h2(doc, "2.1  Install Docker")
add_table(doc,
    ["Operating System", "How to Install"],
    [
        ["Windows 10 / 11",        "Download Docker Desktop from docker.com/products/docker-desktop. Run the installer and restart when prompted. Wait for the whale icon in the taskbar to turn green."],
        ["macOS (Intel / Apple Silicon)", "Download the .dmg for your chip type. Drag Docker to Applications. Open it and wait for the whale icon in the menu bar to be steady."],
        ["Ubuntu / Debian Linux",  "Run: sudo apt-get update && sudo apt-get install -y docker.io && sudo systemctl start docker && sudo usermod -aG docker $USER — then log out and back in."],
    ],
    col_widths=[4, 12.5]
)

add_callout(doc,
    "Verify installation by opening a terminal and running: docker --version\n"
    "You should see something like: Docker version 26.x.x",
    "tip"
)

h2(doc, "2.2  Receive the Docker Image File")
body(doc, (
    "You will receive a file called nlp-topic-api.tar.gz from your team (approximately 1.3 GB). "
    "Save it anywhere on your computer. Do NOT unzip it — Docker loads it directly."
))

doc.add_page_break()

# =============================================================================
# 3. QUICK START
# =============================================================================
h1(doc, "3. Quick Start (10 minutes)")

h2(doc, "Step 1 — Open a Terminal")
add_table(doc,
    ["Operating System", "How to open a terminal"],
    [
        ["Windows", "Press Win + R, type cmd, press Enter. Or search for 'PowerShell'."],
        ["macOS",   "Press Cmd + Space, type Terminal, press Enter."],
        ["Linux",   "Press Ctrl + Alt + T, or right-click the desktop and select 'Open Terminal'."],
    ],
    col_widths=[3.5, 13]
)

h2(doc, "Step 2 — Load the Docker Image")
body(doc, "Run this command once, replacing the path with wherever you saved the file:")
add_code(doc, """\
# Windows example:
docker load < C:\\Users\\YourName\\Downloads\\nlp-topic-api.tar.gz

# Mac / Linux example:
docker load < ~/Downloads/nlp-topic-api.tar.gz\
""", "bash")
body(doc, "You should see: Loaded image: nlp-topic-api:latest")
add_callout(doc,
    "Loading takes 1–3 minutes. The terminal shows no progress bar — just wait.",
    "note"
)

h2(doc, "Step 3 — Start the API")
add_code(doc, "docker run --rm -p 8001:8001 nlp-topic-api", "bash")
body(doc, "Leave this terminal open. You should see the startup message end with:")
add_code(doc, "INFO:     Application startup complete.\nINFO:     Uvicorn running on http://0.0.0.0:8001", "")
add_callout(doc,
    "Leave this terminal open while using the API. Press Ctrl+C to stop.",
    "warning"
)

h2(doc, "Step 4 — Verify the API is Running")
body(doc, "Open a second terminal and run:")
add_code(doc, "curl http://localhost:8001/health", "bash")
body(doc, "You should see all six models loaded:")
add_code(doc, """\
{
    "status": "ok",
    "models_loaded": {
        "bertopic_mini":  true,
        "bertopic_mpnet": true,
        "lda":            true,
        "lsi":            true,
        "hdp":            true,
        "nmf":            true
    }
}\
""", "json")

add_callout(doc,
    "You can also open http://localhost:8001/docs in your browser for the interactive Swagger UI — "
    "try the API without writing any code.",
    "tip"
)

doc.add_page_break()

# =============================================================================
# 4. API REFERENCE
# =============================================================================
h1(doc, "4. API Reference")

h2(doc, "4.1  Base URL")
body(doc, "All requests go to your local machine on port 8001:")
add_code(doc, "http://localhost:8001", "")

add_callout(doc,
    "No API key is required for the local Docker version. "
    "Authentication (x-api-key) is only needed for the cloud AWS version.",
    "note"
)

h2(doc, "4.2  Endpoints")
add_table(doc,
    ["Method", "Path", "Description"],
    [
        ["GET",  "/health",   "Check that the API is running and all models are loaded."],
        ["POST", "/classify", "Classify one or more patient feedback texts into PREM topics."],
        ["GET",  "/docs",     "Interactive Swagger UI — try the API in your browser."],
    ],
    col_widths=[2, 3.5, 11]
)

h2(doc, "4.3  GET /health")
body(doc, "Use this to confirm the API is ready before sending data.")
add_code(doc, "curl http://localhost:8001/health", "bash")
body(doc, "Response (HTTP 200):")
add_code(doc, """\
{
    "status": "ok",
    "models_loaded": {
        "bertopic_mini":  true,
        "bertopic_mpnet": true,
        "lda":            true,
        "lsi":            true,
        "hdp":            true,
        "nmf":            true
    }
}\
""", "json")

h2(doc, "4.4  POST /classify")
body(doc, (
    "This is the main endpoint. Send patient feedback text(s) and receive PREM topic "
    "classifications from all requested models."
))

h3(doc, "Request Body Fields")
add_table(doc,
    ["Field", "Type", "Required?", "Default", "Description"],
    [
        ["texts",  "list of strings", "Yes*", "—",      "List of feedback texts to classify. Min 1, max 500."],
        ["text",   "string",          "Yes*", "—",       "Shortcut for a single text — wrapped in a list automatically."],
        ["models", "list of strings", "No",   "[\"all\"]","Which models to run. Use [\"all\"] or list specific model keys."],
        ["top_n",  "integer",         "No",   "3",       "How many top topics to include per result. Min 1, max 15."],
    ],
    col_widths=[2.5, 3.5, 2.5, 2.5, 5.5]
)
add_callout(doc, "Use either \"text\" or \"texts\" — not both.", "note")

h3(doc, "Available Models")
add_table(doc,
    ["Model Key", "Full Name", "Speed", "Training"],
    [
        ["bertopic_mini",  "BERTopic + MiniLM-L6-v2",         "Fastest", "Zero-shot cosine similarity on 15 PREM topic descriptions"],
        ["bertopic_mpnet", "BERTopic + MPNet-base-v2",         "Medium",  "Zero-shot cosine similarity on 15 PREM topic descriptions"],
        ["lda",            "Latent Dirichlet Allocation",       "Fast",    "Pre-trained on 225-sentence PREM seed corpus (15 topics)"],
        ["lsi",            "Latent Semantic Indexing",          "Fast",    "Pre-trained on 225-sentence PREM seed corpus (15 topics)"],
        ["hdp",            "Hierarchical Dirichlet Process",    "Medium",  "Pre-trained on 225-sentence PREM seed corpus (auto-aligned)"],
        ["nmf",            "Non-negative Matrix Factorization", "Fast",    "Pre-trained on 225-sentence PREM seed corpus (15 topics)"],
    ],
    col_widths=[3.5, 5, 2, 6]
)
add_callout(doc,
    "LDA, LSI, HDP, and NMF are pre-trained at Docker build time on a balanced seed corpus of "
    "225 patient feedback sentences (15 per PREM topic). Their internal topics are automatically "
    "aligned to the 15 PREM categories using keyword-overlap scoring. No training data is needed from you.",
    "note"
)

h3(doc, "Response Fields")
add_table(doc,
    ["Field", "Type", "Description"],
    [
        ["results",                       "list",   "One entry per input text, in the same order sent."],
        ["results[i].text",               "string", "The original input text (echoed back)."],
        ["results[i].classifications",    "object", "One classification result per requested model."],
        ["classifications.<model>.topic_id",   "integer","Numeric ID of the best-matching PREM topic (0–14)."],
        ["classifications.<model>.topic_name", "string", "Human-readable name of the best-matching topic."],
        ["classifications.<model>.confidence", "float",  "Confidence score 0.0–1.0. Higher = more certain."],
        ["classifications.<model>.top_topics", "list",   "Top N topics with their IDs, names, and scores."],
        ["model_status",                  "object", "Whether each model was loaded and ready."],
        ["elapsed_sec",                   "float",  "Total classification time in seconds."],
    ],
    col_widths=[5, 2.5, 9]
)

doc.add_page_break()

# =============================================================================
# 5. CODE EXAMPLES
# =============================================================================
h1(doc, "5. Code Examples")
body(doc, "Make sure the container is running (Step 3) before trying these. Use a second terminal.")

h2(doc, "5.1  cURL (Command Line)")

h3(doc, "Health check")
add_code(doc, "curl http://localhost:8001/health", "bash")

h3(doc, "Classify a single text — all 6 models")
add_code(doc, """\
curl -s -X POST http://localhost:8001/classify \\
  -H "Content-Type: application/json" \\
  -d '{
    "text": "The nurse was incredibly kind and made me feel completely at ease.",
    "models": ["all"],
    "top_n": 3
  }' | python3 -m json.tool\
""", "bash")

h3(doc, "Classify multiple texts at once")
add_code(doc, """\
curl -s -X POST http://localhost:8001/classify \\
  -H "Content-Type: application/json" \\
  -d '{
    "texts": [
      "The doctor was very thorough and answered all my questions.",
      "I waited over two hours past my appointment time.",
      "My telehealth video call worked perfectly and saved me a long drive.",
      "The discharge instructions were unclear about my medications."
    ],
    "models": ["all"],
    "top_n": 2
  }' | python3 -m json.tool\
""", "bash")

h3(doc, "Classify using a specific model only (faster)")
add_code(doc, """\
curl -s -X POST http://localhost:8001/classify \\
  -H "Content-Type: application/json" \\
  -d '{
    "text": "The receptionist was rude and unhelpful when I arrived.",
    "models": ["bertopic_mini"],
    "top_n": 3
  }' | python3 -m json.tool\
""", "bash")

h2(doc, "5.2  Python")
body(doc, "Install the requests library first if you don't have it:")
add_code(doc, "pip install requests", "bash")

h3(doc, "Basic example — classify and print results")
add_code(doc, """\
import requests

BASE_URL = "http://localhost:8001"

# 1. Health check
resp = requests.get(f"{BASE_URL}/health")
print("Health:", resp.json())

# 2. Classify a single text with all 6 models
payload = {
    "text":   "The nurse was incredibly supportive and made me feel at ease.",
    "models": ["all"],
    "top_n":  3,
}
resp = requests.post(f"{BASE_URL}/classify", json=payload)
resp.raise_for_status()

data = resp.json()
for result in data["results"]:
    print(f"\\nText: {result['text']}")
    for model, clf in result["classifications"].items():
        print(f"  {model:20s} -> {clf['topic_name']} (confidence: {clf['confidence']:.2f})")\
""", "python")

h3(doc, "Classify many texts and save to CSV")
add_code(doc, """\
import requests
import csv

BASE_URL = "http://localhost:8001"

feedback_list = [
    "Dr Smith was very thorough and answered all my questions.",
    "I waited over two hours past my appointment time.",
    "The Zedoc portal kept crashing when I tried to book online.",
    "My telehealth video call worked perfectly and saved me a long drive.",
    "I was discharged without clear instructions about my medications.",
]

resp = requests.post(
    f"{BASE_URL}/classify",
    json={"texts": feedback_list, "models": ["bertopic_mini"], "top_n": 3},
)
resp.raise_for_status()

with open("prem_results.csv", "w", newline="", encoding="utf-8") as f:
    writer = csv.writer(f)
    writer.writerow(["text", "topic_id", "topic_name", "confidence"])
    for item in resp.json()["results"]:
        best = item["classifications"]["bertopic_mini"]
        writer.writerow([
            item["text"],
            best["topic_id"],
            best["topic_name"],
            round(best["confidence"], 4),
        ])

print(f"Saved {len(feedback_list)} rows to prem_results.csv")\
""", "python")

h3(doc, "Majority vote across all 6 models")
add_code(doc, """\
import requests
from collections import Counter

resp = requests.post(
    "http://localhost:8001/classify",
    json={
        "texts": [
            "I was very anxious before surgery but the nurse reassured me.",
            "The parking was very difficult and the signage was confusing.",
        ],
        "models": ["all"],
        "top_n": 1,
    },
)
resp.raise_for_status()

for item in resp.json()["results"]:
    votes = [clf["topic_name"] for clf in item["classifications"].values()]
    winner, count = Counter(votes).most_common(1)[0]
    print(f"Text  : {item['text'][:70]}")
    print(f"Winner: {winner}  ({count}/6 models agreed)")
    print()\
""", "python")

h2(doc, "5.3  JavaScript / Node.js")
body(doc, "Using the built-in fetch API (Node.js 18+). No extra packages needed.")
add_code(doc, """\
const BASE_URL = "http://localhost:8001";

async function classifyFeedback(texts, models = ["all"], topN = 3) {
  const response = await fetch(`${BASE_URL}/classify`, {
    method:  "POST",
    headers: { "Content-Type": "application/json" },
    body:    JSON.stringify({ texts, models, top_n: topN }),
  });
  if (!response.ok) throw new Error(`API error: ${response.status}`);
  return response.json();
}

classifyFeedback([
  "The doctor was very kind and explained everything clearly.",
  "I waited over two hours past my appointment time.",
])
  .then(data => {
    data.results.forEach(item => {
      const best = item.classifications.bertopic_mini;
      console.log(`Text:  ${item.text}`);
      console.log(`Topic: ${best.topic_name} (confidence: ${best.confidence.toFixed(2)})`);
      console.log("---");
    });
  })
  .catch(err => console.error("Error:", err));\
""", "javascript")

doc.add_page_break()

# =============================================================================
# 6. UNDERSTANDING THE OUTPUT
# =============================================================================
h1(doc, "6. Understanding the Output")

h2(doc, "6.1  Full Example Response")
add_code(doc, """\
{
  "results": [
    {
      "text": "The nurse was incredibly kind and made me feel completely at ease.",
      "classifications": {
        "bertopic_mini": {
          "topic_id":   11,
          "topic_name": "Emotional Support",
          "confidence": 0.74,
          "top_topics": [
            { "topic_id": 11, "topic_name": "Emotional Support",  "score": 0.74 },
            { "topic_id":  0, "topic_name": "Clinical Staff",     "score": 0.70 },
            { "topic_id":  6, "topic_name": "Care Quality",       "score": 0.69 }
          ]
        },
        "lda": {
          "topic_id":   11,
          "topic_name": "Emotional Support",
          "confidence": 0.68,
          "top_topics": [ ... ]
        }
      }
    }
  ],
  "model_status": {
    "bertopic_mini": true, "bertopic_mpnet": true,
    "lda": true, "lsi": true, "hdp": true, "nmf": true
  },
  "elapsed_sec": 3.2
}\
""", "json")

h2(doc, "6.2  Reading the Results")
for item in [
    "topic_id is the numeric ID of the PREM category (0–14). See Section 1.2 for the full list.",
    "topic_name is the human-readable PREM category name.",
    "confidence is how certain the model is (0.0–1.0). Above 0.70 is generally reliable.",
    "top_topics shows the runner-up categories — useful when confidence is borderline.",
    "elapsed_sec is the total time to run all requested models.",
]:
    bullet(doc, item)

add_callout(doc,
    "If the top-1 confidence is below 0.60, consider showing the top-3 topics to the user "
    "so they can confirm the correct category.",
    "tip"
)

h2(doc, "6.3  Model Agreement (Majority Vote)")
body(doc, (
    "When using models: [\"all\"], you get 6 predictions per text. "
    "The more models that agree, the more reliable the result:"
))
add_table(doc,
    ["Models Agreeing", "Reliability", "Recommendation"],
    [
        ["5 or 6 / 6", "Very high",  "Use the top-1 prediction confidently."],
        ["3 or 4 / 6", "Good",       "Use top-1 but show top-3 as a safety net."],
        ["1 or 2 / 6", "Low",        "Show top-3 topics and ask the user to confirm."],
    ],
    col_widths=[3.5, 3.5, 9.5]
)

doc.add_page_break()

# =============================================================================
# 7. RUNNING OPTIONS
# =============================================================================
h1(doc, "7. Running the Container")

h2(doc, "7.1  Basic Run (foreground)")
add_code(doc, "docker run --rm -p 8001:8001 nlp-topic-api", "bash")

h2(doc, "7.2  Run in Background (detached)")
add_code(doc, """\
# Start in the background
docker run -d --name nlp-api -p 8001:8001 nlp-topic-api

# Check it is running
docker ps

# View logs
docker logs nlp-api

# Stop it
docker stop nlp-api\
""", "bash")

h2(doc, "7.3  Using docker-compose")
body(doc, "If you received a docker-compose.yml alongside the image:")
add_code(doc, """\
docker compose up -d     # start in background
docker compose ps        # check status
docker compose logs -f   # stream logs
docker compose down      # stop\
""", "bash")

h2(doc, "7.4  Using a Different Port")
body(doc, "If port 8001 is already in use, map to a different port:")
add_code(doc, "docker run --rm -p 9000:8001 nlp-topic-api", "bash")
body(doc, "Then send requests to http://localhost:9000 instead.")

doc.add_page_break()

# =============================================================================
# 8. TROUBLESHOOTING
# =============================================================================
h1(doc, "8. Troubleshooting")

add_table(doc,
    ["Problem", "Likely Cause", "Solution"],
    [
        ["docker: command not found",
         "Docker is not installed.",
         "Install Docker Desktop (Section 2.1) and restart your terminal."],
        ["docker load takes a long time with no output",
         "Normal — the file is ~1.3 GB.",
         "Wait 1–3 minutes. Do NOT press Ctrl+C."],
        ["Port 8001 is already in use",
         "Another program is using port 8001.",
         "Use a different port: docker run --rm -p 9000:8001 nlp-topic-api"],
        ["curl: connection refused",
         "The container is not running.",
         "Run docker ps to check. If not listed, start it again."],
        ["Health check shows a model as false",
         "A model failed to load on startup.",
         "Check docker logs nlp-api for the error. Restart the container."],
        ["422 Unprocessable Entity",
         "Missing or invalid request field.",
         "Check that 'text' or 'texts' is present. Check model key spellings."],
        ["First request takes 20–30 seconds",
         "Normal — models load into memory on first request.",
         "Subsequent requests will be much faster. Call /health first to pre-warm."],
        ["docker run: image not found",
         "The image was not loaded yet.",
         "Run docker images to list loaded images. Repeat the docker load step if missing."],
    ],
    col_widths=[4, 4, 8.5]
)

h2(doc, "8.1  Useful Diagnostic Commands")
add_code(doc, """\
docker ps                          # list running containers
docker images                      # list loaded images
docker logs --tail 50 nlp-api      # last 50 lines of logs
docker logs -f nlp-api             # stream live logs (Ctrl+C to stop)

# Check what is using port 8001 (Linux / Mac)
lsof -i :8001

# Check what is using port 8001 (Windows PowerShell)
netstat -ano | findstr :8001\
""", "bash")

doc.add_page_break()

# =============================================================================
# 9. QUICK REFERENCE CARD
# =============================================================================
h1(doc, "9. Quick Reference Card")
body(doc, "Keep this page handy for day-to-day use.")

h2(doc, "Container Commands")
add_table(doc,
    ["Task", "Command"],
    [
        ["Load image (once only)",    "docker load < nlp-topic-api.tar.gz"],
        ["Start (foreground)",        "docker run --rm -p 8001:8001 nlp-topic-api"],
        ["Start (background)",        "docker run -d --name nlp-api -p 8001:8001 nlp-topic-api"],
        ["Start with docker-compose", "docker compose up -d"],
        ["Stop (background)",         "docker stop nlp-api"],
        ["View logs",                 "docker logs nlp-api"],
        ["List running containers",   "docker ps"],
    ],
    col_widths=[5, 11.5]
)

h2(doc, "API Endpoints (base: http://localhost:8001)")
add_table(doc,
    ["Method", "Path", "Use For"],
    [
        ["GET",  "/health",   "Check API is alive and all 6 models are loaded."],
        ["POST", "/classify", "Classify patient feedback into 15 PREM topics."],
        ["GET",  "/docs",     "Interactive Swagger UI — test in browser."],
    ],
    col_widths=[2, 3.5, 11]
)

h2(doc, "Available Models")
add_table(doc,
    ["Key", "Description"],
    [
        ["bertopic_mini",  "BERTopic + MiniLM — fastest, zero-shot PREM classification"],
        ["bertopic_mpnet", "BERTopic + MPNet — higher accuracy, zero-shot"],
        ["lda",            "LDA — pre-trained on 15 PREM topics, interpretable"],
        ["lsi",            "LSI — pre-trained on 15 PREM topics, deterministic"],
        ["hdp",            "HDP — pre-trained on PREM seed corpus, auto-aligned"],
        ["nmf",            "NMF — pre-trained on 15 PREM topics, good for short texts"],
    ],
    col_widths=[3.5, 13]
)

h2(doc, "Minimal cURL — classify one text")
add_code(doc, """\
curl -s -X POST http://localhost:8001/classify \\
  -H "Content-Type: application/json" \\
  -d '{"text": "YOUR FEEDBACK HERE", "models": ["all"], "top_n": 3}' \\
  | python3 -m json.tool\
""", "bash")

h2(doc, "Minimal Python")
add_code(doc, """\
import requests
r = requests.post(
    "http://localhost:8001/classify",
    json={"text": "YOUR FEEDBACK HERE", "models": ["all"], "top_n": 3},
)
for model, clf in r.json()["results"][0]["classifications"].items():
    print(f"{model:20s} -> {clf['topic_name']} ({clf['confidence']:.2f})")\
""", "python")

h2(doc, "HTTP Status Codes")
add_table(doc,
    ["Code", "Meaning", "Fix"],
    [
        ["200", "Success",         "No action needed."],
        ["422", "Bad request",     "Check 'text' or 'texts' is present. Check model key spellings."],
        ["500", "Server error",    "Check docker logs nlp-api for the Python traceback."],
    ],
    col_widths=[1.8, 3.5, 11.2]
)

# =============================================================================
# Save
# =============================================================================
import os
os.makedirs("docs", exist_ok=True)
output_path = "docs/NLP_Topic_Analysis_Docker_Guide.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
