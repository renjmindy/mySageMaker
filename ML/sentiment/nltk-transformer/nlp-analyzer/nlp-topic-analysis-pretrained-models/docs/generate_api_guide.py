"""
Generate: NLP Healthcare Topic/Theme Analysis API Integration Guide
Target audience: Junior software engineer (new graduate)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY       = RGBColor(0x1A, 0x3A, 0x5C)   # headings
TEAL       = RGBColor(0x00, 0x7A, 0x87)   # sub-headings / accents
DARK_GRAY  = RGBColor(0x33, 0x33, 0x33)   # body text
MID_GRAY   = RGBColor(0x66, 0x66, 0x66)   # captions / notes
GREEN      = RGBColor(0x1E, 0x7E, 0x34)   # success / tip callout text
ORANGE     = RGBColor(0xD3, 0x6A, 0x00)   # warning callout text
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
CODE_BG    = RGBColor(0xF4, 0xF6, 0xF8)   # light code background (table shading)
HEADER_BG  = RGBColor(0x1A, 0x3A, 0x5C)   # table header background

# ── Helper: shade a table cell ────────────────────────────────────────────────
def shade_cell(cell, rgb: RGBColor):
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tc_pr.append(shd)

def shade_row(row, rgb: RGBColor):
    for cell in row.cells:
        shade_cell(cell, rgb)

# ── Helper: set paragraph border (callout box) ───────────────────────────────
def add_callout(doc, text, style="tip"):
    """Add a styled callout paragraph (tip / warning / note)."""
    icons   = {"tip": "💡 Tip", "warning": "⚠️  Important", "note": "📝 Note"}
    colors  = {"tip": GREEN, "warning": ORANGE, "note": TEAL}
    label   = icons.get(style, "📝 Note")
    color   = colors.get(style, TEAL)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(f"{label}:  ")
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.color.rgb = DARK_GRAY
    run2.font.size = Pt(10)
    return p

# ── Helper: add code block (table with shaded bg) ────────────────────────────
def add_code(doc, code_text, lang=""):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, CODE_BG)
    cell.paragraphs[0].clear()
    lines = code_text.strip().split("\n")
    for i, line in enumerate(lines):
        if i == 0:
            para = cell.paragraphs[0]
        else:
            para = cell.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        run = para.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    doc.add_paragraph()   # spacing after code block

# ── Helper: styled heading ────────────────────────────────────────────────────
def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(18)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = TEAL
        run.font.size = Pt(14)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def h3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(11.5)
        run.bold = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    return p

def body(doc, text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_GRAY
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_GRAY
    return p

def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_GRAY
    return p

def inline_code(run_or_para, text):
    """Add inline code-styled run to an existing paragraph."""
    if hasattr(run_or_para, 'add_run'):
        r = run_or_para.add_run(text)
    else:
        r = run_or_para
    r.font.name = "Courier New"
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
    return r

# ── Helper: standard table with navy header row ───────────────────────────────
def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hdr = t.rows[0]
    shade_row(hdr, HEADER_BG)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9.5)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        if ri % 2 == 1:
            shade_row(row, RGBColor(0xF0, 0xF4, 0xF8))
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9.5)
            run.font.color.rgb = DARK_GRAY
    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run("NLP Healthcare Topic / Theme Analysis")
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = NAVY

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("API Integration Guide")
run2.font.size = Pt(22)
run2.font.bold = True
run2.font.color.rgb = TEAL

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("For Junior Software Engineers")
run3.font.size = Pt(13)
run3.font.color.rgb = MID_GRAY
run3.font.italic = True

doc.add_paragraph()
doc.add_paragraph()

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
for label, val in [("Version", "1.0"), ("Date", "May 2026"), ("Environment", "AWS ap-southeast-2 (Production)")]:
    r = p4.add_run(f"{label}:  ")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = NAVY
    r2 = p4.add_run(f"{val}     ")
    r2.font.size = Pt(10)
    r2.font.color.rgb = DARK_GRAY

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "1. Overview")

body(doc, (
    "This guide explains how to integrate with the Healthcare Topic / Theme Analysis API. "
    "The API takes free-text patient feedback — such as comments left on a survey or online form — "
    "and automatically identifies which healthcare topic or theme the comment is about."
))
body(doc, (
    "For example, if a patient writes \"I waited two hours past my appointment time\", "
    "the API will classify that as the topic Efficiency and Flow. "
    "If they write \"The nurse was incredibly kind and reassuring\", it will be classified as "
    "Clinical Staff."
))

add_callout(doc, (
    "You do not need to understand how the machine learning models work internally. "
    "This guide focuses entirely on how to call the API and use the results in your application."
), "note")

h2(doc, "1.1  What the API Does")
for item in [
    "Accepts one or more patient feedback texts in a single request.",
    "Runs them through up to 6 different NLP (Natural Language Processing) models in parallel.",
    "Returns each text's most likely topic, a confidence score, and the top N matching topics.",
    "Supports 15 predefined PREM (Patient Reported Experience Measure) topic categories.",
]:
    bullet(doc, item)

h2(doc, "1.2  Key Terms")
add_table(doc,
    ["Term", "What it means"],
    [
        ["API", "Application Programming Interface — a way for your code to talk to another service over the internet."],
        ["REST API", "A common style of API that uses standard HTTP methods (GET, POST) and returns JSON data."],
        ["JSON", "JavaScript Object Notation — a text format for sending structured data. Looks like: {\"key\": \"value\"}."],
        ["Endpoint", "A specific URL you send a request to. Like a specific page on a website, but for data."],
        ["API Key", "A secret password string that proves your request is authorised. Must be sent in a header."],
        ["Header", "Extra metadata attached to your HTTP request (not part of the URL or body)."],
        ["Request Body", "The JSON data you send with a POST request — this is where your texts go."],
        ["Response", "The JSON data the API sends back after processing your request."],
        ["NLP", "Natural Language Processing — AI technology that understands human language."],
        ["PREM", "Patient Reported Experience Measure — a framework for categorising patient feedback."],
        ["Confidence", "A score from 0.0 to 1.0 showing how sure the model is. 0.9 = very confident, 0.5 = uncertain."],
    ],
    col_widths=[4, 12.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. QUICK START
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "2. Quick Start (5 minutes)")

body(doc, "Follow these four steps to make your first successful API call.")

h2(doc, "Step 1 — Note Your Credentials")
body(doc, "You will need these two values for every request you make:")

add_table(doc,
    ["Item", "Value"],
    [
        ["Base URL",        "https://5kzn638wee.execute-api.ap-southeast-2.amazonaws.com/prod"],
        ["API Key",         "Retrieve from AWS — see Step 2 below (never hardcode)"],
        ["Key Header Name", "x-api-key"],
        ["AWS Region",      "ap-southeast-2 (Sydney)"],
    ],
    col_widths=[4.5, 12]
)

add_callout(doc,
    "NEVER paste the API key directly into your source code or commit it to Git. "
    "Always retrieve it from AWS at runtime or store it in an environment variable. "
    "Hardcoded keys in code repositories are automatically detected by security scanners "
    "and must be rotated immediately.",
    "warning"
)

h2(doc, "Step 2 — Retrieve Your API Key from AWS & Set Shell Variables")
body(doc, (
    "The API key is stored securely in AWS API Gateway — not in any file. "
    "Run the following commands in your terminal to retrieve it and save it "
    "to shell variables for use in subsequent steps:"
))
add_code(doc, """\
# Retrieve the API key from AWS (requires AWS CLI configured with ap-southeast-2 access)
export HEALTHCARE_API_KEY=$(aws apigateway get-api-keys \\
  --region ap-southeast-2 \\
  --include-values \\
  --query "items[?contains(name,'healthcare-topic-classifier')].value | [0]" \\
  --output text)

# Set the base URL
export HEALTHCARE_API_BASE="$HEALTHCARE_API_BASE"

# Verify both are set
echo "Base URL : $HEALTHCARE_API_BASE"
echo "Key set  : ${HEALTHCARE_API_KEY:0:4}************************"  # shows only first 4 chars\
""", "bash")

add_callout(doc,
    "The echo command above intentionally shows only the first 4 characters of the key "
    "so you can confirm it loaded without revealing the full value in your terminal history.",
    "tip"
)

h2(doc, "Step 3 — Run a Health Check")
body(doc, (
    "A health check tells you the API is running and all models are loaded. "
    "Run this in your terminal:"
))
add_code(doc, """\
curl -s "$HEALTHCARE_API_BASE/health" \\
  -H "x-api-key: $HEALTHCARE_API_KEY" \\
  | python3 -m json.tool\
""", "bash")

body(doc, "You should see:")
add_code(doc, """\
{
    "status": "ok",
    "models_loaded": {
        "bertopic_mini":  true,
        "bertopic_mpnet": true,
        "lda":            true,
        "lsi":            true,
        "hdp":            true,
        "nmf":            true
    }
}\
""", "json")

add_callout(doc,
    "If you see \"status\": \"ok\" and all six models show true, the API is healthy and ready.",
    "tip"
)

h2(doc, "Step 4 — Classify Your First Text")
body(doc, "Send a piece of patient feedback to the /classify endpoint:")
add_code(doc, """\
curl -s -X POST "$HEALTHCARE_API_BASE/classify" \\
  -H "Content-Type: application/json" \\
  -H "x-api-key: $HEALTHCARE_API_KEY" \\
  -d '{
    "text": "The doctor was very kind and explained everything clearly.",
    "models": ["bertopic_mini"],
    "top_n": 3
  }' | python3 -m json.tool\
""", "bash")

body(doc, "The API will return the topic it detected, its name, and a confidence score. "
         "See Section 5 for the full response format.")

add_callout(doc,
    "Use \"text\" (singular) when sending one piece of feedback. "
    "Use \"texts\" (plural, as a list) when sending multiple at once. Both work.",
    "tip"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 3. API REFERENCE
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "3. API Reference")

h2(doc, "3.1  Base URL")
add_code(doc, "https://5kzn638wee.execute-api.ap-southeast-2.amazonaws.com/prod", "")

h2(doc, "3.2  Authentication")
body(doc, (
    "Every request must include your API key as an HTTP header. "
    "If the header is missing or incorrect, the API returns HTTP 403 Forbidden."
))
add_table(doc,
    ["Header Name", "Header Value", "Required?"],
    [
        ["x-api-key", "$HEALTHCARE_API_KEY  (retrieve from AWS — see Section 2, Step 2)", "Yes — on every request"],
        ["Content-Type", "application/json", "Yes — on POST requests"],
    ],
    col_widths=[4, 9, 3.5]
)

h2(doc, "3.3  Endpoints")
add_table(doc,
    ["Method", "Path", "Description", "Auth"],
    [
        ["GET",  "/health",   "Returns API status and which models are loaded.", "API Key"],
        ["POST", "/classify", "Classifies one or more texts into PREM topics.",  "API Key"],
    ],
    col_widths=[2, 3, 10, 2.5]
)

h2(doc, "3.4  GET /health")
body(doc, "Use this endpoint to check whether the API is running before sending data.")
body(doc, "Request:")
add_code(doc, """\
curl -s "$HEALTHCARE_API_BASE/health" \\
  -H "x-api-key: $HEALTHCARE_API_KEY"\
""", "bash")
body(doc, "Response (HTTP 200):")
add_code(doc, """\
{
    "status": "ok",
    "models_loaded": {
        "bertopic_mini":  true,
        "bertopic_mpnet": true,
        "lda":            true,
        "lsi":            true,
        "hdp":            true,
        "nmf":            true
    }
}\
""", "json")

h2(doc, "3.5  POST /classify")
body(doc, (
    "This is the main endpoint. Send patient feedback text(s) and receive topic classifications. "
    "All fields go inside the JSON request body."
))

h3(doc, "Request Body Fields")
add_table(doc,
    ["Field", "Type", "Required?", "Default", "Description"],
    [
        ["texts",   "list of strings", "Yes*", "—",       "A list of feedback texts to classify. Max 500 items per request."],
        ["text",    "string",          "Yes*", "—",       "Shortcut for a single text. The API wraps it in a list automatically."],
        ["models",  "list of strings", "No",   "[\"all\"]", "Which models to run. Use [\"all\"] for all 6, or list specific model names."],
        ["top_n",   "integer",         "No",   "3",       "How many top topics to include in each result. Min 1, sensible max is 5."],
    ],
    col_widths=[2.5, 3.5, 2.5, 2.5, 5.5]
)

add_callout(doc,
    "Use either \"text\" or \"texts\" — not both. "
    "\"text\" is a convenience shortcut for a single string. "
    "\"texts\" is a list and lets you classify many items in one API call.",
    "note"
)

h3(doc, "Available Models")
add_table(doc,
    ["Model Key", "Full Name", "Speed", "Best For"],
    [
        ["bertopic_mini",  "BERTopic + MiniLM-L6-v2",    "Fastest",  "General use, quick prototyping"],
        ["bertopic_mpnet", "BERTopic + MPNet-base-v2",   "Medium",   "Higher accuracy, production use"],
        ["lda",            "Latent Dirichlet Allocation", "Fast",     "Traditional NLP, interpretable results"],
        ["lsi",            "Latent Semantic Indexing",    "Fast",     "Keyword-based, deterministic output"],
        ["hdp",            "Hierarchical Dirichlet Process", "Medium", "Automatic topic count discovery"],
        ["nmf",            "Non-negative Matrix Factorization", "Fast", "Short texts, clear keyword topics"],
    ],
    col_widths=[3.5, 5.5, 2.5, 5]
)

add_callout(doc,
    "When in doubt, start with models: [\"all\"] to compare all 6 results side by side. "
    "Once you know which model performs best for your data, switch to that one to save processing time.",
    "tip"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. THE 15 PREM TOPICS
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "4. The 15 PREM Topic Categories")

body(doc, (
    "Every piece of patient feedback is classified into one of 15 predefined PREM "
    "(Patient Reported Experience Measure) topics. These cover the full spectrum of "
    "patient experience in a healthcare setting."
))

add_table(doc,
    ["ID", "Topic Name", "Example Patient Comment"],
    [
        ["0",  "Clinical Staff",                           "\"My GP Dr Smith was very thorough and spent time answering all my questions.\""],
        ["1",  "Administrative Staff",                     "\"The receptionist was very helpful and checked me in quickly.\""],
        ["2",  "Efficiency and Flow",                      "\"I waited over two hours past my appointment time with no explanation.\""],
        ["3",  "Booking & Access",                         "\"The Zedoc booking portal kept crashing when I tried to schedule online.\""],
        ["4",  "Information & Shared Decision-Making",     "\"The doctor explained all my options clearly and asked for my preference.\""],
        ["5",  "Dignity & Respect",                        "\"I felt humiliated when the doctor spoke about me as if I wasn't there.\""],
        ["6",  "Care Quality",                             "\"The standard of care was outstanding — I couldn't fault anything.\""],
        ["7",  "Care Coordination",                        "\"My care was well coordinated between my GP and the specialist.\""],
        ["8",  "Results & Information",                    "\"I waited three weeks for blood test results and had to call multiple times.\""],
        ["9",  "Facilities, Environment & Hygiene",        "\"The hospital parking was very limited and difficult to navigate.\""],
        ["10", "Telehealth",                               "\"My video consultation worked smoothly and saved me a two-hour drive.\""],
        ["11", "Emotional Support",                        "\"I was very anxious but the nurse reassured me and helped me feel calm.\""],
        ["12", "Cultural & Accessibility Needs",           "\"An interpreter was provided, which allowed me to communicate effectively.\""],
        ["13", "Discharge & Follow-up",                    "\"The discharge instructions were unclear — I didn't know what medications to take.\""],
        ["14", "Feeling in Good Hands",                    "\"I felt completely confident in my healthcare team throughout my treatment.\""],
    ],
    col_widths=[0.8, 5, 10.7]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. REQUEST & RESPONSE FORMAT
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "5. Request & Response Format")

h2(doc, "5.1  Request Examples")

h3(doc, "Single text — default model")
add_code(doc, """\
POST /classify
Content-Type: application/json
x-api-key: $HEALTHCARE_API_KEY

{
  "text": "The nurse was incredibly supportive and made me feel at ease.",
  "models": ["all"],
  "top_n": 3
}\
""", "http")

h3(doc, "Multiple texts at once")
add_code(doc, """\
{
  "texts": [
    "Dr Smith was very thorough and answered all my questions.",
    "I waited over two hours past my appointment time.",
    "The Zedoc portal kept crashing when I tried to book online.",
    "The car park was very limited and difficult to navigate."
  ],
  "models": ["all"],
  "top_n": 3
}\
""", "json")

h3(doc, "Specific models only (faster)")
add_code(doc, """\
{
  "texts": ["My telehealth video call worked perfectly and saved me a long drive."],
  "models": ["bertopic_mini", "lda"],
  "top_n": 2
}\
""", "json")

h2(doc, "5.2  Response Format")
body(doc, "A successful response (HTTP 200) always has this structure:")
add_code(doc, """\
{
  "results": [
    {
      "text": "Dr Smith was very thorough and answered all my questions.",
      "classifications": {
        "bertopic_mini": {
          "topic_id":   0,
          "topic_name": "Clinical Staff",
          "confidence": 0.85,
          "top_topics": [
            { "topic_id": 0, "topic_name": "Clinical Staff",        "confidence": 0.85 },
            { "topic_id": 6, "topic_name": "Care Quality",          "confidence": 0.09 },
            { "topic_id": 4, "topic_name": "Information & Shared Decision-Making", "confidence": 0.04 }
          ]
        },
        "lda": {
          "topic_id":   0,
          "topic_name": "Clinical Staff",
          "confidence": 0.76,
          "top_topics": [ ... ]
        }
      }
    }
  ],
  "model_status": {
    "bertopic_mini": true, "bertopic_mpnet": true,
    "lda": true, "lsi": true, "hdp": true, "nmf": true
  },
  "elapsed_sec": 1.23
}\
""", "json")

h3(doc, "Response Fields Explained")
add_table(doc,
    ["Field", "Type", "Description"],
    [
        ["results",                    "list",   "One entry per input text, in the same order you sent them."],
        ["results[i].text",            "string", "The original input text (echoed back for easy matching)."],
        ["results[i].classifications", "object", "One classification result per model you requested."],
        ["classifications.<model>",    "object", "The result from one model (see fields below)."],
        ["topic_id",                   "integer","The numeric ID of the best-matching topic (0–14)."],
        ["topic_name",                 "string", "The human-readable name of the best-matching topic."],
        ["confidence",                 "float",  "How confident the model is. Range 0.0 (no idea) to 1.0 (certain)."],
        ["top_topics",                 "list",   "The top N topics with their IDs, names, and confidence scores."],
        ["model_status",               "object", "Whether each model was loaded and ready at the time of the request."],
        ["elapsed_sec",                "float",  "How long the classification took in seconds."],
    ],
    col_widths=[4.5, 2.5, 9.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 6. CODE EXAMPLES
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "6. Code Examples")

h2(doc, "6.1  cURL (Command Line)")
body(doc, "The simplest way to test. Copy-paste these directly into your terminal after setting BASE and KEY.")

h3(doc, "Health check")
add_code(doc, """\
curl -s "$HEALTHCARE_API_BASE/health" \\
  -H "x-api-key: $HEALTHCARE_API_KEY" \\
  | python3 -m json.tool\
""", "bash")

h3(doc, "Classify a single text")
add_code(doc, """\
curl -s -X POST "$HEALTHCARE_API_BASE/classify" \\
  -H "Content-Type: application/json" \\
  -H "x-api-key: $HEALTHCARE_API_KEY" \\
  -d '{
    "text": "The discharge instructions were unclear and confusing.",
    "models": ["all"],
    "top_n": 3
  }' | python3 -m json.tool\
""", "bash")

h3(doc, "Classify multiple texts at once")
add_code(doc, """\
curl -s -X POST "$HEALTHCARE_API_BASE/classify" \\
  -H "Content-Type: application/json" \\
  -H "x-api-key: $HEALTHCARE_API_KEY" \\
  -d '{
    "texts": [
      "Dr Smith was very thorough and answered all my questions.",
      "I waited over two hours past my appointment time.",
      "The Zedoc portal kept crashing when I tried to book online."
    ],
    "models": ["all"],
    "top_n": 3
  }' | python3 -m json.tool\
""", "bash")

h2(doc, "6.2  Python")
body(doc, "Install the requests library first if you don't have it:")
add_code(doc, "pip install requests", "bash")

add_callout(doc,
    "All Python examples load the API key from the HEALTHCARE_API_KEY environment variable. "
    "Never hardcode the key as a string in your source code.",
    "warning"
)

h3(doc, "Basic example")
add_code(doc, """\
import os
import requests

# ── Load credentials from environment variables (never hardcode!) ───────────
BASE_URL = os.environ["HEALTHCARE_API_BASE"]   # set in Step 2 of Quick Start
API_KEY  = os.environ["HEALTHCARE_API_KEY"]    # set in Step 2 of Quick Start

HEADERS = {
    "Content-Type": "application/json",
    "x-api-key":    API_KEY,
}

# ── 1. Health check ─────────────────────────────────────────────────────────
response = requests.get(f"{BASE_URL}/health", headers=HEADERS)
print("Health:", response.json())

# ── 2. Classify a single text ───────────────────────────────────────────────
payload = {
    "text":   "The nurse was incredibly supportive and made me feel at ease.",
    "models": ["all"],
    "top_n":  3,
}

response = requests.post(f"{BASE_URL}/classify", headers=HEADERS, json=payload)
response.raise_for_status()   # raises an error if status is not 200

data = response.json()
for result in data["results"]:
    print(f"\\nText: {result['text']}")
    for model, classification in result["classifications"].items():
        print(f"  {model:20s}  ->  {classification['topic_name']}  "
              f"(confidence: {classification['confidence']:.2f})")\
""", "python")

h3(doc, "Classifying many texts in a loop")
add_code(doc, """\
import os
import requests

BASE_URL = os.environ["HEALTHCARE_API_BASE"]
API_KEY  = os.environ["HEALTHCARE_API_KEY"]
HEADERS  = {"Content-Type": "application/json", "x-api-key": API_KEY}

feedback_list = [
    "Dr Smith was very thorough and answered all my questions.",
    "I waited over two hours past my appointment time.",
    "The Zedoc portal kept crashing when I tried to book online.",
    "The car park was very limited and difficult to navigate.",
    "I was anxious but the nurse reassured me and helped me feel calm.",
]

# Send all texts in ONE request (more efficient than one request per text)
payload = {
    "texts":  feedback_list,
    "models": ["bertopic_mini"],   # use one model for speed
    "top_n":  1,                   # only need the top topic
}

response = requests.post(f"{BASE_URL}/classify", headers=HEADERS, json=payload)
response.raise_for_status()

for item in response.json()["results"]:
    best = item["classifications"]["bertopic_mini"]
    print(f"Topic: {best['topic_name']:45s}  | {item['text'][:60]}")\
""", "python")

h3(doc, "Saving results to a CSV file")
add_code(doc, """\
import os
import requests
import csv

BASE_URL = os.environ["HEALTHCARE_API_BASE"]
API_KEY  = os.environ["HEALTHCARE_API_KEY"]
HEADERS  = {"Content-Type": "application/json", "x-api-key": API_KEY}

feedback_list = [
    "Dr Smith was very thorough and answered all my questions.",
    "I waited over two hours past my appointment time.",
    "The Zedoc portal kept crashing when I tried to book online.",
]

response = requests.post(
    f"{BASE_URL}/classify",
    headers=HEADERS,
    json={"texts": feedback_list, "models": ["bertopic_mini"], "top_n": 3},
)
response.raise_for_status()

with open("classification_results.csv", "w", newline="") as f:
    writer = csv.writer(f)
    writer.writerow(["text", "topic_id", "topic_name", "confidence"])
    for item in response.json()["results"]:
        best = item["classifications"]["bertopic_mini"]
        writer.writerow([
            item["text"],
            best["topic_id"],
            best["topic_name"],
            round(best["confidence"], 4),
        ])

print("Results saved to classification_results.csv")\
""", "python")

h2(doc, "6.3  JavaScript / Node.js")
body(doc, "Using the built-in fetch API (Node.js 18+). Load the key from process.env — never hardcode it.")

h3(doc, "Using fetch (Node.js 18+)")
add_code(doc, """\
// Load from environment variables (set before running: export HEALTHCARE_API_KEY=...)
const BASE_URL = process.env.HEALTHCARE_API_BASE;
const API_KEY  = process.env.HEALTHCARE_API_KEY;

if (!BASE_URL || !API_KEY) {
  throw new Error("Set HEALTHCARE_API_BASE and HEALTHCARE_API_KEY env vars first.");
}

async function classifyFeedback(texts) {
  const response = await fetch(`${BASE_URL}/classify`, {
    method:  "POST",
    headers: {
      "Content-Type": "application/json",
      "x-api-key":    API_KEY,
    },
    body: JSON.stringify({ texts, models: ["all"], top_n: 3 }),
  });

  if (!response.ok) {
    throw new Error(`API error: ${response.status} ${response.statusText}`);
  }
  return response.json();
}

// Usage
classifyFeedback([
  "The doctor was very kind and explained everything clearly.",
  "I waited over two hours past my appointment time.",
])
  .then(data => {
    data.results.forEach(item => {
      const best = item.classifications.bertopic_mini;
      console.log(`Text:  ${item.text}`);
      console.log(`Topic: ${best.topic_name} (confidence: ${best.confidence.toFixed(2)})`);
      console.log("---");
    });
  })
  .catch(err => console.error("Error:", err));\
""", "javascript")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 7. TESTING GUIDE
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "7. Testing Guide")

body(doc, (
    "This section provides copy-paste test cases that cover each of the 15 PREM topics, "
    "edge cases, and error scenarios. Run these to verify the API is working correctly."
))

h2(doc, "7.1  Set Up (run once)")
add_code(doc, """\
BASE="$HEALTHCARE_API_BASE"
KEY="$HEALTHCARE_API_KEY"\
""", "bash")

h2(doc, "7.2  Test 1 — Health Check")
body(doc, "Verifies the API is live and all 6 models are loaded.")
add_code(doc, """\
curl -s \\
  "$HEALTHCARE_API_BASE/health" \\
  -H "x-api-key: $HEALTHCARE_API_KEY" \\
  | python3 -m json.tool\
""", "bash")
body(doc, "✅  Pass if: status is \"ok\" and all 6 models show true.")

h2(doc, "7.3  Test 2 — Single Text, All Models")
body(doc, "Verifies the classify endpoint works with a single text across all 6 models.")
add_code(doc, """\
curl -s -X POST \\
  "$HEALTHCARE_API_BASE/classify" \\
  -H "Content-Type: application/json" \\
  -H "x-api-key: $HEALTHCARE_API_KEY" \\
  -d '{
    "text": "Dr Smith was very thorough and answered all my questions.",
    "models": ["all"],
    "top_n": 3
  }' | python3 -m json.tool\
""", "bash")
body(doc, "✅  Pass if: response contains classifications for all 6 models and topic_name is \"Clinical Staff\".")

h2(doc, "7.4  Test 3 — All 15 PREM Topics")
body(doc, "Sends one representative text per topic to verify all categories can be detected.")
add_code(doc, """\
curl -s -X POST \\
  "$HEALTHCARE_API_BASE/classify" \\
  -H "Content-Type: application/json" \\
  -H "x-api-key: $HEALTHCARE_API_KEY" \\
  -d '{
    "texts": [
      "My GP was very thorough and spent time answering all my questions.",
      "The receptionist was helpful and checked me in quickly.",
      "I waited over two hours past my appointment time with no explanation.",
      "The Zedoc booking portal kept crashing when I tried to schedule.",
      "The doctor explained all my treatment options and asked for my preference.",
      "I felt humiliated when the doctor spoke about me as if I was not there.",
      "The standard of care was outstanding and I could not fault anything.",
      "My care was well coordinated between my GP and the specialist.",
      "I waited three weeks for blood test results and had to call multiple times.",
      "The hospital parking was very limited and difficult to navigate.",
      "My telehealth video consultation worked smoothly and saved me a long drive.",
      "I was very anxious but the nurse reassured me and helped me feel calm.",
      "An interpreter was provided which allowed me to communicate effectively.",
      "I was discharged without clear instructions about my medications.",
      "I felt completely confident in my healthcare team throughout my treatment."
    ],
    "models": ["bertopic_mini"],
    "top_n": 1
  }' | python3 -m json.tool\
""", "bash")
body(doc, "✅  Pass if: 15 results are returned, each with a topic_id from 0 to 14.")

h2(doc, "7.5  Test 4 — Specific Models Only")
body(doc, "Verifies you can select a subset of models instead of running all 6.")
add_code(doc, """\
curl -s -X POST \\
  "$HEALTHCARE_API_BASE/classify" \\
  -H "Content-Type: application/json" \\
  -H "x-api-key: $HEALTHCARE_API_KEY" \\
  -d '{
    "texts": ["My telehealth video call worked perfectly."],
    "models": ["bertopic_mini", "lda"],
    "top_n": 2
  }' | python3 -m json.tool\
""", "bash")
body(doc, "✅  Pass if: only bertopic_mini and lda appear under classifications.")

h2(doc, "7.6  Test 5 — Error Cases")
body(doc, "These tests verify the API returns the correct error codes for bad requests.")

h3(doc, "Missing API key → 403 Forbidden")
add_code(doc, """\
# Deliberately omit the key — expect 403
curl -s -o /dev/null -w "%{http_code}" \\
  -X POST \\
  "$HEALTHCARE_API_BASE/classify" \\
  -H "Content-Type: application/json" \\
  -d '{"text": "test"}'\
""", "bash")
body(doc, "✅  Pass if: output is 403.")

h3(doc, "Wrong API key → 403 Forbidden")
add_code(doc, """\
curl -s -o /dev/null -w "%{http_code}" \\
  -X POST \\
  "$HEALTHCARE_API_BASE/classify" \\
  -H "Content-Type: application/json" \\
  -H "x-api-key: wrong-key-here" \\
  -d '{"text": "test"}'\
""", "bash")
body(doc, "✅  Pass if: output is 403.")

h3(doc, "Missing texts field → 422 Unprocessable")
add_code(doc, """\
curl -s -X POST \\
  "$HEALTHCARE_API_BASE/classify" \\
  -H "Content-Type: application/json" \\
  -H "x-api-key: $HEALTHCARE_API_KEY" \\
  -d '{"models": ["all"]}' \\
  | python3 -m json.tool\
""", "bash")
body(doc, "✅  Pass if: response contains an error message about missing \"texts\" field.")

h3(doc, "Invalid model name → 422 Unprocessable")
add_code(doc, """\
curl -s -X POST \\
  "$HEALTHCARE_API_BASE/classify" \\
  -H "Content-Type: application/json" \\
  -H "x-api-key: $HEALTHCARE_API_KEY" \\
  -d '{"text": "test", "models": ["invalid_model_name"]}' \\
  | python3 -m json.tool\
""", "bash")
body(doc, "✅  Pass if: response contains an error message listing valid model names.")

h2(doc, "7.7  Direct Lambda Invoke (AWS CLI)")
body(doc, (
    "You can also invoke the Lambda function directly via the AWS CLI, "
    "bypassing API Gateway entirely. This is useful for debugging."
))
add_code(doc, """\
aws lambda invoke \\
  --function-name healthcare-topic-classifier \\
  --region ap-southeast-2 \\
  --payload '{
    "texts": ["The nursing staff were attentive and caring throughout my stay."],
    "models": ["all"],
    "top_n": 3
  }' \\
  --cli-binary-format raw-in-base64-out \\
  response.json && cat response.json | python3 -m json.tool\
""", "bash")
add_callout(doc,
    "Direct Lambda invoke does NOT require the x-api-key header — it uses your AWS IAM credentials instead.",
    "note"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 8. ERROR REFERENCE
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "8. Error Reference")

body(doc, "The API uses standard HTTP status codes. Here is what each one means and how to fix it:")

add_table(doc,
    ["HTTP Status", "Name", "Cause", "How to Fix"],
    [
        ["200", "OK",                   "Request succeeded.",                                   "No action needed."],
        ["403", "Forbidden",            "x-api-key header is missing or has the wrong value.",  "Add -H \"x-api-key: <your-key>\" to every request."],
        ["400", "Bad Request",          "The request body is not valid JSON.",                  "Check for missing quotes, commas, or brackets in your JSON."],
        ["422", "Unprocessable Entity", "A required field is missing or has an invalid value.", "Check that \"text\" or \"texts\" is present. Check model names are spelled correctly."],
        ["429", "Too Many Requests",    "Rate limit exceeded: 50 requests/sec or 10,000/day.",  "Slow down your requests. Add a small delay between calls if batching."],
        ["500", "Internal Server Error","An unexpected error occurred in the Lambda function.", "Check Lambda logs. Retry the request. Contact the team if it persists."],
        ["504", "Gateway Timeout",      "Request took longer than 29 seconds (API Gateway limit).", "Reduce the number of texts per request. Use models: [\"bertopic_mini\"] for speed."],
    ],
    col_widths=[2.2, 3.5, 5.8, 5]
)

h2(doc, "8.1  Viewing Lambda Logs")
body(doc, "If you get a 500 error, check the Lambda logs with this command:")
add_code(doc, """\
aws logs tail /aws/lambda/healthcare-topic-classifier \\
  --region ap-southeast-2 \\
  --follow\
""", "bash")

add_callout(doc,
    "The --follow flag streams logs in real time, like \"tail -f\" on a local file. Press Ctrl+C to stop.",
    "tip"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 9. BEST PRACTICES
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "9. Best Practices")

h2(doc, "9.1  Performance")
for item in [
    "Batch your texts — send up to 500 texts in a single request rather than one request per text. This is much faster.",
    "Use models: [\"bertopic_mini\"] when you need speed. Use [\"all\"] when you need to compare results.",
    "The first request after a long idle period may take 30–60 seconds (cold start). Subsequent requests are fast. "
    "An EventBridge rule pings the API every 5 minutes to keep it warm during business hours.",
    "Keep texts under 300 words. Longer texts are truncated by the models anyway.",
]:
    bullet(doc, item)

h2(doc, "9.2  Security")
for item in [
    "Never hardcode the API key directly in your source code. Use environment variables instead.",
    "Do not commit the API key to Git. Add it to your .gitignore or use a .env file.",
    "Rotate the key periodically. Contact the platform team to issue a new key.",
    "In production, call the API from your server — never from a browser or mobile app directly (the key would be visible).",
]:
    bullet(doc, item)

body(doc, "Store the key safely in Python like this:")
add_code(doc, """\
import os

# Retrieve the key at runtime — never hardcode it:
#   export HEALTHCARE_API_KEY=$(aws apigateway get-api-keys \\
#     --region ap-southeast-2 --include-values \\
#     --query "items[?contains(name,'healthcare-topic-classifier')].value | [0]" \\
#     --output text)

API_KEY = os.environ.get("HEALTHCARE_API_KEY")
if not API_KEY:
    raise ValueError("HEALTHCARE_API_KEY environment variable is not set!")\
""", "python")

h2(doc, "9.3  Error Handling")
body(doc, "Always handle errors in your code. Here is a robust Python example:")
add_code(doc, """\
import os, requests

def classify_feedback(texts, models=None, top_n=3):
    \"\"\"Classify feedback texts. Returns results or raises a clear error.\"\"\"
    base_url = os.environ["HEALTHCARE_API_BASE"]   # e.g. https://<id>.execute-api.<region>.amazonaws.com/prod
    api_key  = os.environ["HEALTHCARE_API_KEY"]    # retrieved via AWS CLI, never hardcoded
    url      = f"{base_url}/classify"
    headers  = {
        "Content-Type": "application/json",
        "x-api-key":    api_key,
    }

    try:
        response = requests.post(
            url,
            headers=headers,
            json={"texts": texts, "models": models or ["all"], "top_n": top_n},
            timeout=60,    # wait up to 60 seconds (handles cold starts)
        )
        response.raise_for_status()
        return response.json()

    except requests.exceptions.Timeout:
        print("Request timed out. The API may be cold-starting. Try again.")
        raise
    except requests.exceptions.HTTPError as e:
        if e.response.status_code == 403:
            print("Authentication failed. Check your API key.")
        elif e.response.status_code == 422:
            print("Invalid request:", e.response.json())
        elif e.response.status_code == 429:
            print("Rate limit exceeded. Slow down your requests.")
        elif e.response.status_code >= 500:
            print("Server error. Check Lambda logs.")
        raise\
""", "python")

h2(doc, "9.4  Rate Limits")
add_table(doc,
    ["Limit", "Value", "What happens when exceeded"],
    [
        ["Requests per second",  "50 rps",         "HTTP 429 — slow down and retry after 1 second."],
        ["Burst limit",          "100 requests",   "HTTP 429 — brief spike limit."],
        ["Requests per day",     "10,000",         "HTTP 429 — quota resets at midnight UTC."],
        ["Max texts per request","500",            "HTTP 422 — split into smaller batches."],
        ["Max text length",      "5,000 characters","Longer texts are silently truncated by the model."],
    ],
    col_widths=[4, 3, 9.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 10. INFRASTRUCTURE OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "10. Infrastructure Overview")

body(doc, (
    "This section gives you a high-level picture of how the API is deployed on AWS. "
    "You don't need to manage any of this — it is here so you understand what happens "
    "when you make a request."
))

add_table(doc,
    ["AWS Service", "Resource Name", "What It Does"],
    [
        ["API Gateway (REST)", "healthcare-topic-classifier-api", "Receives your HTTP requests and forwards them to the Lambda function."],
        ["API Key",            "healthcare-topic-classifier-api-key-prod", "Your x-api-key value. Enforces rate limits and quotas."],
        ["Usage Plan",         "healthcare-topic-classifier-usage-plan-prod", "50 rps / 10,000 req/day quota settings."],
        ["AWS Lambda",         "healthcare-topic-classifier", "The Python function that runs all 6 NLP models and returns classifications."],
        ["Amazon ECR",         "healthcare-topic-classifier", "Stores the Docker container image (includes trained models baked in)."],
        ["EventBridge Rule",   "healthcare-topic-classifier-warmup", "Pings the Lambda every 5 minutes to prevent cold starts."],
        ["IAM Role",           "healthcare-topic-classifier-lambda-role", "Grants the Lambda permission to pull from ECR and write logs."],
        ["CloudFormation",     "healthcare-topic-classifier-stack", "Manages all the above resources as a single deployable unit."],
    ],
    col_widths=[3.5, 5.5, 7.5]
)

body(doc, "The request flow when you call the API:")
for step in [
    "Your code sends an HTTPS request to the API Gateway URL.",
    "API Gateway checks the x-api-key header. If invalid → 403.",
    "API Gateway forwards the request to the Lambda function.",
    "Lambda loads the requested NLP models (already in memory if warm).",
    "Lambda classifies the texts and builds the JSON response.",
    "Lambda returns the response to API Gateway.",
    "API Gateway sends the JSON response back to your code.",
]:
    numbered(doc, step)

add_callout(doc,
    "Cold start: if the Lambda has been idle for more than ~15 minutes, the first request triggers "
    "a cold start where all models are loaded from disk. This takes 30–60 seconds. "
    "The EventBridge warmup rule prevents this during business hours by pinging the API every 5 minutes.",
    "note"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 11. TROUBLESHOOTING
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "11. Troubleshooting")

add_table(doc,
    ["Problem", "Likely Cause", "Solution"],
    [
        ["403 Forbidden on every request",
         "x-api-key header is missing or wrong.",
         "Double-check the header name (lowercase: x-api-key) and the key value. No extra spaces."],
        ["curl: command not found",
         "curl is not installed.",
         "On Ubuntu/WSL: sudo apt install curl. On Mac: brew install curl."],
        ["First request takes 60 seconds",
         "Lambda cold start — models loading from disk.",
         "Normal behaviour. Run GET /health first to pre-warm. Subsequent calls are fast."],
        ["504 Gateway Timeout",
         "Request took longer than 29 seconds.",
         "Reduce batch size (fewer texts per request) or use a faster model like bertopic_mini."],
        ["json.decoder.JSONDecodeError in Python",
         "The response body is not JSON (e.g. an HTML error page).",
         "Print response.text before calling response.json() to see the raw response."],
        ["top_topics list is empty",
         "top_n was set to 0.",
         "Use top_n: 1 or higher. top_n: 3 is the recommended default."],
        ["model_status shows false for some models",
         "A model failed to load on cold start.",
         "Call GET /warmup to reload all models. Check Lambda logs if it persists."],
        ["Confidence scores all equal for every topic",
         "Input text is too short or contains only stop words (e.g. \"ok\").",
         "Ensure texts are at least 4 meaningful words long."],
        ["429 Too Many Requests",
         "Sending more than 50 requests per second.",
         "Add time.sleep(0.1) between requests or send larger batches instead of many small ones."],
    ],
    col_widths=[4, 4.5, 8]
)

h2(doc, "11.1  Useful Debugging Commands")
h3(doc, "Check the Lambda is running")
add_code(doc, """\
aws lambda get-function \\
  --function-name healthcare-topic-classifier \\
  --region ap-southeast-2 \\
  --query "Configuration.{State:State,LastStatus:LastUpdateStatus}" \\
  --output table\
""", "bash")

h3(doc, "View recent Lambda logs")
add_code(doc, """\
aws logs tail /aws/lambda/healthcare-topic-classifier \\
  --region ap-southeast-2 \\
  --since 1h\
""", "bash")

h3(doc, "Check the CloudFormation stack status")
add_code(doc, """\
aws cloudformation describe-stacks \\
  --stack-name healthcare-topic-classifier-stack \\
  --region ap-southeast-2 \\
  --query "Stacks[0].{Status:StackStatus,Updated:LastUpdatedTime}" \\
  --output table\
""", "bash")

h3(doc, "Retrieve the API key again if you lose it")
add_code(doc, """\
aws apigateway get-api-keys \\
  --region ap-southeast-2 \\
  --include-values \\
  --query "items[?contains(name,'healthcare-topic-classifier')].{name:name,value:value}" \\
  --output table\
""", "bash")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 12. QUICK REFERENCE CARD
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "12. Quick Reference Card")

body(doc, "Tear out and keep this page handy.")

h2(doc, "Credentials")
add_code(doc, """\
BASE="$HEALTHCARE_API_BASE"
KEY="$HEALTHCARE_API_KEY"\
""", "bash")

h2(doc, "Endpoints")
add_table(doc,
    ["Method", "Path", "Use For"],
    [
        ["GET",  "/health",   "Check API is alive and all 6 models are loaded."],
        ["POST", "/classify", "Classify one or more patient feedback texts into PREM topics."],
    ],
    col_widths=[2, 3, 11.5]
)

h2(doc, "Models")
add_table(doc,
    ["Key", "Description"],
    [
        ["bertopic_mini",  "BERTopic + MiniLM  — fastest, good general accuracy"],
        ["bertopic_mpnet", "BERTopic + MPNet   — slower, highest accuracy"],
        ["lda",            "Latent Dirichlet Allocation — traditional NLP"],
        ["lsi",            "Latent Semantic Indexing — deterministic, keyword-based"],
        ["hdp",            "Hierarchical Dirichlet Process — auto topic count"],
        ["nmf",            "Non-negative Matrix Factorization — good for short texts"],
    ],
    col_widths=[3.5, 13]
)

h2(doc, "Minimal cURL — classify one text")
add_code(doc, """\
curl -s -X POST \\
  "$HEALTHCARE_API_BASE/classify" \\
  -H "Content-Type: application/json" \\
  -H "x-api-key: $HEALTHCARE_API_KEY" \\
  -d '{"text": "YOUR TEXT HERE", "models": ["all"], "top_n": 3}' \\
  | python3 -m json.tool\
""", "bash")

h2(doc, "Minimal Python")
add_code(doc, """\
import os, requests
r = requests.post(
    os.environ["HEALTHCARE_API_BASE"] + "/classify",
    headers={"Content-Type": "application/json",
             "x-api-key": os.environ["HEALTHCARE_API_KEY"]},
    json={"text": "YOUR TEXT HERE", "models": ["all"], "top_n": 3},
)
print(r.json())\
""", "python")

h2(doc, "HTTP Status Codes")
add_table(doc,
    ["Code", "Meaning", "Fix"],
    [
        ["200", "Success",            "No action needed."],
        ["403", "Bad / missing key",  "Add -H \"x-api-key: <key>\" to every request."],
        ["422", "Bad request data",   "Check \"text\" or \"texts\" field exists. Check model names."],
        ["429", "Rate limit",         "Slow down. Max 50 rps, 10,000/day."],
        ["500", "Server error",       "Retry. Check Lambda logs if it persists."],
        ["504", "Timeout",            "Reduce batch size or use bertopic_mini."],
    ],
    col_widths=[1.8, 3.5, 11.2]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 13. MODEL EVALUATION
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "13. Topic Model Classifier Evaluation")

body(doc, (
    "This section explains how the six NLP models are evaluated, what the evaluation "
    "metrics mean, how to run evaluations yourself via the API, and how to interpret "
    "the results to choose the best model for your use case."
))

add_callout(doc,
    "Evaluation tells you how accurately each model classifies patient feedback into "
    "the correct PREM topic. A higher accuracy means fewer misclassifications in "
    "your application.",
    "note"
)

# ── 13.1 ──────────────────────────────────────────────────────────────────────
h2(doc, "13.1  Evaluation Metrics Explained")

body(doc, "The following metrics are used to assess each model:")

add_table(doc,
    ["Metric", "What It Measures", "How to Read It"],
    [
        ["Top-1 Accuracy",
         "The % of test texts where the model's #1 predicted topic matches the correct topic.",
         "80% means 8 out of 10 texts were correctly classified. Higher is better."],
        ["Top-3 Accuracy",
         "The % of test texts where the correct topic appears anywhere in the top 3 predictions.",
         "Always ≥ Top-1 Accuracy. Useful when you present multiple suggestions to a user."],
        ["Confidence Score",
         "How sure the model is about its top prediction (0.0 = no idea, 1.0 = certain).",
         "A high accuracy model with low confidence may be unreliable on unseen data."],
        ["Average Confidence (Correct)",
         "Mean confidence score across all correctly classified texts.",
         "High value (e.g. 0.80+) means the model is confident when it is right."],
        ["Average Confidence (Wrong)",
         "Mean confidence score across misclassified texts.",
         "If this is high (e.g. 0.75+), the model is confidently wrong — a red flag."],
        ["Per-Topic Accuracy",
         "Accuracy broken down by each of the 15 PREM topics.",
         "Reveals which topics each model handles well and which it confuses."],
    ],
    col_widths=[4, 6, 6.5]
)

# ── 13.2 ──────────────────────────────────────────────────────────────────────
h2(doc, "13.2  Standard 15-Case Evaluation Dataset")

body(doc, (
    "The standard evaluation uses 15 representative test cases — one per PREM topic. "
    "Each test case has a known correct topic (the \"ground truth\") so you can measure "
    "whether each model predicted correctly."
))

add_table(doc,
    ["ID", "Expected Topic", "Test Text"],
    [
        ["0",  "Clinical Staff",                       "Dr Smith was incredibly thorough and explained my entire care plan step by step."],
        ["1",  "Administrative Staff",                 "The receptionist at the front desk resolved my billing issue very efficiently."],
        ["2",  "Efficiency and Flow",                  "I waited two hours past my appointment time with no explanation given."],
        ["3",  "Booking & Access",                     "The Zedoc patient portal crashed every time I tried to book an appointment."],
        ["4",  "Information & Shared Decision-Making", "The specialist explained all treatment options and asked for my preference."],
        ["5",  "Dignity & Respect",                    "I felt humiliated when the doctor discussed my case without acknowledging me."],
        ["6",  "Care Quality",                         "The standard of care was outstanding — thorough, careful, and highly professional."],
        ["7",  "Care Coordination",                    "My referral was seamlessly coordinated between my GP and the specialist team."],
        ["8",  "Results & Information",                "I waited three weeks for my blood test results and couldn't understand the report."],
        ["9",  "Facilities, Environment & Hygiene",    "The car park was full and I had to walk from Westfield with difficulty."],
        ["10", "Telehealth",                           "My telehealth video consultation was smooth and saved a long drive to the clinic."],
        ["11", "Emotional Support",                    "I was very anxious before surgery but the nurse reassured me and I felt safe."],
        ["12", "Cultural & Accessibility Needs",       "An interpreter was arranged which allowed me to communicate with my doctor properly."],
        ["13", "Discharge & Follow-up",                "I was discharged without clear instructions about my medications or follow-up date."],
        ["14", "Feeling in Good Hands",                "I cannot praise this team enough — I left every appointment feeling in great hands."],
    ],
    col_widths=[0.8, 4.8, 10.9]
)

# ── 13.3 ──────────────────────────────────────────────────────────────────────
h2(doc, "13.3  Running an Evaluation via the API")

body(doc, (
    "You can run the full 15-case evaluation against the live AWS API using the "
    "Python script below. It sends all 15 test texts in a single request and "
    "computes per-model accuracy automatically."
))

h3(doc, "Step 1 — Install dependencies")
add_code(doc, "pip install requests", "bash")

h3(doc, "Step 2 — Run the evaluation script")
add_code(doc, """\
import os, requests

# ── Configuration — credentials from environment (never hardcoded) ─────────────
# Set before running:
#   export HEALTHCARE_API_BASE=$(aws cloudformation describe-stacks \\
#     --stack-name healthcare-topic-classifier-stack --region ap-southeast-2 \\
#     --query "Stacks[0].Outputs[?OutputKey=='ApiEndpoint'].OutputValue" --output text)
#   export HEALTHCARE_API_KEY=$(aws apigateway get-api-keys \\
#     --region ap-southeast-2 --include-values \\
#     --query "items[?contains(name,'healthcare-topic-classifier')].value | [0]" \\
#     --output text)
BASE_URL = os.environ["HEALTHCARE_API_BASE"]
API_KEY  = os.environ["HEALTHCARE_API_KEY"]
HEADERS  = {"Content-Type": "application/json", "x-api-key": API_KEY}

# ── Ground-truth test cases (text, expected topic name) ────────────────────────
TEST_CASES = [
    (0,  "Clinical Staff",                       "Dr Smith was incredibly thorough and explained my entire care plan step by step."),
    (1,  "Administrative Staff",                 "The receptionist resolved my billing issue very efficiently."),
    (2,  "Efficiency and Flow",                  "I waited two hours past my appointment time with no explanation given."),
    (3,  "Booking & Access",                     "The Zedoc patient portal crashed every time I tried to book an appointment."),
    (4,  "Information & Shared Decision-Making", "The specialist explained all treatment options and asked for my preference."),
    (5,  "Dignity & Respect",                    "I felt humiliated when the doctor discussed my case without acknowledging me."),
    (6,  "Care Quality",                         "The standard of care was outstanding — thorough, careful, highly professional."),
    (7,  "Care Coordination",                    "My referral was seamlessly coordinated between my GP and the specialist team."),
    (8,  "Results & Information",                "I waited three weeks for blood test results and couldn't understand the report."),
    (9,  "Facilities, Environment & Hygiene",    "The car park was full and I had to walk from Westfield with difficulty."),
    (10, "Telehealth",                           "My telehealth video consultation was smooth and saved a long drive."),
    (11, "Emotional Support",                    "I was anxious before surgery but the nurse reassured me and I felt safe."),
    (12, "Cultural & Accessibility Needs",       "An interpreter was arranged so I could communicate with my doctor properly."),
    (13, "Discharge & Follow-up",                "I was discharged without clear instructions about my medications."),
    (14, "Feeling in Good Hands",                "I cannot praise this team enough — I left every visit feeling in great hands."),
]

texts         = [text for _, _, text in TEST_CASES]
expected_ids  = [tid for tid, _, _ in TEST_CASES]
expected_names = [name for _, name, _ in TEST_CASES]

# ── Call the API ────────────────────────────────────────────────────────────────
print("Calling API...")
response = requests.post(
    f"{BASE_URL}/classify",
    headers=HEADERS,
    json={"texts": texts, "models": ["all"], "top_n": 3},
    timeout=90,
)
response.raise_for_status()
results = response.json()["results"]

# ── Compute per-model metrics ───────────────────────────────────────────────────
MODELS = ["bertopic_mini", "bertopic_mpnet", "lda", "lsi", "hdp", "nmf"]

metrics = {m: {"top1": 0, "top3": 0, "conf_correct": [], "conf_wrong": []}
           for m in MODELS}

for i, (item, exp_name) in enumerate(zip(results, expected_names)):
    for model in MODELS:
        clf   = item["classifications"][model]
        pred  = clf["topic_name"]
        conf  = clf["confidence"]
        top3  = [t["topic_name"] for t in clf["top_topics"]]

        if pred == exp_name:
            metrics[model]["top1"] += 1
            metrics[model]["conf_correct"].append(conf)
        else:
            metrics[model]["conf_wrong"].append(conf)

        if exp_name in top3:
            metrics[model]["top3"] += 1

# ── Print results ───────────────────────────────────────────────────────────────
n = len(TEST_CASES)
print(f"\\n{'Model':<20} {'Top-1 Acc':>10} {'Top-3 Acc':>10} "
      f"{'Avg Conf ✅':>12} {'Avg Conf ❌':>12}")
print("-" * 70)
for model in MODELS:
    m      = metrics[model]
    top1   = m["top1"] / n * 100
    top3   = m["top3"] / n * 100
    avg_c  = sum(m["conf_correct"]) / len(m["conf_correct"]) if m["conf_correct"] else 0
    avg_w  = sum(m["conf_wrong"])   / len(m["conf_wrong"])   if m["conf_wrong"]   else 0
    bar    = "█" * m["top1"] + "░" * (n - m["top1"])
    print(f"{model:<20} {top1:>9.0f}%  {top3:>9.0f}%  "
          f"{avg_c:>11.2f}   {avg_w:>11.2f}   {bar}")
print("-" * 70)\
""", "python")

add_callout(doc,
    "Save this script as evaluate_models.py and run it with: python3 evaluate_models.py",
    "tip"
)

# ── 13.4 ──────────────────────────────────────────────────────────────────────
h2(doc, "13.4  Interpreting Evaluation Results")

body(doc, (
    "Once you have run the evaluation, use this guide to interpret what the numbers mean "
    "and decide which model to use."
))

h3(doc, "Reading the accuracy table")
add_table(doc,
    ["Top-1 Accuracy", "What It Means", "Recommendation"],
    [
        ["90–100%", "Excellent — the model almost always picks the right topic.",
         "Use this model in production with confidence."],
        ["75–89%",  "Good — correct most of the time, occasional misclassifications.",
         "Suitable for production. Show top-3 topics to users as a safety net."],
        ["60–74%",  "Fair — wrong roughly 1 in 4 times.",
         "Use with caution. Consider combining with another model (ensemble)."],
        ["Below 60%","Poor — struggles with this dataset.",
         "Do not rely on this model alone. Use it alongside a higher-accuracy model."],
    ],
    col_widths=[3.5, 7, 6]
)

h3(doc, "Understanding confidence scores")
body(doc, (
    "The confidence score (0.0–1.0) shows how certain the model is. "
    "It is important to look at confidence alongside accuracy:"
))
for item in [
    "High accuracy + high confidence (e.g. 0.80+): The model is reliable — use its top-1 prediction.",
    "High accuracy + low confidence (e.g. below 0.50): The model gets it right but is uncertain — always show top-3 options.",
    "Low accuracy + high confidence: The model is confidently wrong — do not trust it on this data.",
    "Low accuracy + low confidence: The model is uncertain and often wrong — consider a different model.",
]:
    bullet(doc, item)

add_callout(doc,
    "A good rule of thumb: if the top-1 confidence is below 0.60, show the top-3 topics "
    "to the user instead of just one. This gives them a chance to pick the correct one.",
    "tip"
)

h3(doc, "Comparing models side by side")
body(doc, (
    "When running models: [\"all\"], you get all 6 predictions for every text. "
    "You can use majority voting to pick the most reliable answer:"
))
add_code(doc, """\
from collections import Counter

# After calling the API with models: ["all"] ...
for item in results:
    # Collect all model predictions for this text
    votes = [clf["topic_name"]
             for clf in item["classifications"].values()]

    # Pick the topic that the most models agreed on
    winner, count = Counter(votes).most_common(1)[0]
    print(f"Majority vote: {winner}  ({count}/6 models agreed)")
    print(f"Text: {item['text'][:60]}")
    print()\
""", "python")

# ── 13.5 ──────────────────────────────────────────────────────────────────────
h2(doc, "13.5  Per-Topic Accuracy Analysis")

body(doc, (
    "Some topics are easier to classify than others. "
    "Run this script to see a breakdown of which topics each model handles well "
    "and which ones it tends to confuse."
))

add_code(doc, """\
# Continuing from the evaluation script in Section 13.3 ...
# (results, TEST_CASES, expected_names, MODELS must already be defined)

from collections import defaultdict

# Per-topic accuracy for each model
print("\\nPer-Topic Accuracy (bertopic_mini):")
print(f"  {'Topic':<44} {'Result':<8} Predicted")
print("  " + "-" * 75)

for i, (item, (tid, exp_name, text)) in enumerate(zip(results, TEST_CASES)):
    clf  = item["classifications"]["bertopic_mini"]
    pred = clf["topic_name"]
    conf = clf["confidence"]
    hit  = "✅ PASS" if pred == exp_name else "❌ FAIL"
    print(f"  {exp_name:<44} {hit}   {pred}  ({conf:.0%})")\
""", "python")

# ── 13.6 ──────────────────────────────────────────────────────────────────────
h2(doc, "13.6  Model Selection Guide")

body(doc, (
    "Use this table to choose the right model for your use case. "
    "You do not need to run all 6 models every time — "
    "pick the one that best fits your requirements."
))

add_table(doc,
    ["Use Case", "Recommended Model", "Why"],
    [
        ["Quick prototype or demo",
         "bertopic_mini",
         "Fastest response time. Good accuracy on clear, well-written feedback."],
        ["Production — maximum accuracy",
         "bertopic_mpnet",
         "Highest quality embeddings. Best for ambiguous or nuanced text."],
        ["Short feedback (1–2 sentences)",
         "nmf",
         "Non-negative Matrix Factorization handles short texts well."],
        ["Need deterministic/repeatable results",
         "lsi",
         "LSI always produces the same output for the same input — no randomness."],
        ["Unknown number of topics in your data",
         "hdp",
         "HDP automatically determines the number of topics — no fixed N required."],
        ["Traditional NLP pipeline",
         "lda",
         "Interpretable topic-word distributions. Familiar to data scientists."],
        ["Highest confidence across the board",
         "bertopic_mini + bertopic_mpnet (majority vote)",
         "Two BERTopic models agreeing gives high reliability on most inputs."],
        ["Full comparison / research",
         "all",
         "Run all 6 and use majority voting or pick the consensus prediction."],
    ],
    col_widths=[4.5, 4, 8]
)

# ── 13.7 ──────────────────────────────────────────────────────────────────────
h2(doc, "13.7  Known Limitations")

body(doc, "Be aware of these known limitations when evaluating or using the models:")

add_table(doc,
    ["Limitation", "Details", "Workaround"],
    [
        ["Short or vague texts",
         "Texts under 4 words or consisting only of generic phrases (e.g. \"it was ok\") "
         "produce low-confidence, unreliable results.",
         "Filter out texts with fewer than 4 meaningful words before sending to the API."],
        ["Overlapping topics",
         "Some topics are semantically close (e.g. Clinical Staff vs Care Quality). "
         "Models may swap these.",
         "Use top_n: 3 and show all top topics. Consider combining adjacent topic IDs in your reporting."],
        ["Seed corpus bias",
         "Classical models (LDA, LSI, HDP, NMF) were trained on 225 seed sentences. "
         "They may perform worse on text that is very different in style or vocabulary.",
         "Evaluate on your actual patient feedback before going to production."],
        ["BERTopic cold-start",
         "The first request after a long idle period loads model weights from disk (~30–60 s).",
         "Call GET /health before your first classify request to pre-warm the Lambda."],
        ["Language",
         "All models are trained on English text only. Non-English feedback will be misclassified.",
         "Translate non-English texts to English before calling the API."],
        ["Maximum text length",
         "Texts longer than 5,000 characters are silently truncated before classification.",
         "Keep feedback under 300 words for best results."],
    ],
    col_widths=[3.5, 6.5, 6.5]
)

add_callout(doc,
    "Always evaluate the models on a sample of your own real patient feedback data "
    "before deploying to production. The 15 seed test cases are a starting point — "
    "your actual data may have different characteristics.",
    "warning"
)

# ── Save to docs/ ─────────────────────────────────────────────────────────────
import os as _os
_os.makedirs("docs", exist_ok=True)
output_path = "docs/NLP_Healthcare_API_Integration_Guide.docx"
doc.save(output_path)
print(f"✅  Saved: {output_path}")
