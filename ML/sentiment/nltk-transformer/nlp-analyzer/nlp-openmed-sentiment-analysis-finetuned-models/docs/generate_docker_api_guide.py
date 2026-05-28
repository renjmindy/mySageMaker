"""
Generate: NLP OpenMed Sentiment Analysis — Docker & API Guide
Target audience: Junior software engineer (new graduate)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY       = RGBColor(0x1A, 0x3A, 0x5C)
TEAL       = RGBColor(0x00, 0x7A, 0x87)
DARK_GRAY  = RGBColor(0x33, 0x33, 0x33)
MID_GRAY   = RGBColor(0x66, 0x66, 0x66)
GREEN      = RGBColor(0x1E, 0x7E, 0x34)
ORANGE     = RGBColor(0xD3, 0x6A, 0x00)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
CODE_BG    = RGBColor(0xF4, 0xF6, 0xF8)
HEADER_BG  = RGBColor(0x1A, 0x3A, 0x5C)

# ── Helper: shade a table cell ────────────────────────────────────────────────
def shade_cell(cell, rgb: RGBColor):
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tc_pr.append(shd)

def shade_row(row, rgb: RGBColor):
    for cell in row.cells:
        shade_cell(cell, rgb)

# ── Helper: callout box ───────────────────────────────────────────────────────
def add_callout(doc, text, style="tip"):
    icons  = {"tip": "💡 Tip", "warning": "⚠️  Important", "note": "📝 Note"}
    colors = {"tip": GREEN, "warning": ORANGE, "note": TEAL}
    label  = icons.get(style, "📝 Note")
    color  = colors.get(style, TEAL)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(f"{label}:  ")
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.color.rgb = DARK_GRAY
    run2.font.size = Pt(10)
    return p

# ── Helper: code block (shaded table) ────────────────────────────────────────
def add_code(doc, code_text, lang=""):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, CODE_BG)
    cell.paragraphs[0].clear()
    lines = code_text.strip().split("\n")
    for i, line in enumerate(lines):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        run = para.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    doc.add_paragraph()

# ── Helper: headings ─────────────────────────────────────────────────────────
def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(18)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = TEAL
        run.font.size = Pt(14)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def h3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(11.5)
        run.bold = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    return p

def body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_GRAY
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_GRAY
    return p

def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_GRAY
    return p

def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0]
    shade_row(hdr, HEADER_BG)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9.5)
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        if ri % 2 == 1:
            shade_row(row, RGBColor(0xF0, 0xF4, 0xF8))
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9.5)
            run.font.color.rgb = DARK_GRAY
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run("NLP OpenMed Sentiment Analysis")
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = NAVY

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Docker & API Guide")
run2.font.size = Pt(22)
run2.font.bold = True
run2.font.color.rgb = TEAL

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("For Junior Software Engineers")
run3.font.size = Pt(13)
run3.font.color.rgb = MID_GRAY
run3.font.italic = True

doc.add_paragraph()
doc.add_paragraph()

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
for label, val in [("Version", "1.0"), ("Date", "May 2026"), ("Environment", "Docker / Local API")]:
    r = p4.add_run(f"{label}:  ")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = NAVY
    r2 = p4.add_run(f"{val}     ")
    r2.font.size = Pt(10)
    r2.font.color.rgb = DARK_GRAY

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "Table of Contents")

toc_items = [
    ("1.", "Overview", "3"),
    ("2.", "Prerequisites", "4"),
    ("3.", "Understanding the Dockerfile", "5"),
    ("4.", "Building the Docker Image", "7"),
    ("5.", "Running the Container", "9"),
    ("6.", "Making Your First API Call", "11"),
    ("7.", "API Reference", "13"),
    ("8.", "Code Examples", "16"),
    ("9.", "Troubleshooting", "20"),
    ("10.", "Quick Reference Card", "22"),
]

for num, title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{num}  {title}")
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = NAVY

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "1. Overview")

body(doc, (
    "This guide walks you through two things: building a Docker image for the NLP OpenMed "
    "Sentiment Analysis application, and then using that running container to get results by "
    "calling its REST API."
))
body(doc, (
    "Docker packages the entire application — Python runtime, all 13 NLP models, dependencies, "
    "and configuration — into a single portable image. Once built, the image runs identically "
    "on any machine regardless of the operating system or what is installed locally. "
    "This means you can run the API on your laptop, a teammate's machine, or a cloud server "
    "without any setup differences."
))

add_callout(doc, (
    "You do not need to understand how Docker works internally to follow this guide. "
    "Every command is explained step by step. If something goes wrong, Section 9 "
    "(Troubleshooting) covers the most common problems."
), "note")

h2(doc, "1.1  What You Will Be Able to Do After This Guide")
for item in [
    "Build the NLP sentiment analysis Docker image from source.",
    "Run the containerised API on your local machine.",
    "Send text to the API and receive a sentiment prediction in return.",
    "Write a simple Python or cURL script that calls the API in a loop.",
    "Understand the API response fields and how to use them.",
    "Diagnose and fix the most common Docker and API errors.",
]:
    bullet(doc, item)

h2(doc, "1.2  Key Concepts (Plain English)")
add_table(doc,
    ["Term", "What it means in this context"],
    [
        ["Docker",         "A tool that packages an application and everything it needs into a self-contained box called an image."],
        ["Image",          "The packaged box — a snapshot of the application and all its dependencies. Think of it as a USB drive with everything pre-installed."],
        ["Container",      "A running instance of an image. One image can produce many containers, just like one USB drive can be plugged into many computers."],
        ["docker build",   "The command that reads the Dockerfile and creates the image from scratch."],
        ["docker run",     "The command that starts a container from an existing image."],
        ["Port mapping",   "A way to make a port inside the container accessible from your machine. -p 8000:8000 means: connect your machine's port 8000 to the container's port 8000."],
        ["REST API",       "A web service that accepts HTTP requests and returns JSON data. The sentiment analysis results come back as JSON."],
        ["JSON",           "A text format for structured data. Looks like: {\"sentiment\": \"POSITIVE\", \"confidence\": 0.93}."],
        ["Endpoint",       "A specific URL path on the API, like /api/v1/analyze. Each endpoint does something different."],
        [".env file",      "A plain text file that stores secret configuration values (like API keys) without putting them in your code."],
        ["Environment variable", "A named value passed to the container at startup. Equivalent to the values in your .env file."],
    ],
    col_widths=[3.5, 13]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. PREREQUISITES
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "2. Prerequisites")

body(doc, (
    "Before you start, make sure you have the following installed and ready. "
    "Each item below includes a command to verify it is working."
))

h2(doc, "2.1  Required Software")
add_table(doc,
    ["Software", "Minimum Version", "How to Verify", "Where to Get It"],
    [
        ["Docker Desktop", "24.x or later",
         "docker --version",
         "https://www.docker.com/products/docker-desktop"],
        ["Git", "2.x",
         "git --version",
         "https://git-scm.com/downloads"],
        ["A terminal", "—",
         "macOS: Terminal or iTerm2\nWindows: PowerShell or WSL2\nLinux: any shell",
         "Built into your OS"],
        ["curl (optional)", "Any",
         "curl --version",
         "Built into macOS/Linux. Windows: install via winget or use PowerShell Invoke-RestMethod"],
        ["Python 3.x (optional)", "3.8+",
         "python3 --version",
         "https://www.python.org/downloads — only needed to run the Python examples in Section 8"],
    ],
    col_widths=[3.5, 3, 5, 5]
)

add_callout(doc,
    "On Windows, use WSL2 (Windows Subsystem for Linux) or PowerShell. "
    "All commands in this guide are shown in bash/sh style. "
    "PowerShell users: replace single quotes ' with double quotes \" in JSON strings.",
    "note"
)

h2(doc, "2.2  Verify Docker is Running")
body(doc, "Open your terminal and run:")
add_code(doc, "docker --version\ndocker info", "bash")
body(doc, (
    "You should see a version number (e.g. Docker version 24.0.5) and a block of system "
    "information. If you see an error like 'Cannot connect to the Docker daemon', "
    "make sure Docker Desktop is open and running."
))

h2(doc, "2.3  The OpenMed PII API Key")
body(doc, (
    "The application calls an external service to redact patient names, dates, and identifiers "
    "from text before running sentiment analysis. You need an API key for that service. "
    "Ask your team lead for the value — it looks like a long random string."
))
body(doc, "Create a file called .env in the project root directory with these two lines:")
add_code(doc, """\
OPENMED_PII_BASE_URL=https://i6di6fkeqi.execute-api.ap-southeast-2.amazonaws.com/prod
OPENMED_PII_API_KEY=<your-openmed-api-key>""", "bash")

add_callout(doc,
    "Never share this file or commit it to Git. The .gitignore file already excludes .env. "
    "If you accidentally commit it, contact your team lead immediately to rotate the key.",
    "warning"
)

h2(doc, "2.4  Get the Project Code")
body(doc, "Clone the repository and navigate into it:")
add_code(doc, """\
git clone <repo-url>
cd nlp-openmed-sentiment-analysis-finetuned-models""", "bash")
body(doc, "Confirm the Dockerfile is present:")
add_code(doc, "ls Dockerfile", "bash")
body(doc, "You should see: Dockerfile")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 3. UNDERSTANDING THE DOCKERFILE
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "3. Understanding the Dockerfile")

body(doc, (
    "You do not need to edit the Dockerfile. This section explains what it does so you "
    "understand what happens during the build and why the first build takes longer than usual."
))

h2(doc, "3.1  What a Dockerfile Is")
body(doc, (
    "A Dockerfile is a plain text recipe that tells Docker how to build the image. "
    "It is a series of instructions executed in order, from top to bottom. "
    "Each instruction creates a new layer in the image — Docker caches these layers "
    "so that if you rebuild and only your code changed, Docker reuses all the earlier "
    "layers (dependencies, models) without re-downloading them."
))

h2(doc, "3.2  The Dockerfile Explained Line by Line")
body(doc, "The project Dockerfile has five main stages:")

add_table(doc,
    ["Stage", "What it does", "Why it matters"],
    [
        ["FROM python:3.11-slim",
         "Starts from an official Python 3.11 image based on a minimal Linux system.",
         "Gives us Python without including unnecessary system packages, keeping the image smaller."],
        ["ENV (environment variables)",
         "Sets default environment variables: HF_HOME (HuggingFace model cache location), "
         "NLTK_DATA (NLTK data location), MODE=api, PORT=8000.",
         "These control where models are stored and which mode the app starts in at runtime."],
        ["RUN apt-get install",
         "Installs Linux system libraries needed by the app (image processing, compilers).",
         "Without these, certain Python packages (e.g. Pillow, wordcloud) cannot be installed."],
        ["RUN pip install -r requirements.txt",
         "Installs all Python packages listed in requirements.txt.",
         "This is the longest step — it downloads PyTorch, Transformers, spaCy, FastAPI, etc."],
        ["RUN python -m spacy download",
         "Downloads the spaCy English language model (en_core_web_md).",
         "Used by the NLP preprocessing pipeline to clean and tag text."],
        ["RUN nltk.download (NLTK data)",
         "Downloads NLTK datasets: punkt, wordnet, averaged_perceptron_tagger.",
         "Used for tokenisation and lemmatisation during text preprocessing."],
        ["RUN pipeline(...) — HuggingFace models",
         "Pre-downloads the HuggingFace Transformer model weights for the three baseline models "
         "(DistilBERT SST-2, Twitter RoBERTa, GoEmotions).",
         "Without this, the first API call to each model would trigger a download and take "
         "30–90 seconds. Pre-downloading at build time means zero cold-start delays at runtime."],
        ["COPY src/ api/ ui/ ...",
         "Copies the application source code into the image.",
         "This is always the last COPY so that code changes don't invalidate the expensive "
         "dependency and model download layers above."],
        ["CMD",
         "Defines the default command to run when the container starts. "
         "If MODE=ui, launches the Gradio web interface. Otherwise, starts the FastAPI server.",
         "You can override this at runtime — this guide always uses the default MODE=api."],
    ],
    col_widths=[4, 6, 6.5]
)

h2(doc, "3.3  Build Cache — Why the Second Build is Fast")
body(doc, (
    "Docker caches each instruction layer. On the first build, everything is downloaded "
    "from scratch. On subsequent builds, Docker only re-runs instructions whose inputs "
    "have changed. Because the application code (COPY src/ ...) comes last in the Dockerfile, "
    "changing a Python file only re-runs the final COPY and CMD steps — all the expensive "
    "package and model download steps are served from cache and complete in seconds."
))

add_callout(doc,
    "If you only changed application code (e.g. api/routes.py), a rebuild takes under "
    "30 seconds because Docker reuses the cached layers for pip install, spaCy, NLTK, "
    "and HuggingFace models.",
    "tip"
)

h2(doc, "3.4  Runtime Modes")
body(doc, "The container supports two modes, controlled by the MODE environment variable:")
add_table(doc,
    ["MODE value", "What starts", "Port", "Use case"],
    [
        ["api (default)", "FastAPI REST API via uvicorn", "8000", "Programmatic access — your code calls the API."],
        ["ui",            "Gradio web interface",         "7860", "Browser-based interactive testing."],
    ],
    col_widths=[3, 4, 2, 7.5]
)
body(doc, "This guide focuses entirely on MODE=api (the default), which is what you will use to call the API programmatically.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. BUILDING THE DOCKER IMAGE
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "4. Building the Docker Image")

body(doc, (
    "Building the image converts the Dockerfile and the application code into a single "
    "portable image stored on your machine. You only need to do this once (or again after "
    "code changes)."
))

h2(doc, "4.1  The Build Command")
body(doc, "From the project root directory (where the Dockerfile lives), run:")
add_code(doc, "docker build -t nlp-sentiment .", "bash")

body(doc, "Breaking this down:")
add_table(doc,
    ["Part of command", "What it means"],
    [
        ["docker build", "Tells Docker to build a new image."],
        ["-t nlp-sentiment",
         "Names the image 'nlp-sentiment'. The -t flag stands for 'tag'. "
         "You use this name later to run the container."],
        [".",
         "The dot means 'use the current directory as the build context' — "
         "Docker will read the Dockerfile and all project files from here."],
    ],
    col_widths=[4, 12.5]
)

h2(doc, "4.2  What You Will See During the Build")
body(doc, (
    "The build prints each step as it runs. The first build will take "
    "10–30 minutes depending on your internet speed and machine speed. "
    "Here is what to expect:"
))

add_code(doc, """\
[+] Building 0.0s (2/18) STARTED
 => [internal] load build definition from Dockerfile                  0.0s
 => [internal] load .dockerignore                                      0.0s
 => [1/14] FROM python:3.11-slim                                       2.1s
 => [2/14] RUN apt-get update && apt-get install -y ...               15.3s
 => [3/14] RUN pip install --upgrade pip && pip install -r ...      4m 22s
 => [4/14] RUN python -m spacy download en_core_web_md               45.8s
 => [5/14] RUN mkdir -p /app/nltk_data && python - ...               12.4s
 => [6/14] RUN mkdir -p /app/.cache/huggingface && python - ...    3m 11s
 => [7/14] COPY src/ ./src/                                            0.1s
 ...
 => exporting to image                                                  8.2s
 => => naming to docker.io/library/nlp-sentiment                       0.0s""", "")

add_callout(doc,
    "The long steps are [3/14] pip install (downloading PyTorch and Transformers) and "
    "[6/14] HuggingFace model download. Both are normal and expected on the first build. "
    "Subsequent rebuilds use the Docker cache and skip these steps entirely.",
    "note"
)

h2(doc, "4.3  Confirming the Build Succeeded")
body(doc, "After the build completes, verify the image exists:")
add_code(doc, "docker images nlp-sentiment", "bash")
body(doc, "You should see a table like this:")
add_code(doc, """\
REPOSITORY      TAG       IMAGE ID       CREATED          SIZE
nlp-sentiment   latest    a1b2c3d4e5f6   2 minutes ago    6.8GB""", "")
body(doc, (
    "The image size of approximately 6–7 GB is expected — it includes PyTorch, "
    "all Transformer model weights, and the Python environment."
))

add_callout(doc,
    "If the build fails with an error message, note the step number (e.g. [3/14]) "
    "and the error text, then check Section 9 (Troubleshooting) for a solution.",
    "warning"
)

h2(doc, "4.4  Rebuilding After Code Changes")
body(doc, (
    "If you change application code (e.g. a file in api/ or src/), run the same "
    "build command again:"
))
add_code(doc, "docker build -t nlp-sentiment .", "bash")
body(doc, (
    "Docker will reuse all cached layers up to the COPY instruction and only "
    "re-copy the changed files. The rebuild typically completes in under 30 seconds."
))

add_callout(doc,
    "If you add or remove a Python package from requirements.txt, Docker will re-run "
    "the pip install step and all subsequent steps. This takes 5–15 minutes.",
    "note"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. RUNNING THE CONTAINER
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "5. Running the Container")

body(doc, (
    "Once the image is built, you start a container from it. The container runs the "
    "FastAPI server and exposes the API on port 8000."
))

h2(doc, "5.1  Basic Run Command")
add_code(doc, """\
docker run -p 8000:8000 \\
  -e OPENMED_PII_BASE_URL=https://i6di6fkeqi.execute-api.ap-southeast-2.amazonaws.com/prod \\
  -e OPENMED_PII_API_KEY=<your-openmed-api-key> \\
  nlp-sentiment""", "bash")

body(doc, "Breaking this down:")
add_table(doc,
    ["Part of command", "What it means"],
    [
        ["docker run",
         "Starts a new container."],
        ["-p 8000:8000",
         "Maps port 8000 on your machine to port 8000 inside the container. "
         "Without this, you cannot reach the API from your browser or code."],
        ["-e OPENMED_PII_BASE_URL=...",
         "Sets the OpenMed PII service URL as an environment variable inside the container. "
         "The -e flag stands for 'environment variable'."],
        ["-e OPENMED_PII_API_KEY=...",
         "Sets your OpenMed API key. Replace <your-openmed-api-key> with the actual key value."],
        ["nlp-sentiment",
         "The name of the image to run (the same name you used in docker build -t)."],
    ],
    col_widths=[4.5, 12]
)

add_callout(doc,
    "Replace <your-openmed-api-key> with your actual key. The key is a long string "
    "that looks like: sk-abc123xyz.... Ask your team lead for the correct value.",
    "warning"
)

h2(doc, "5.2  Using a .env File (Simpler)")
body(doc, (
    "Typing the API key in the command every time is error-prone. "
    "A cleaner approach is to load the values from your .env file:"
))
add_code(doc, """\
# Load values from your .env file and pass them to the container
docker run -p 8000:8000 \\
  --env-file .env \\
  nlp-sentiment""", "bash")

add_callout(doc,
    "The --env-file flag reads every KEY=VALUE line from .env and passes them all "
    "to the container automatically. This is the recommended approach — keep the key "
    "in .env and never type it on the command line.",
    "tip"
)

h2(doc, "5.3  What You Will See When the Container Starts")
body(doc, "If the container starts successfully, you will see output like this:")
add_code(doc, """\
INFO:     Started server process [1]
INFO:     Waiting for application startup.
INFO:     Application startup complete.
INFO:     Uvicorn running on http://0.0.0.0:8000 (Press CTRL+C to quit)""", "")
body(doc, (
    "The API is now live and accepting requests. Open a second terminal window to "
    "run the API calls in the following sections — do not close this terminal window "
    "or the container will stop."
))

h2(doc, "5.4  Running in the Background (Detached Mode)")
body(doc, (
    "If you want the container to run in the background so you can continue using "
    "the same terminal window, add the -d flag:"
))
add_code(doc, """\
docker run -d -p 8000:8000 --env-file .env --name nlp-api nlp-sentiment""", "bash")

add_table(doc,
    ["New flag", "What it does"],
    [
        ["-d",            "Runs the container in the background (detached mode). Returns the container ID immediately."],
        ["--name nlp-api","Gives the container a friendly name so you can reference it easily in later commands."],
    ],
    col_widths=[3, 13.5]
)

body(doc, "Useful commands when running in detached mode:")
add_code(doc, """\
# View live logs from the running container
docker logs -f nlp-api

# Stop the container
docker stop nlp-api

# Remove the stopped container
docker rm nlp-api

# Stop and remove in one step
docker rm -f nlp-api""", "bash")

h2(doc, "5.5  Running the Gradio UI Instead")
body(doc, (
    "If you want to test the application interactively in a browser rather than "
    "calling the API from code, run the container in UI mode:"
))
add_code(doc, """\
docker run -p 7860:7860 --env-file .env -e MODE=ui nlp-sentiment""", "bash")
body(doc, "Then open http://localhost:7860 in your browser.")

add_callout(doc,
    "The Gradio UI is useful for quick manual tests. For automated or programmatic "
    "use — writing scripts, integrating with other systems — use the API (MODE=api).",
    "note"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 6. MAKING YOUR FIRST API CALL
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "6. Making Your First API Call")

body(doc, (
    "With the container running (from Section 5), you are ready to call the API. "
    "Follow these steps in order."
))

h2(doc, "Step 1 — Check the API is Alive")
body(doc, "Open a new terminal and run:")
add_code(doc, """\
curl -s http://localhost:8000/api/v1/health | python3 -m json.tool""", "bash")
body(doc, "Expected response:")
add_code(doc, """\
{
    "status": "ok",
    "models_loaded": []
}""", "json")
body(doc, (
    "status: \"ok\" confirms the API is running. "
    "models_loaded is empty on a fresh start — models are loaded the first time "
    "they are used."
))

h2(doc, "Step 2 — Warm Up All 13 Models")
body(doc, (
    "The first call to any model triggers a one-time load from disk, which can take "
    "10–30 seconds. To avoid this delay on your first real request, pre-load all models now:"
))
add_code(doc, """\
curl -s http://localhost:8000/api/v1/warmup | python3 -m json.tool""", "bash")
body(doc, "Expected response (abbreviated):")
add_code(doc, """\
{
    "status": "ok",
    "models": {
        "bert_hc_v2":          "loaded in 4.2s",
        "distilroberta_hc_v2": "loaded in 3.1s",
        "distilbert_hc_v2":    "loaded in 2.8s",
        "bert_hc":             "loaded in 4.0s",
        "distilbert_hc":       "loaded in 2.9s",
        "distilroberta_hc":    "loaded in 3.0s",
        "default":             "loaded in 2.5s",
        "roberta":             "loaded in 4.1s",
        "emotion":             "loaded in 3.3s",
        "amazon":              "loaded in 2.6s",
        "twitter":             "loaded in 3.5s",
        "sst2":                "loaded in 4.0s",
        "zeroshot":            "loaded in 8.7s"
    }
}""", "json")

add_callout(doc,
    "You only need to call /warmup once after the container starts. "
    "All 13 models stay loaded in memory until the container is stopped.",
    "tip"
)

h2(doc, "Step 3 — Analyse Your First Text")
body(doc, "Send a piece of text for sentiment analysis using the default model (bert_hc_v2):")
add_code(doc, """\
curl -s -X POST http://localhost:8000/api/v1/analyze \\
  -H "Content-Type: application/json" \\
  -d '{
    "text": "The nurse was incredibly kind and made me feel completely at ease."
  }' | python3 -m json.tool""", "bash")

body(doc, "You will receive a response like this:")
add_code(doc, """\
{
    "sentiment":       "POSITIVE",
    "probabilities":   [0.02, 0.05, 0.93],
    "labels":          ["NEGATIVE", "NEUTRAL", "POSITIVE"],
    "model_type":      "bert_hc_v2",
    "redacted_text":   "The nurse was incredibly kind and made me feel completely at ease.",
    "cleaned_text":    "nurse incredibly kind made feel completely ease",
    "tokenized_text":  ["nurse", "incredibly", "kind", "made", "feel", "completely", "ease"],
    "lemmatized_text": ["nurse", "incredibly", "kind", "make", "feel", "completely", "ease"],
    "ner":             [],
    "word_distribution": {
        "NEGATIVE": 0,
        "NEUTRAL":  2,
        "POSITIVE": 5
    },
    "total_words": 7
}""", "json")

h2(doc, "6.1  Reading the Response")
body(doc, "Here is what the important fields mean:")
add_table(doc,
    ["Field", "Example value", "What it means"],
    [
        ["sentiment",       "\"POSITIVE\"",
         "The model's top prediction. This is the main result you will use."],
        ["probabilities",   "[0.02, 0.05, 0.93]",
         "Confidence scores for each label. Aligned with the labels field. Sum to 1.0."],
        ["labels",          "[\"NEGATIVE\", \"NEUTRAL\", \"POSITIVE\"]",
         "The label names for this model, in the same order as probabilities."],
        ["model_type",      "\"bert_hc_v2\"",
         "Which model was used (echoed back from your request)."],
        ["redacted_text",   "\"The nurse was...\"",
         "The text after PII/PHI removal. This is what the model actually analysed."],
        ["cleaned_text",    "\"nurse incredibly kind...\"",
         "After stop-word removal and punctuation cleaning."],
        ["lemmatized_text", "[\"nurse\", \"make\", ...]",
         "Tokens reduced to their dictionary root form ('made' → 'make')."],
        ["word_distribution","{ \"POSITIVE\": 5, ... }",
         "How many words lean positive, neutral, or negative individually."],
        ["total_words",     "7",
         "Number of tokens after preprocessing."],
    ],
    col_widths=[3.5, 3.5, 9.5]
)

add_callout(doc,
    "The most important fields for most use cases are: sentiment (the predicted label), "
    "probabilities (the confidence scores), and redacted_text (to confirm PII was removed "
    "before storage or logging).",
    "tip"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 7. API REFERENCE
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "7. API Reference")

h2(doc, "7.1  Base URL")
add_code(doc, "http://localhost:8000/api/v1", "")
body(doc, "All three endpoints are accessed at paths relative to this base URL.")

h2(doc, "7.2  Authentication")
body(doc, (
    "No API key is required to call the local API. All three endpoints are open "
    "when the container is running locally. The only credential you need is the "
    "OpenMed PII key — but that is configured as an environment variable at container "
    "startup and used server-side. You never include it in your API calls."
))

add_callout(doc,
    "The API should only be exposed on your local machine during development. "
    "Do not make port 8000 publicly accessible without adding authentication at "
    "a reverse proxy (e.g. nginx, AWS API Gateway).",
    "warning"
)

h2(doc, "7.3  Endpoints Summary")
add_table(doc,
    ["Method", "Path", "Description", "When to Use"],
    [
        ["GET",  "/api/v1/health",
         "Returns API status and list of loaded models.",
         "To verify the API is live. Run this first after the container starts."],
        ["GET",  "/api/v1/warmup",
         "Pre-loads all 13 models into memory.",
         "Run once after each container start to eliminate first-request delays."],
        ["POST", "/api/v1/analyze",
         "Redacts PII, preprocesses text, and runs sentiment analysis.",
         "Every time you want to analyse a piece of text."],
    ],
    col_widths=[2, 4, 6, 4.5]
)

h2(doc, "7.4  POST /api/v1/analyze — Request")
h3(doc, "Request body fields")
add_table(doc,
    ["Field", "Type", "Required?", "Default", "Description"],
    [
        ["text",
         "string", "Yes", "—",
         "The text to analyse. Must be between 4 and 300 words (counted after stripping whitespace)."],
        ["model_type",
         "string", "No", "\"bert_hc_v2\"",
         "Which of the 13 models to use. See Section 7.5 for all valid values."],
    ],
    col_widths=[3, 2.5, 2.5, 3, 5.5]
)

h3(doc, "Example request (cURL)")
add_code(doc, """\
curl -s -X POST http://localhost:8000/api/v1/analyze \\
  -H "Content-Type: application/json" \\
  -d '{
    "text": "I waited over three hours with no update on my appointment.",
    "model_type": "bert_hc_v2"
  }' | python3 -m json.tool""", "bash")

h2(doc, "7.5  Available Models")
body(doc, (
    "The model_type field accepts one of 13 values. The six fine-tuned healthcare "
    "models (prefix bert_hc, distilroberta_hc, distilbert_hc) are the best choice "
    "for clinical and patient feedback text."
))

h3(doc, "Fine-Tuned Healthcare Models (recommended for clinical text)")
add_table(doc,
    ["model_type value", "Labels returned", "Notes"],
    [
        ["bert_hc_v2",          "NEGATIVE / NEUTRAL / POSITIVE", "Best accuracy. Production default. Use this unless you have a specific reason not to."],
        ["distilroberta_hc_v2", "NEGATIVE / NEUTRAL / POSITIVE", "Fast and accurate. Good alternative to bert_hc_v2 at high volume."],
        ["distilbert_hc_v2",    "NEGATIVE / NEUTRAL / POSITIVE", "Smallest and fastest fine-tuned model."],
        ["bert_hc",             "NEGATIVE / NEUTRAL / POSITIVE", "v1 baseline. Use for benchmarking against v2."],
        ["distilbert_hc",       "NEGATIVE / NEUTRAL / POSITIVE", "v1 lightweight baseline."],
        ["distilroberta_hc",    "NEGATIVE / NEUTRAL / POSITIVE", "v1 RoBERTa baseline."],
    ],
    col_widths=[4, 4, 8.5]
)

h3(doc, "General-Purpose Pretrained Models")
add_table(doc,
    ["model_type value", "Labels returned", "Notes"],
    [
        ["default",  "POSITIVE / NEGATIVE",                                         "DistilBERT SST-2. Fast binary baseline. No NEUTRAL class."],
        ["roberta",  "1 STAR / 2 STARS / 3 STARS / 4 STARS / 5 STARS",            "BERT Multilingual. Returns star-rating style output."],
        ["emotion",  "ANGER / DISGUST / FEAR / JOY / NEUTRAL / SADNESS / SURPRISE","GoEmotions. Use when detecting specific emotions, not just polarity."],
        ["amazon",   "POSITIVE / NEGATIVE",                                         "Amazon Reviews BERT. Review-style feedback. No NEUTRAL class."],
        ["twitter",  "NEGATIVE / NEUTRAL / POSITIVE",                               "RoBERTa trained on tweets. Handles informal language."],
        ["sst2",     "POSITIVE / NEGATIVE",                                         "BERT SST-2. Standard benchmark baseline. No NEUTRAL class."],
        ["zeroshot", "POSITIVE / NEGATIVE / NEUTRAL",                               "BART MNLI. Most flexible but slowest model."],
    ],
    col_widths=[2.8, 5, 8.7]
)

add_callout(doc,
    "When in doubt, use bert_hc_v2. It is the recommended default for healthcare text "
    "and handles positive, neutral, and negative sentiments.",
    "tip"
)

h2(doc, "7.6  Error Responses")
add_table(doc,
    ["HTTP Code", "Meaning", "Common Cause", "How to Fix"],
    [
        ["200", "OK",                   "—",
         "No action needed. Check the sentiment field in the response."],
        ["422", "Unprocessable Entity", "text is under 4 words, over 300 words, or model_type is not a valid value.",
         "Check word count. Verify model_type is spelled exactly as shown in Section 7.5."],
        ["400", "Bad Request",          "The request body is not valid JSON.",
         "Check for missing quotes, commas, or unmatched brackets in your JSON."],
        ["500", "Internal Server Error","OpenMed PII service unreachable, or a model failed to load.",
         "Check container logs (docker logs nlp-api). Verify OPENMED_PII_API_KEY is correct."],
        ["503", "Service Unavailable",  "Container is not running.",
         "Run docker run ... to start the container again."],
    ],
    col_widths=[2.2, 3.5, 5.5, 5.3]
)

h2(doc, "7.7  Interactive API Documentation")
body(doc, (
    "FastAPI automatically generates an interactive browser-based API explorer. "
    "While the container is running, open this URL in your browser:"
))
add_code(doc, "http://localhost:8000/docs", "")
body(doc, (
    "From there you can try every endpoint without writing any code: "
    "click an endpoint, click 'Try it out', fill in the fields, and click 'Execute'. "
    "This is the fastest way to test a request manually."
))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 8. CODE EXAMPLES
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "8. Code Examples")

body(doc, (
    "All examples below assume the container is running and the API is reachable at "
    "http://localhost:8000/api/v1. Start the container first (Section 5) before running these."
))

h2(doc, "8.1  cURL (Command Line)")
body(doc, "No installation needed — cURL is built into macOS and Linux.")

h3(doc, "Health check")
add_code(doc, """\
curl -s http://localhost:8000/api/v1/health | python3 -m json.tool""", "bash")

h3(doc, "Warm up all models")
add_code(doc, """\
curl -s http://localhost:8000/api/v1/warmup | python3 -m json.tool""", "bash")

h3(doc, "Analyse a positive text")
add_code(doc, """\
curl -s -X POST http://localhost:8000/api/v1/analyze \\
  -H "Content-Type: application/json" \\
  -d '{
    "text": "The specialist was outstanding and I felt fully supported throughout."
  }' | python3 -m json.tool""", "bash")

h3(doc, "Analyse with a specific model")
add_code(doc, """\
curl -s -X POST http://localhost:8000/api/v1/analyze \\
  -H "Content-Type: application/json" \\
  -d '{
    "text": "I am so relieved — the test results came back clear.",
    "model_type": "emotion"
  }' | python3 -m json.tool""", "bash")

h3(doc, "Extract just the sentiment label from the response")
add_code(doc, """\
curl -s -X POST http://localhost:8000/api/v1/analyze \\
  -H "Content-Type: application/json" \\
  -d '{"text": "The ward was clean and the staff were attentive."}' \\
  | python3 -c "import sys,json; print(json.load(sys.stdin)['sentiment'])" """, "bash")

h2(doc, "8.2  Python")
body(doc, "Install the requests library if you haven't already:")
add_code(doc, "pip install requests", "bash")

h3(doc, "Single text — basic")
add_code(doc, """\
import requests

BASE_URL = "http://localhost:8000/api/v1"
HEADERS  = {"Content-Type": "application/json"}

# Health check
resp = requests.get(f"{BASE_URL}/health")
print("API status:", resp.json()["status"])

# Analyse a text
payload = {
    "text":       "The doctor explained my diagnosis clearly and I felt reassured.",
    "model_type": "bert_hc_v2",
}
resp = requests.post(f"{BASE_URL}/analyze", headers=HEADERS, json=payload)
resp.raise_for_status()

data = resp.json()
print(f"Sentiment:     {data['sentiment']}")
print(f"Confidence:    {max(data['probabilities']):.1%}")
print(f"Redacted text: {data['redacted_text']}")""", "python")

h3(doc, "Analysing a list of texts in a loop")
add_code(doc, """\
import requests

BASE_URL = "http://localhost:8000/api/v1"
HEADERS  = {"Content-Type": "application/json"}
MODEL    = "bert_hc_v2"

feedback_list = [
    "Dr Smith was thorough and answered all my questions.",
    "I waited two hours past my appointment time with no explanation.",
    "The discharge instructions were clear and easy to follow.",
    "The nurse called me promptly with my test results.",
    "I could not get through on the phone to reschedule.",
]

for text in feedback_list:
    resp = requests.post(
        f"{BASE_URL}/analyze",
        headers=HEADERS,
        json={"text": text, "model_type": MODEL},
    )
    resp.raise_for_status()
    result = resp.json()
    label      = result["sentiment"]
    confidence = max(result["probabilities"])
    print(f"{label:10s}  ({confidence:.0%})  |  {text[:60]}")""", "python")

h3(doc, "Saving results to a CSV file")
add_code(doc, """\
import requests
import csv

BASE_URL = "http://localhost:8000/api/v1"
HEADERS  = {"Content-Type": "application/json"}

feedback_list = [
    "The nursing team were warm, professional, and incredibly caring.",
    "I waited four hours with no update on my appointment.",
    "My discharge letter explained everything very clearly.",
]

with open("sentiment_results.csv", "w", newline="") as f:
    writer = csv.writer(f)
    writer.writerow(["original_text", "redacted_text", "sentiment", "confidence"])

    for text in feedback_list:
        resp = requests.post(
            f"{BASE_URL}/analyze",
            headers=HEADERS,
            json={"text": text, "model_type": "bert_hc_v2"},
        )
        resp.raise_for_status()
        r = resp.json()
        writer.writerow([
            text,
            r["redacted_text"],
            r["sentiment"],
            round(max(r["probabilities"]), 4),
        ])

print("Saved to sentiment_results.csv")""", "python")

h3(doc, "Robust wrapper with error handling")
add_code(doc, """\
import requests

BASE_URL = "http://localhost:8000/api/v1"
HEADERS  = {"Content-Type": "application/json"}

def analyse(text: str, model_type: str = "bert_hc_v2") -> dict:
    \"\"\"Call the sentiment API and return the parsed JSON response.\"\"\"
    try:
        resp = requests.post(
            f"{BASE_URL}/analyze",
            headers=HEADERS,
            json={"text": text, "model_type": model_type},
            timeout=60,
        )
        resp.raise_for_status()
        return resp.json()

    except requests.exceptions.ConnectionError:
        print("Cannot connect — is the Docker container running on port 8000?")
        raise
    except requests.exceptions.Timeout:
        print("Request timed out — model may still be loading. Try again.")
        raise
    except requests.exceptions.HTTPError as e:
        if e.response.status_code == 422:
            print("Invalid request:", e.response.json().get("detail"))
        elif e.response.status_code >= 500:
            print("Server error — check the container logs: docker logs nlp-api")
        raise

# Example usage
result = analyse("The consultation was excellent and I felt listened to.")
print(result["sentiment"], result["redacted_text"])""", "python")

h2(doc, "8.3  JavaScript / Node.js")
body(doc, "Using the built-in fetch API (Node.js 18+ or any modern browser).")

h3(doc, "Single text")
add_code(doc, """\
const BASE_URL = "http://localhost:8000/api/v1";

async function analyseSentiment(text, modelType = "bert_hc_v2") {
  const response = await fetch(`${BASE_URL}/analyze`, {
    method:  "POST",
    headers: { "Content-Type": "application/json" },
    body: JSON.stringify({ text, model_type: modelType }),
  });

  if (!response.ok) {
    const err = await response.json();
    throw new Error(`API error ${response.status}: ${JSON.stringify(err)}`);
  }

  return response.json();
}

// Usage
analyseSentiment("The nurse was incredibly supportive during my stay.", "bert_hc_v2")
  .then(data => {
    console.log("Sentiment:    ", data.sentiment);
    console.log("Confidence:   ", Math.max(...data.probabilities).toFixed(2));
    console.log("Redacted text:", data.redacted_text);
  })
  .catch(err => console.error("Error:", err.message));""", "javascript")

h2(doc, "8.4  Comparing Multiple Models on the Same Text")
body(doc, "Use this Python snippet to run all 13 models and compare their outputs side by side:")
add_code(doc, """\
import requests

BASE_URL   = "http://localhost:8000/api/v1"
HEADERS    = {"Content-Type": "application/json"}
ALL_MODELS = [
    "bert_hc_v2", "distilroberta_hc_v2", "distilbert_hc_v2",
    "bert_hc",    "distilbert_hc",        "distilroberta_hc",
    "default",    "roberta",              "emotion",
    "amazon",     "twitter",              "sst2",    "zeroshot",
]

text = "The specialist was knowledgeable and I felt confident in the diagnosis."
print(f"Text: {text}\\n")
print(f"{'Model':<25} {'Sentiment':<15} {'Confidence'}")
print("-" * 55)

for model in ALL_MODELS:
    resp = requests.post(
        f"{BASE_URL}/analyze",
        headers=HEADERS,
        json={"text": text, "model_type": model},
        timeout=60,
    )
    if resp.status_code == 200:
        r = resp.json()
        print(f"{model:<25} {r['sentiment']:<15} {max(r['probabilities']):.1%}")
    else:
        print(f"{model:<25} ERROR {resp.status_code}")""", "python")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 9. TROUBLESHOOTING
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "9. Troubleshooting")

body(doc, (
    "This section covers the most common problems you may encounter when building "
    "the image or calling the API. Find your error message below for a step-by-step fix."
))

h2(doc, "9.1  Build Errors")
add_table(doc,
    ["Error / Symptom", "Cause", "Fix"],
    [
        ["Cannot connect to the Docker daemon",
         "Docker Desktop is not running.",
         "Open Docker Desktop from your Applications folder (macOS) or Start menu (Windows) "
         "and wait for it to fully start before re-running the build."],
        ["ERROR [3/14] pip install — could not find a version that satisfies...",
         "A package name in requirements.txt is misspelled or does not exist on PyPI.",
         "Check requirements.txt for typos. Run the failing pip install command locally "
         "to see the full error message."],
        ["ERROR — no space left on device",
         "Your machine's disk is full. A 7 GB image requires substantial free space.",
         "Free disk space. Run docker system prune to remove unused images and containers. "
         "Docker Desktop → Settings → Resources → Disk image size may also need increasing."],
        ["COPY failed: file not found in build context",
         "A file listed in the Dockerfile's COPY instruction doesn't exist in the project.",
         "Make sure you are running docker build from the project root directory "
         "and that all required files are present."],
        ["Build hangs at RUN pip install for more than 30 minutes",
         "Slow or intermittent internet connection during package download.",
         "Wait — very large packages (PyTorch ~2 GB) can take a long time on slow connections. "
         "If it doesn't progress after 60 minutes, press Ctrl+C and retry."],
    ],
    col_widths=[4.5, 4, 8]
)

h2(doc, "9.2  Container Startup Errors")
add_table(doc,
    ["Error / Symptom", "Cause", "Fix"],
    [
        ["docker: Error response from daemon: driver failed programming external connectivity "
         "— Bind for 0.0.0.0:8000 failed: port is already allocated",
         "Something else is already using port 8000 on your machine.",
         "Either stop the other process using port 8000, or map to a different host port: "
         "docker run -p 8080:8000 ... then access the API at localhost:8080."],
        ["Container exits immediately with no output",
         "A startup error (e.g. missing file, bad Python syntax) caused the process to crash.",
         "Run without -d to see the output: docker run -p 8000:8000 --env-file .env nlp-sentiment. "
         "Or view logs: docker logs nlp-api."],
        ["Application startup complete never appears",
         "Server is still starting (may take 10–20 seconds) or crashed silently.",
         "Wait 30 seconds. If it still doesn't appear, check logs: docker logs nlp-api."],
    ],
    col_widths=[5.5, 3.5, 7.5]
)

h2(doc, "9.3  API Call Errors")
add_table(doc,
    ["Error / Symptom", "Cause", "Fix"],
    [
        ["curl: (7) Failed to connect to localhost port 8000",
         "The container is not running or is starting up.",
         "Verify the container is running: docker ps. "
         "If it's not listed, start it again with docker run ...."],
        ["HTTP 500 on every /analyze request",
         "OPENMED_PII_API_KEY is missing or incorrect.",
         "Check your .env file has the correct key. "
         "Restart the container after fixing .env. "
         "Check logs: docker logs nlp-api | grep -i error"],
        ["HTTP 422: Text must be at least 4 words",
         "The submitted text has fewer than 4 words after whitespace stripping.",
         "Ensure the text field contains at least 4 meaningful words."],
        ["HTTP 422: model_type must be one of ...",
         "The model_type value is misspelled or not in the supported list.",
         "Check Section 7.5 for the exact string values. They are case-sensitive."],
        ["First /analyze call takes 20–30 seconds",
         "The model is being loaded from disk for the first time.",
         "This is normal. Call GET /api/v1/warmup after each container start "
         "to pre-load all 13 models."],
        ["redacted_text is identical to the original",
         "No PII was detected in the text (this is normal for non-PII text), "
         "or the OpenMed service is unreachable.",
         "Check that OPENMED_PII_BASE_URL is correct in .env. "
         "Test with text that contains a clear name: e.g. 'Dr Jane Smith was helpful'."],
        ["Python: requests.exceptions.ConnectionError",
         "The container is not running or the port mapping is wrong.",
         "Check docker ps. Verify you used -p 8000:8000 in the docker run command."],
    ],
    col_widths=[4.5, 4, 8]
)

h2(doc, "9.4  Useful Diagnostic Commands")
add_code(doc, """\
# List all running containers
docker ps

# List all containers (including stopped ones)
docker ps -a

# View live logs from the container
docker logs -f nlp-api

# View the last 50 lines of logs
docker logs --tail 50 nlp-api

# Check container resource usage (CPU, memory)
docker stats nlp-api

# List all local Docker images
docker images

# Remove stopped containers and unused images to free disk space
docker system prune""", "bash")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 10. QUICK REFERENCE CARD
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "10. Quick Reference Card")

body(doc, "Keep this page handy as a cheat sheet for daily use.")

h2(doc, "Build the Image")
add_code(doc, """\
# One-time build (run from the project root directory)
docker build -t nlp-sentiment .""", "bash")

h2(doc, "Start the API Container")
add_code(doc, """\
# Foreground (you see all logs in this terminal)
docker run -p 8000:8000 --env-file .env nlp-sentiment

# Background (returns immediately; use logs to monitor)
docker run -d -p 8000:8000 --env-file .env --name nlp-api nlp-sentiment""", "bash")

h2(doc, "Container Management")
add_code(doc, """\
docker ps                   # list running containers
docker logs -f nlp-api      # stream live logs
docker stop nlp-api         # stop the container
docker rm -f nlp-api        # stop and remove""", "bash")

h2(doc, "API Calls")
add_code(doc, """\
# Health check
curl -s http://localhost:8000/api/v1/health | python3 -m json.tool

# Warm up all 13 models (run once per container start)
curl -s http://localhost:8000/api/v1/warmup | python3 -m json.tool

# Analyse text — default model (bert_hc_v2)
curl -s -X POST http://localhost:8000/api/v1/analyze \\
  -H "Content-Type: application/json" \\
  -d '{"text": "YOUR TEXT HERE"}' | python3 -m json.tool

# Analyse text — specific model
curl -s -X POST http://localhost:8000/api/v1/analyze \\
  -H "Content-Type: application/json" \\
  -d '{"text": "YOUR TEXT HERE", "model_type": "emotion"}' | python3 -m json.tool""", "bash")

h2(doc, "Minimal Python")
add_code(doc, """\
import requests
resp = requests.post(
    "http://localhost:8000/api/v1/analyze",
    json={"text": "YOUR TEXT HERE", "model_type": "bert_hc_v2"},
)
resp.raise_for_status()
print(resp.json()["sentiment"])          # e.g. POSITIVE
print(resp.json()["redacted_text"])      # PII-redacted version""", "python")

h2(doc, "Endpoints")
add_table(doc,
    ["Method", "Path", "Use for"],
    [
        ["GET",  "/api/v1/health",  "Confirm API is alive."],
        ["GET",  "/api/v1/warmup",  "Pre-load all 13 models (run once after startup)."],
        ["POST", "/api/v1/analyze", "Analyse sentiment of a text using a chosen model."],
    ],
    col_widths=[2, 4.5, 10]
)

h2(doc, "Model Quick Reference")
add_table(doc,
    ["Key", "Labels", "Best for"],
    [
        ["bert_hc_v2",          "NEGATIVE / NEUTRAL / POSITIVE",                              "Clinical notes — production default."],
        ["distilroberta_hc_v2", "NEGATIVE / NEUTRAL / POSITIVE",                              "High-volume, fast + accurate."],
        ["distilbert_hc_v2",    "NEGATIVE / NEUTRAL / POSITIVE",                              "Fastest HC fine-tuned model."],
        ["emotion",             "ANGER / DISGUST / FEAR / JOY / NEUTRAL / SADNESS / SURPRISE","Detecting patient emotions."],
        ["roberta",             "1 STAR / 2 STARS / 3 STARS / 4 STARS / 5 STARS",            "Star-rating style reporting."],
        ["twitter",             "NEGATIVE / NEUTRAL / POSITIVE",                              "Informal or short-form feedback."],
        ["zeroshot",            "POSITIVE / NEGATIVE / NEUTRAL",                              "Zero-shot — no training needed."],
    ],
    col_widths=[4, 5.5, 7]
)

h2(doc, "HTTP Status Codes")
add_table(doc,
    ["Code", "Meaning", "Fix"],
    [
        ["200", "Success",                  "No action needed."],
        ["422", "Invalid request data",     "Check text is 4–300 words. Verify model_type is a valid key."],
        ["500", "Server / PII error",       "Check docker logs nlp-api. Verify .env has the correct OpenMed key."],
        ["503", "Container not running",    "Start the container: docker run -p 8000:8000 --env-file .env nlp-sentiment"],
    ],
    col_widths=[1.8, 4, 10.7]
)

h2(doc, "Interactive Docs")
body(doc, "While the container is running, open this in your browser:")
add_code(doc, "http://localhost:8000/docs", "")
body(doc, "Try every endpoint directly from your browser — no code needed.")

# ── Save ──────────────────────────────────────────────────────────────────────
import os as _os
_os.makedirs("docs", exist_ok=True)
output_path = "docs/NLP_OpenMed_Docker_API_Guide.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
